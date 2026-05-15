# -*- coding: utf-8 -*-
"""
임진강김치 MES 시스템 구축 - 아키텍처설계서 생성 스크립트
실행: python generate_architecture_doc.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"C:\gerardo\01 SmallSF\Imjingang-Kimchi\산출물\설계단계\05_아키텍처설계서.docx"

# ── 색상 상수 ──────────────────────────────────────────────
COLOR_TITLE_BG   = RGBColor(0x1F, 0x49, 0x7D)   # 진한 네이비
COLOR_H1_BG      = RGBColor(0x2E, 0x74, 0xB5)   # 파랑
COLOR_H2_BG      = RGBColor(0xD6, 0xE4, 0xF0)   # 연한 파랑
COLOR_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_BLACK      = RGBColor(0x00, 0x00, 0x00)
COLOR_TH_BG      = RGBColor(0x2E, 0x74, 0xB5)
COLOR_ALT_ROW    = RGBColor(0xEA, 0xF2, 0xFB)
COLOR_DIAGRAM_BG = RGBColor(0xF2, 0xF2, 0xF2)
COLOR_ACCENT     = RGBColor(0xC0, 0x00, 0x00)   # 강조 빨강


# ── 헬퍼 함수 ──────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(rgb)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """셀 테두리 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_para_border_bottom(para, color="2E74B5", sz=12):
    """단락 아래쪽 테두리"""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_run_with_font(para, text, bold=False, italic=False,
                      size_pt=11, color=None, font_name="맑은 고딕"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color
    return run


def heading1(doc, text, number):
    """1수준 제목"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # 배경 효과를 위해 테이블 1x1 사용
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_H1_BG)
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    add_run_with_font(p, f"{number}. {text}", bold=True, size_pt=14, color=COLOR_WHITE)
    doc.add_paragraph()  # 여백
    return tbl


def heading2(doc, text):
    """2수준 제목"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    add_run_with_font(para, text, bold=True, size_pt=12, color=COLOR_H1_BG)
    set_para_border_bottom(para, color="2E74B5", sz=8)
    return para


def body_para(doc, text, indent=False):
    """본문 단락"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    if indent:
        para.paragraph_format.left_indent = Cm(0.5)
    add_run_with_font(para, text, size_pt=10.5)
    return para


def bullet_para(doc, text, level=1):
    """불릿 단락"""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Cm(level * 0.5)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    add_run_with_font(para, text, size_pt=10)
    return para


def add_table(doc, headers, rows_data, col_widths=None, alt_row=True):
    """스타일 테이블 생성"""
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, COLOR_TH_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_font(p, h, bold=True, size_pt=10, color=COLOR_WHITE)

    # 데이터 행
    for r_idx, row_data in enumerate(rows_data):
        row = tbl.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            if alt_row and r_idx % 2 == 1:
                set_cell_bg(cell, COLOR_ALT_ROW)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run_with_font(p, str(cell_text), size_pt=9.5)

    # 컬럼 너비
    if col_widths:
        for i, row in enumerate(tbl.rows):
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Cm(col_widths[j])

    doc.add_paragraph()
    return tbl


def add_diagram_box(doc, lines, title=None):
    """다이어그램 텍스트 박스 (단일 테이블)"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_DIAGRAM_BG)
    cell.width = Cm(15.5)

    if title:
        p0 = cell.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_font(p0, title, bold=True, size_pt=10, color=COLOR_H1_BG)
        cell.add_paragraph()  # 빈 줄

    for line in lines:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run_with_font(p, line, size_pt=9, font_name="Courier New")

    doc.add_paragraph()
    return tbl


# ══════════════════════════════════════════════════════════
#  문서 생성 시작
# ══════════════════════════════════════════════════════════
doc = Document()

# ── 페이지 여백 설정 ──
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ══════════════════════════════════════════════════════════
#  표지
# ══════════════════════════════════════════════════════════
# 표지 배경 테이블
cover_tbl = doc.add_table(rows=1, cols=1)
cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_cell = cover_tbl.cell(0, 0)
set_cell_bg(cover_cell, COLOR_TITLE_BG)
cover_cell.width = Cm(16)

for _ in range(5):
    cover_cell.add_paragraph()

p_title = cover_cell.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run_with_font(p_title, "아키텍처설계서", bold=True, size_pt=28, color=COLOR_WHITE)

cover_cell.add_paragraph()

p_sub = cover_cell.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run_with_font(p_sub, "임진강김치 MES 시스템 구축", bold=True, size_pt=16, color=RGBColor(0xBD, 0xD7, 0xEE))

cover_cell.add_paragraph()

p_line = cover_cell.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run_with_font(p_line, "─" * 40, size_pt=10, color=RGBColor(0x70, 0xA0, 0xD0))

cover_cell.add_paragraph()

info_lines = [
    ("문서번호", "IK-MES-ARCH-001"),
    ("작성일",   "2026-05-12"),
    ("버전",     "V1.0"),
    ("작성기관", "임진강김치 MES 구축 프로젝트팀"),
    ("보안등급", "대외비"),
]
for label, value in info_lines:
    pi = cover_cell.add_paragraph()
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_with_font(pi, f"{label}:  ", bold=True, size_pt=11, color=RGBColor(0xBD, 0xD7, 0xEE))
    add_run_with_font(pi, value, size_pt=11, color=COLOR_WHITE)

for _ in range(5):
    cover_cell.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  목차 페이지
# ══════════════════════════════════════════════════════════
p_toc_title = doc.add_paragraph()
p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run_with_font(p_toc_title, "목  차", bold=True, size_pt=16, color=COLOR_H1_BG)
set_para_border_bottom(p_toc_title, color="2E74B5", sz=12)

doc.add_paragraph()

toc_items = [
    ("1.", "개요", "3"),
    ("  1.1", "문서 목적", "3"),
    ("  1.2", "아키텍처 설계 원칙", "3"),
    ("  1.3", "시스템 범위", "3"),
    ("2.", "전체 시스템 아키텍처", "4"),
    ("  2.1", "5계층 아키텍처 개요", "4"),
    ("  2.2", "계층별 역할과 책임", "5"),
    ("3.", "HW 아키텍처", "6"),
    ("  3.1", "현장 설비 및 센서 구성", "6"),
    ("  3.2", "네트워크 구성", "7"),
    ("  3.3", "서버 구성", "7"),
    ("4.", "SW 아키텍처", "8"),
    ("  4.1", "Frontend 아키텍처", "8"),
    ("  4.2", "Backend 아키텍처", "9"),
    ("  4.3", "데이터 아키텍처", "10"),
    ("  4.4", "AI Agent 아키텍처", "11"),
    ("5.", "보안 아키텍처", "12"),
    ("6.", "배포 아키텍처", "13"),
    ("7.", "시스템 연계 아키텍처", "14"),
    ("8.", "성능 및 가용성 설계", "15"),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    is_main = not num.startswith("  ")
    add_run_with_font(p, f"{num}  {title}", bold=is_main, size_pt=10.5 if is_main else 10)
    tab_run = p.add_run()
    tab_run.font.size = Pt(10)
    # 점선 탭
    dots = "." * max(1, 60 - len(num) - len(title))
    add_run_with_font(p, f" {dots} {page}", size_pt=10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  1. 개요
# ══════════════════════════════════════════════════════════
heading1(doc, "개요", "1")

heading2(doc, "1.1 문서 목적")
body_para(doc,
    "본 문서는 임진강김치 MES(Manufacturing Execution System) 시스템 구축 프로젝트의 "
    "전체 시스템 아키텍처를 정의하고 기술합니다. 시스템의 물리적 구성, 소프트웨어 구조, "
    "데이터 흐름, 보안 및 배포 전략을 포함하여 개발팀과 운영팀이 일관된 아키텍처 기준을 "
    "공유할 수 있도록 합니다.")
body_para(doc,
    "본 설계서는 사업계획서의 기술 스택 및 요구사항을 기반으로 작성되었으며, "
    "이후 상세설계 및 개발 단계의 기준 문서로 활용됩니다.")

heading2(doc, "1.2 아키텍처 설계 원칙")
principles = [
    "계층 분리 원칙: Physical → Edge → Cloud → Data → Presentation의 명확한 5계층 구조 유지",
    "MSA(Micro Service Architecture): 공정별/기능별 독립 서비스 분리로 확장성 확보",
    "Edge-First: 현장 네트워크 장애 시에도 Edge 버퍼를 통한 데이터 무결성 보장",
    "보안 내재화(Security by Design): SSL/TLS, OAuth2, RBAC를 설계 단계부터 적용",
    "표준 프로토콜 준수: OPC-UA, RS-485, Modbus 등 산업 표준 통신 프로토콜 채택",
    "클라우드 네이티브: Docker 컨테이너 기반 배포, CI/CD 자동화",
    "데이터 이중화: RDBMS(운영) + InfluxDB(시계열) 이원화로 성능과 이력관리 동시 달성",
    "AI 통합 설계: RAG 기반 AI Agent를 아키텍처 최상위 계층으로 독립 설계",
]
for p in principles:
    bullet_para(doc, p)

heading2(doc, "1.3 시스템 범위")
body_para(doc,
    "본 시스템의 범위는 김치 제조 전 공정(입고→절단→세척→절임→혼합→금속검출→포장→출고)의 "
    "MES 통합 관리를 대상으로 하며, 현장 설비 데이터 수집부터 클라우드 기반 분석 및 AI 기반 "
    "품질 관리까지 포함합니다.")

scope_headers = ["구분", "범위", "비고"]
scope_rows = [
    ["현장 설비", "PLC, 센서, 금속검출기, 포장기 등 16종", "HW 연계"],
    ["Edge 수집", "Mini PC 기반 데이터 수집 및 전처리", "현장 설치"],
    ["MES Core", "전 공정 통합관리 (8개 공정)", "Cloud 운영"],
    ["데이터 관리", "RDBMS + InfluxDB + Data Mart", "Cloud DB"],
    ["사용자 인터페이스", "Web, POP, 모바일, 현황판", "4종 UI"],
    ["AI Agent", "RAG 기반 품질/공정 지원 Agent", "AI 플랫폼"],
    ["ERP 연계", "더존 ERP 데이터 연계", "외부 시스템"],
]
add_table(doc, scope_headers, scope_rows, col_widths=[3, 9, 4])

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  2. 전체 시스템 아키텍처
# ══════════════════════════════════════════════════════════
heading1(doc, "전체 시스템 아키텍처", "2")

heading2(doc, "2.1 5계층 아키텍처 개요")
body_para(doc,
    "임진강김치 MES 시스템은 현장의 물리 설비에서 AI 분석까지 6개 계층으로 구성됩니다. "
    "아래 다이어그램은 각 계층의 구성 요소와 데이터 흐름을 나타냅니다.")

diagram_lines = [
    "┌─────────────────────────────────────────────────────────────────┐",
    "│          [ LAYER 6 ]  AI Agent Layer                            │",
    "│   RAG Engine │ LLM │ HACCP지식베이스 │ 작업표준서 │ 레시피DB    │",
    "├─────────────────────────────────────────────────────────────────┤",
    "│          [ LAYER 5 ]  Presentation Layer                        │",
    "│  관리자 Web(PC) │ POP 단말(터치PC) │ 모바일/패드 │ 대형현황판  │",
    "├─────────────────────────────────────────────────────────────────┤",
    "│          [ LAYER 4 ]  Data Layer                                │",
    "│   MySQL(운영DB) │ InfluxDB(시계열) │ Data Mart(공정/LOT/제품)   │",
    "├─────────────────────────────────────────────────────────────────┤",
    "│          [ LAYER 3 ]  Cloud Layer (MES Core)                    │",
    "│  입고 │ 절단 │ 세척 │ 절임 │ 혼합 │ 금속검출 │ 포장 │ 출고    │",
    "│       Spring Boot / FastAPI  ·  RESTful API  ·  MSA             │",
    "├─────────────────────────────────────────────────────────────────┤",
    "│          [ LAYER 2 ]  Edge Layer                                │",
    "│    Mini PC  │ Edge Collector │ 프로토콜 변환 │ 로컬 버퍼        │",
    "│    OPC-UA / RS-485 / Analog(4-20mA) / Ethernet 통합 수집        │",
    "├─────────────────────────────────────────────────────────────────┤",
    "│          [ LAYER 1 ]  Physical Layer (현장 설비)                │",
    "│  PLC Master/Slave │ 아이스박스포장기 │ 금속검출기 │ 자동포장기  │",
    "│  온습도센서×10 │ 온도조절기×4 │ 염도센서×2 │ 소독수 IF Module  │",
    "└─────────────────────────────────────────────────────────────────┘",
    "",
    "  데이터 흐름 방향:  현장설비 ──▶ Edge ──▶ Cloud MES ──▶ DB/AI",
    "  제어 흐름 방향:    Cloud MES ──▶ Edge ──▶ PLC/설비",
]
add_diagram_box(doc, diagram_lines, title="[그림 2-1] 임진강김치 MES 6계층 아키텍처 전체도")

heading2(doc, "2.2 계층별 역할과 책임")
layer_headers = ["계층", "구성 요소", "주요 역할", "통신/기술"]
layer_rows = [
    ["Physical Layer\n(L1)", "PLC, 센서, 포장기,\n금속검출기", "현장 설비 제어,\n실시간 데이터 생성", "RS-485, OPC-UA,\n4-20mA Analog"],
    ["Edge Layer\n(L2)", "Mini PC,\nEdge Collector", "프로토콜 변환, 데이터\n전처리, 로컬 버퍼링", "OPC-UA Client,\nSerial, Ethernet"],
    ["Cloud MES\n(L3)", "Spring Boot/FastAPI\nMSA 서비스", "공정 통합관리, 비즈니스\n로직, API 제공", "REST API,\nMessage Queue"],
    ["Data Layer\n(L4)", "MySQL, InfluxDB,\nData Mart", "운영데이터 저장, 시계열\n데이터, 분석용 Mart", "JDBC, InfluxDB\nLine Protocol"],
    ["Presentation\n(L5)", "Web, POP, 모바일,\n현황판", "사용자 인터페이스,\n실시간 모니터링", "React/Vue.js,\nWebSocket"],
    ["AI Agent\n(L6)", "RAG Engine,\nLLM, 지식베이스", "품질 이상 감지, HACCP\n기준 위반 알림, 공정 최적화", "Python, Vector DB,\nLLM API"],
]
add_table(doc, layer_headers, layer_rows, col_widths=[3, 4, 5, 4])

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  3. HW 아키텍처
# ══════════════════════════════════════════════════════════
heading1(doc, "HW 아키텍처", "3")

heading2(doc, "3.1 현장 설비 및 센서 구성")
body_para(doc,
    "현장에 설치되는 HW 장비의 목록과 각 장비의 용도 및 통신 방식은 아래 표와 같습니다.")

hw_headers = ["품목", "수량", "설치 위치/용도", "통신 방식", "수집 데이터"]
hw_rows = [
    ["아이스박스 자동포장기", "1대", "포장 공정 자동화", "Ethernet / OPC-UA", "포장완료수량, 가동상태"],
    ["PLC Master", "1대", "설비 제어 마스터", "Ethernet / OPC-UA", "제어 명령/상태"],
    ["PLC Slave", "1대", "설비 제어 슬레이브", "Ethernet / OPC-UA", "설비 제어 상태"],
    ["온습도 센서", "10개", "각 공정실 환경 감시", "RS-485", "온도(℃), 습도(%)"],
    ["온도조절기", "4대", "절임/숙성 온도 제어", "RS-485", "현재온도, 설정온도"],
    ["염도센서", "2개", "절임 공정 염도 측정", "Analog 4-20mA", "염도(%)"],
    ["소독수 Interface Module", "1개", "소독수 농도 관리", "Ethernet / OPC-UA", "소독수농도(ppm), 시간"],
    ["금속검출기 PLC", "2대", "이물질 검출", "RS-485", "검출결과(OK/NG), 검사수량"],
    ["대형 현황판", "1대", "공장 내 실시간 현황 표시", "Ethernet", "화면 출력"],
    ["터치 PC (POP 단말)", "3대", "현장 작업 입력/조회", "Ethernet / Wi-Fi", "작업 데이터 입출력"],
    ["Mini PC (Edge Collector)", "1대", "현장 데이터 수집 서버", "Ethernet + USB", "모든 센서 데이터 집계"],
    ["관리용 컴퓨터", "1대", "사무실 관리자 단말", "Ethernet", "MES Web 접속"],
    ["스마트패드", "2대", "이동 중 공정 확인", "Wi-Fi", "MES 모바일 접속"],
    ["무선 AP", "6개", "공장 전역 Wi-Fi 구성", "802.11ac", "무선 네트워크 제공"],
]
add_table(doc, hw_headers, hw_rows, col_widths=[3.8, 1.5, 4, 3, 3.7])

heading2(doc, "3.2 네트워크 구성")
body_para(doc,
    "현장 네트워크는 유선 Ethernet 기간망과 무선 Wi-Fi 보조망으로 이중 구성됩니다. "
    "무선 AP 6개를 공장 전역에 분산 배치하여 스마트패드 및 POP 단말의 안정적인 접속을 보장합니다.")

net_diagram = [
    "  [ 클라우드 (AWS/Azure) ]",
    "          │  인터넷/전용선",
    "  ┌───────┴───────────────────────┐",
    "  │  방화벽 / UTM                 │",
    "  └───────┬───────────────────────┘",
    "          │",
    "  ┌───────┴─────────────────────────────────────────┐",
    "  │  코어 스위치 (L3)                               │",
    "  └───┬──────────┬────────────┬─────────────────────┘",
    "      │          │            │",
    "  [서버존]  [관리자존]  [현장 네트워크 존]",
    "  Mini PC    관리PC       ├─ 터치PC×3 (Ethernet)",
    "                          ├─ 현황판 (Ethernet)",
    "                          ├─ PLC Master/Slave (Ethernet)",
    "                          ├─ 포장기 (Ethernet)",
    "                          ├─ 소독수 Module (Ethernet)",
    "                          └─ 무선 AP×6 ──▶ 스마트패드×2",
    "",
    "  RS-485 버스:  금속검출기PLC × 2  +  온습도센서 × 10  +  온도조절기 × 4",
    "  Analog 루프:  염도센서 × 2  (4-20mA  →  Mini PC AI/O 모듈 변환)",
]
add_diagram_box(doc, net_diagram, title="[그림 3-1] 현장 네트워크 구성도")

heading2(doc, "3.3 서버 구성")
srv_headers = ["서버 유형", "사양(권장)", "역할", "위치"]
srv_rows = [
    ["Cloud Application Server", "vCPU 4Core, RAM 16GB", "MES Core 서비스 운영", "AWS/Azure"],
    ["Cloud DB Server", "vCPU 4Core, RAM 32GB, SSD 1TB", "MySQL + InfluxDB 운영", "AWS/Azure"],
    ["Mini PC (Edge)", "Intel i5, RAM 8GB, SSD 256GB", "현장 데이터 수집·전처리", "현장 설치"],
    ["관리용 컴퓨터", "Intel i5, RAM 8GB", "관리자 MES 접속", "사무실"],
]
add_table(doc, srv_headers, srv_rows, col_widths=[4, 5, 5, 2])

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  4. SW 아키텍처
# ══════════════════════════════════════════════════════════
heading1(doc, "SW 아키텍처", "4")

# 4.1 Frontend
heading2(doc, "4.1 Frontend 아키텍처")
body_para(doc, "Frontend는 React 또는 Vue.js 기반으로 개발하며, 단일 코드베이스에서 "
    "PC 웹, POP 단말, 모바일/스마트패드, 대형 현황판 등 4종의 화면을 반응형으로 제공합니다.")

fe_diagram = [
    "  ┌──────────────────────────────────────────────────────────┐",
    "  │                 Frontend (React / Vue.js)                │",
    "  ├────────────┬──────────────┬────────────┬─────────────────┤",
    "  │ 관리자 Web │  POP 단말    │  모바일    │  대형 현황판    │",
    "  │ (PC 1920px)│ (Touch 1024) │ (모바일)   │ (현황판 4K)     │",
    "  ├────────────┴──────────────┴────────────┴─────────────────┤",
    "  │  공통 컴포넌트 라이브러리 (Design System)                │",
    "  │  상태관리: Redux/Pinia  │  라우팅: React Router/Vue Router│",
    "  │  실시간 통신: WebSocket / SSE                            │",
    "  │  API Client: Axios + JWT Token 자동 갱신                 │",
    "  └──────────────────────────────────────────────────────────┘",
    "                          │ REST API / WebSocket",
    "                   [ Backend API Gateway ]",
]
add_diagram_box(doc, fe_diagram, title="[그림 4-1] Frontend 아키텍처")

fe_headers = ["화면 유형", "해상도/디바이스", "주요 기능", "특이사항"]
fe_rows = [
    ["관리자 Web", "PC / 1920×1080", "전 공정 모니터링, LOT 관리, 보고서", "RBAC 권한 관리"],
    ["POP 단말", "터치PC / 1024×768", "작업 지시 확인, 실적 입력", "터치 최적화 UI"],
    ["모바일/패드", "스마트패드 / 768px~", "이동 중 현황 조회, 알림 수신", "반응형 레이아웃"],
    ["대형 현황판", "대형 모니터 / 4K", "공장 전체 KPI 실시간 표시", "키오스크 모드, 자동 새로고침"],
]
add_table(doc, fe_headers, fe_rows, col_widths=[3, 4, 6, 4])

# 4.2 Backend
heading2(doc, "4.2 Backend 아키텍처")
body_para(doc, "Backend는 Spring Boot(Java) 또는 FastAPI(Python)를 기반으로 MSA 구조로 구현합니다. "
    "각 공정 또는 기능 도메인을 독립 마이크로서비스로 분리하여 독립 배포 및 스케일 아웃을 지원합니다.")

be_diagram = [
    "  ┌──────────────────────────────────────────────────────────┐",
    "  │              API Gateway (인증·라우팅·로드밸런싱)        │",
    "  └──────────────────────────────────────────────────────────┘",
    "       │           │          │           │          │",
    "  ┌────┴──┐  ┌─────┴──┐ ┌────┴────┐ ┌────┴───┐ ┌───┴──────┐",
    "  │ 입고  │  │공정관리│ │ 품질관리 │ │ 설비관리 │ │ 보고서   │",
    "  │Service│  │Service │ │ Service  │ │ Service  │ │ Service  │",
    "  └───────┘  └────────┘ └──────────┘ └──────────┘ └──────────┘",
    "       │           │          │           │          │",
    "  ┌──────────────────────────────────────────────────────────┐",
    "  │              Message Broker (RabbitMQ / Kafka)           │",
    "  └──────────────────────────────────────────────────────────┘",
    "       │           │          │",
    "  ┌────┴──┐  ┌─────┴──┐ ┌────┴────┐",
    "  │ MySQL │  │InfluxDB│ │Data Mart │",
    "  └───────┘  └────────┘ └──────────┘",
]
add_diagram_box(doc, be_diagram, title="[그림 4-2] Backend MSA 구조")

msa_headers = ["서비스명", "기술스택", "주요 기능", "DB"]
msa_rows = [
    ["API Gateway", "Spring Cloud Gateway", "인증, 라우팅, Rate Limiting, 로깅", "-"],
    ["입고관리 Service", "Spring Boot / FastAPI", "원재료 입고, 검수, LOT 생성", "MySQL"],
    ["공정관리 Service", "Spring Boot / FastAPI", "8개 공정 실적 관리, 작업지시", "MySQL"],
    ["품질관리 Service", "Spring Boot / FastAPI", "금속검출, HACCP 기준 검증", "MySQL + InfluxDB"],
    ["설비관리 Service", "Spring Boot / FastAPI", "설비 가동률, 이상 감지, Edge 수신", "InfluxDB"],
    ["보고서 Service", "Spring Boot / FastAPI", "생산 실적, 품질 보고서, ERP 연계", "Data Mart"],
    ["AI Agent Service", "Python (FastAPI)", "RAG 엔진, LLM 연계, 품질 예측", "Vector DB"],
    ["알림 Service", "Spring Boot", "이메일, SMS, Push 알림 발송", "MySQL"],
]
add_table(doc, msa_headers, msa_rows, col_widths=[4, 3.5, 5.5, 3])

# 4.3 데이터
heading2(doc, "4.3 데이터 아키텍처")
body_para(doc, "데이터 계층은 운영 데이터용 RDBMS(MySQL), 센서 시계열 데이터용 InfluxDB, "
    "분석용 Data Mart의 3원 구조로 설계합니다. 각 저장소는 데이터 특성에 최적화된 "
    "스키마와 보존 정책을 적용합니다.")

data_diagram = [
    "  [Edge Collector]",
    "       │",
    "       ├──────────────────────────────────────────────┐",
    "       │                                              │",
    "       ▼                                              ▼",
    "  ┌────────────┐                             ┌────────────────┐",
    "  │   MySQL    │  (운영 트랜잭션 데이터)      │   InfluxDB     │",
    "  │            │                             │                │",
    "  │ ▪ LOT 마스터│                             │ ▪ 센서 원시값  │",
    "  │ ▪ 작업 지시 │                             │ ▪ 설비 상태    │",
    "  │ ▪ 품질 검사 │                             │ ▪ 환경 이력    │",
    "  │ ▪ 입출고    │                             │ (1초~1분 단위) │",
    "  │ ▪ 사용자    │                             │ 보존정책: 2년  │",
    "  └────────────┘                             └────────────────┘",
    "       │                                              │",
    "       └──────────────────┬───────────────────────────┘",
    "                          ▼",
    "                  ┌───────────────┐",
    "                  │  Data Mart    │",
    "                  │               │",
    "                  │ ▪ 공정별 Mart │",
    "                  │ ▪ LOT별 Mart  │",
    "                  │ ▪ 제품별 Mart │",
    "                  │ ▪ 일/월 집계  │",
    "                  └───────────────┘",
    "                          │",
    "              [보고서 / AI 분석 / ERP 연계]",
]
add_diagram_box(doc, data_diagram, title="[그림 4-3] 데이터 아키텍처")

db_headers = ["DB", "기술", "저장 데이터", "보존 정책"]
db_rows = [
    ["운영 DB", "MySQL 8.0", "LOT, 작업지시, 품질, 입출고, 사용자", "영구 (Archive 정책)"],
    ["시계열 DB", "InfluxDB 2.x", "센서값, 설비상태, 환경데이터", "2년 (Downsampling 적용)"],
    ["Data Mart", "MySQL (별도 스키마)", "공정/LOT/제품별 집계 데이터", "5년"],
    ["Vector DB", "ChromaDB / Weaviate", "HACCP 기준, 작업표준서 임베딩", "영구 (갱신 관리)"],
]
add_table(doc, db_headers, db_rows, col_widths=[3, 4, 6, 3])

# 4.4 AI Agent
heading2(doc, "4.4 AI Agent 아키텍처")
body_para(doc, "AI Agent는 RAG(Retrieval-Augmented Generation) 기반으로 구현하며, "
    "HACCP 기준서, 작업표준서, 레시피, 공정이력 데이터를 학습하여 "
    "품질 이상 감지, 공정 최적화 제안, HACCP 위반 사전 알림 기능을 제공합니다.")

ai_diagram = [
    "  ┌─────────────────────────────────────────────────────┐",
    "  │               AI Agent Layer                        │",
    "  │                                                     │",
    "  │  사용자 질의 / 시스템 트리거                        │",
    "  │          │                                          │",
    "  │  ┌───────▼────────┐    ┌─────────────────────┐     │",
    "  │  │ Query Analyzer │    │  Knowledge Base      │     │",
    "  │  │  (의도 파악)   │    │ ▪ HACCP 기준서       │     │",
    "  │  └───────┬────────┘    │ ▪ 작업표준서         │     │",
    "  │          │ Retrieval   │ ▪ 레시피 DB           │     │",
    "  │  ┌───────▼────────┐    │ ▪ 공정이력 (벡터화)  │     │",
    "  │  │ Vector Search  │◀───│                      │     │",
    "  │  │ (유사 문서 검색)│    └─────────────────────┘     │",
    "  │  └───────┬────────┘                                 │",
    "  │          │ Context + Query                          │",
    "  │  ┌───────▼────────┐                                 │",
    "  │  │  LLM (Claude / │                                 │",
    "  │  │  GPT-4 / 로컬) │                                 │",
    "  │  └───────┬────────┘                                 │",
    "  │          │ Response                                 │",
    "  │  ┌───────▼────────┐                                 │",
    "  │  │ Action Engine  │ ──▶ 알림/보고서/제어 명령        │",
    "  │  └────────────────┘                                 │",
    "  └─────────────────────────────────────────────────────┘",
]
add_diagram_box(doc, ai_diagram, title="[그림 4-4] AI Agent RAG 아키텍처")

ai_headers = ["기능", "설명", "학습 데이터"]
ai_rows = [
    ["품질 이상 감지", "센서 데이터 패턴 분석 → 이상 사전 탐지", "공정이력, 불량이력"],
    ["HACCP 위반 알림", "HACCP 기준과 실측값 비교 → 위반 시 즉시 알림", "HACCP 기준서, 측정값"],
    ["공정 최적화 제안", "과거 우수 공정 데이터 기반 최적 조건 추천", "레시피 DB, 공정이력"],
    ["작업 가이드", "작업자 질의에 대한 표준 작업 절차 안내", "작업표준서"],
    ["이상 원인 분석", "불량 발생 시 연관 공정/원재료 추적 분석", "LOT 이력, 공정이력"],
]
add_table(doc, ai_headers, ai_rows, col_widths=[4, 7, 5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  5. 보안 아키텍처
# ══════════════════════════════════════════════════════════
heading1(doc, "보안 아키텍처", "5")

heading2(doc, "5.1 전송 보안: SSL/TLS")
body_para(doc, "모든 클라이언트-서버 간 통신은 SSL/TLS 1.3을 적용하여 전송 중 데이터 암호화를 보장합니다. "
    "Edge-Cloud 간 통신에도 mTLS(상호인증)를 적용하여 인가된 Edge 단말만 연결을 허용합니다.")

tls_items = [
    "공용 도메인: Let's Encrypt 또는 상용 인증서 적용 (1년 갱신)",
    "Edge-Cloud 구간: mTLS 클라이언트 인증서 발급 (Edge 장치별 고유 인증서)",
    "내부 서비스 간(MSA): 서비스 메시(Istio) 또는 API Gateway를 통한 TLS 종단처리",
    "DB 연결: SSL 연결 필수화, 평문 연결 차단",
]
for item in tls_items:
    bullet_para(doc, item)

heading2(doc, "5.2 인증/인가: OAuth2")
body_para(doc, "사용자 인증은 OAuth2 + JWT(JSON Web Token) 방식으로 구현합니다. "
    "액세스 토큰 유효기간은 1시간, 리프레시 토큰은 7일로 설정하며, "
    "토큰 재발급 시 이전 토큰은 즉시 무효화합니다.")

auth_headers = ["구분", "방식", "유효기간", "적용 범위"]
auth_rows = [
    ["사용자 인증", "OAuth2 + JWT (Access Token)", "1시간", "모든 API 엔드포인트"],
    ["토큰 갱신", "Refresh Token", "7일", "Token 갱신 엔드포인트"],
    ["Edge 인증", "API Key + mTLS 인증서", "1년 (갱신)", "Edge → Cloud 수집 API"],
    ["AI Agent", "Service Account + JWT", "24시간", "AI 서비스 내부 호출"],
]
add_table(doc, auth_headers, auth_rows, col_widths=[3.5, 5, 3, 4.5])

heading2(doc, "5.3 역할 기반 접근제어 (RBAC)")
body_para(doc, "시스템 접근 권한은 역할(Role) 기반으로 관리하며, 최소 권한 원칙을 적용합니다.")

rbac_headers = ["역할", "접근 가능 메뉴", "읽기", "쓰기", "삭제", "시스템 관리"]
rbac_rows = [
    ["시스템 관리자", "전체", "O", "O", "O", "O"],
    ["공장장", "전 공정 + 보고서", "O", "O", "X", "X"],
    ["공정 담당자", "담당 공정 + 공통", "O", "O", "X", "X"],
    ["품질 담당자", "품질관리 + 검사이력", "O", "O", "X", "X"],
    ["현장 작업자", "POP 화면 (담당 공정)", "O", "제한", "X", "X"],
    ["조회자", "보고서 + 현황판", "O", "X", "X", "X"],
]
add_table(doc, rbac_headers, rbac_rows, col_widths=[3.5, 5, 1.5, 1.5, 1.5, 3])

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  6. 배포 아키텍처
# ══════════════════════════════════════════════════════════
heading1(doc, "배포 아키텍처", "6")

heading2(doc, "6.1 Docker 컨테이너 구성")
body_para(doc, "모든 서비스는 Docker 컨테이너로 패키징하여 환경 일관성을 보장합니다. "
    "Docker Compose(개발/스테이징) 및 Kubernetes 또는 ECS(운영)로 오케스트레이션합니다.")

docker_diagram = [
    "  [ 운영 환경 - Cloud ]",
    "  ┌──────────────────────────────────────────────────────┐",
    "  │  Kubernetes Cluster / AWS ECS                        │",
    "  │                                                      │",
    "  │  ┌──────────┐  ┌──────────┐  ┌──────────┐           │",
    "  │  │ api-gw   │  │process-  │  │quality-  │  ...      │",
    "  │  │ :8080    │  │svc :8081 │  │svc :8082 │           │",
    "  │  └──────────┘  └──────────┘  └──────────┘           │",
    "  │  ┌──────────┐  ┌──────────┐  ┌──────────┐           │",
    "  │  │frontend  │  │ai-agent  │  │influxdb  │           │",
    "  │  │ :3000    │  │svc :8090 │  │ :8086    │           │",
    "  │  └──────────┘  └──────────┘  └──────────┘           │",
    "  │                                                      │",
    "  │  Shared: MySQL :3306  │  Redis Cache :6379           │",
    "  └──────────────────────────────────────────────────────┘",
    "",
    "  [ Edge 환경 - Mini PC ]",
    "  ┌──────────────────────────┐",
    "  │  Docker Compose          │",
    "  │  ▪ edge-collector :9000  │",
    "  │  ▪ local-buffer (SQLite) │",
    "  │  ▪ protocol-adapter      │",
    "  └──────────────────────────┘",
]
add_diagram_box(doc, docker_diagram, title="[그림 6-1] Docker 컨테이너 배포 구성")

heading2(doc, "6.2 클라우드 배포 전략 (AWS/Azure)")
cloud_headers = ["구성 요소", "AWS 서비스", "Azure 서비스", "역할"]
cloud_rows = [
    ["컴퓨팅", "ECS Fargate / EKS", "AKS / ACI", "컨테이너 오케스트레이션"],
    ["로드밸런서", "ALB (Application LB)", "Azure Load Balancer", "트래픽 분산"],
    ["관계형 DB", "RDS MySQL", "Azure Database for MySQL", "운영 DB"],
    ["시계열 DB", "EC2 + InfluxDB", "VM + InfluxDB", "센서 데이터"],
    ["파일 스토리지", "S3", "Azure Blob Storage", "첨부파일, 백업"],
    ["CDN", "CloudFront", "Azure CDN", "정적 리소스 배포"],
    ["보안", "WAF + Shield", "Azure DDoS + WAF", "웹 방화벽"],
    ["모니터링", "CloudWatch", "Azure Monitor", "로그, 메트릭"],
]
add_table(doc, cloud_headers, cloud_rows, col_widths=[3.5, 4, 4, 4.5])

heading2(doc, "6.3 CI/CD 파이프라인")
body_para(doc, "소스코드 변경 → 자동 빌드·테스트 → 스테이징 배포 → 승인 → 운영 배포의 "
    "CI/CD 파이프라인을 구성합니다.")

cicd_diagram = [
    "  개발자 Push",
    "     │",
    "     ▼",
    "  [GitHub / GitLab]  ──▶  CI Pipeline (GitHub Actions / GitLab CI)",
    "                           │",
    "                    ┌──────┴──────┐",
    "                    │  Build      │  ▪ Docker 이미지 빌드",
    "                    │  Unit Test  │  ▪ 단위 테스트 실행",
    "                    │  SonarQube  │  ▪ 코드 품질 검사",
    "                    └──────┬──────┘",
    "                           │ (성공 시)",
    "                    ┌──────▼──────┐",
    "                    │  Staging    │  ▪ 스테이징 환경 자동 배포",
    "                    │  배포 + 테스트│  ▪ 통합 테스트 실행",
    "                    └──────┬──────┘",
    "                           │ (수동 승인)",
    "                    ┌──────▼──────┐",
    "                    │  Production │  ▪ Blue-Green 배포",
    "                    │  배포       │  ▪ 롤백 전략 포함",
    "                    └─────────────┘",
]
add_diagram_box(doc, cicd_diagram, title="[그림 6-2] CI/CD 파이프라인")

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  7. 시스템 연계 아키텍처
# ══════════════════════════════════════════════════════════
heading1(doc, "시스템 연계 아키텍처", "7")

heading2(doc, "7.1 ERP 연계 (더존 ERP)")
body_para(doc, "더존 ERP와의 연계는 REST API 또는 DB 직접 연계(뷰/프로시저) 방식을 검토하며, "
    "양방향 데이터 동기화를 지원합니다. 연계 주기는 실시간(이벤트 기반) 또는 배치(1시간 단위)를 "
    "데이터 유형에 따라 적용합니다.")

erp_headers = ["연계 데이터", "방향", "연계 방식", "주기"]
erp_rows = [
    ["원재료 입고 정보", "ERP → MES", "REST API / DB View", "실시간"],
    ["생산 실적", "MES → ERP", "REST API", "배치 (1시간)"],
    ["품질 검사 결과", "MES → ERP", "REST API", "실시간"],
    ["LOT 추적 이력", "MES → ERP", "REST API", "배치 (일 1회)"],
    ["재고 현황", "MES ↔ ERP", "DB 동기화", "배치 (30분)"],
    ["출고 정보", "MES → ERP", "REST API", "실시간"],
]
add_table(doc, erp_headers, erp_rows, col_widths=[4, 2.5, 4.5, 5])

heading2(doc, "7.2 현장 설비 연계 (PLC, 센서)")
body_para(doc, "Edge Collector(Mini PC)가 현장 설비와 다양한 프로토콜로 통신하여 "
    "데이터를 수집하고, 표준화된 형태로 Cloud MES에 전송합니다.")

device_headers = ["설비/센서", "프로토콜", "수집 주기", "수집 데이터", "처리 방식"]
device_rows = [
    ["염도센서 × 2", "Analog 4-20mA", "1초", "염도(%)", "AI/O 모듈 변환"],
    ["소독수 PLC × 1", "Ethernet / OPC-UA", "1초", "소독수농도(ppm), 시간", "OPC-UA Client"],
    ["금속검출기 PLC × 2", "RS-485", "이벤트", "검출결과(OK/NG), 수량", "Modbus RTU 폴링"],
    ["자동포장기 × 1", "Ethernet / OPC-UA", "1초", "포장수량, 가동상태", "OPC-UA Client"],
    ["온습도센서 × 10", "RS-485", "30초", "온도(℃), 습도(%)", "Modbus RTU 폴링"],
    ["온도조절기 × 4", "RS-485", "30초", "현재온도, 설정온도", "Modbus RTU 폴링"],
]
add_table(doc, device_headers, device_rows, col_widths=[3.5, 3, 2, 4, 3.5])

heading2(doc, "7.3 Edge-Cloud 데이터 흐름")
edge_cloud_diagram = [
    "  [ 현장 설비 ]          [ Edge Layer ]           [ Cloud MES ]",
    "",
    "  PLC/센서/포장기  ──▶  Protocol Adapter    ──▶  API Gateway",
    "  (RS-485/OPC-UA/       (OPC-UA Client,           (HTTPS/TLS)",
    "   4-20mA)               Modbus RTU,                   │",
    "                         Analog 변환)           ┌──────┴──────┐",
    "                              │                 │             │",
    "                      ┌───────▼────────┐  설비관리Svc   공정관리Svc",
    "                      │  Local Buffer  │        │             │",
    "                      │  (SQLite/Queue)│        ▼             ▼",
    "                      └───────┬────────┘   InfluxDB       MySQL",
    "                              │",
    "              [네트워크 정상] ─┤─ [네트워크 장애]",
    "                              │         │",
    "                        Cloud 전송   로컬 버퍼 적재",
    "                                    (복구 후 일괄 전송)",
    "",
    "  버퍼 용량: 최대 72시간 데이터 로컬 보관 (약 500MB 예상)",
]
add_diagram_box(doc, edge_cloud_diagram, title="[그림 7-1] Edge-Cloud 데이터 흐름도")

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  8. 성능 및 가용성 설계
# ══════════════════════════════════════════════════════════
heading1(doc, "성능 및 가용성 설계", "8")

heading2(doc, "8.1 성능 목표")
body_para(doc, "시스템의 성능 목표를 아래와 같이 정의하며, 부하 테스트를 통해 목표 달성 여부를 검증합니다.")

perf_headers = ["성능 지표", "목표값", "측정 방법", "비고"]
perf_rows = [
    ["API 응답 시간 (평균)", "< 500ms", "APM 모니터링 (p95)", "일반 조회 기준"],
    ["API 응답 시간 (최대)", "< 2,000ms", "APM 모니터링 (p99)", "피크 부하 기준"],
    ["센서 데이터 수집 지연", "< 3초", "Edge ↔ Cloud 타임스탬프", "실시간 공정 기준"],
    ["동시 사용자", "50명 이상", "부하 테스트 (JMeter)", "내부 사용자 전원"],
    ["데이터 처리 TPS", "1,000 TPS", "InfluxDB 입력 속도", "센서 데이터 피크"],
    ["시스템 가용성", "99.5% 이상", "월별 다운타임 계산", "연 44시간 이내"],
    ["데이터 백업 RTO", "< 4시간", "복구 훈련 결과", "장애 복구 목표"],
    ["데이터 백업 RPO", "< 1시간", "백업 주기 설정", "데이터 손실 허용"],
]
add_table(doc, perf_headers, perf_rows, col_widths=[4.5, 3, 4.5, 4])

heading2(doc, "8.2 장애 대응 (Edge 버퍼, 이중화)")
body_para(doc, "단일 장애점(SPOF)을 최소화하고, 계층별 장애 대응 전략을 수립합니다.")

ha_headers = ["장애 유형", "영향 범위", "대응 전략", "복구 목표"]
ha_rows = [
    ["Cloud 서버 장애", "MES 전체 기능 중단", "Auto Scaling + 다중 AZ 배포,\nBlue-Green 전환", "RTO < 30분"],
    ["Cloud-Edge 네트워크 단절", "실시간 수집 중단", "Edge 로컬 버퍼 72시간 보관,\n복구 후 자동 일괄 전송", "데이터 손실 없음"],
    ["Mini PC(Edge) 장애", "현장 수집 중단", "설비별 직접 데이터 저장 (일부),\nEdge 교체 시 설정 자동 복구", "RTO < 2시간"],
    ["DB 장애 (MySQL)", "운영 데이터 접근 불가", "Read Replica + 자동 Failover,\n정기 스냅샷 백업", "RTO < 1시간"],
    ["DB 장애 (InfluxDB)", "시계열 데이터 저장 불가", "Edge 버퍼 축적 후 복구 시 전송", "RPO < 1시간"],
    ["개별 센서 장애", "해당 센서 값 누락", "알림 발생 + 마지막 정상값 표시,\n수동 입력 대체", "작업 지속 가능"],
    ["사이버 공격 (DDoS)", "서비스 접근 불가", "WAF + Shield (클라우드),\nRate Limiting", "자동 완화"],
]
add_table(doc, ha_headers, ha_rows, col_widths=[3.5, 3.5, 5.5, 3.5])

body_para(doc, "")
body_para(doc,
    "[이중화 구성 요약]",
    indent=False)

ha_items = [
    "Cloud 서버: 다중 가용영역(Multi-AZ) 배포, Auto Scaling Group 적용",
    "DB (MySQL): Primary + Read Replica 구성, 자동 Failover 5분 이내",
    "로드밸런서: 클라우드 관리형 LB (단일장애점 제거)",
    "Edge Collector: 로컬 버퍼 72시간 + 복구 후 자동 재전송",
    "네트워크: 유선 기간망 + 무선 AP 이중 경로 (AP 6개 분산)",
    "백업: 일 1회 전체 백업 + 1시간 단위 증분 백업 (S3/Blob Storage)",
]
for item in ha_items:
    bullet_para(doc, item)

# ── 하단 여백 ──
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run_with_font(p_footer,
    "── 임진강김치 MES 아키텍처설계서 V1.0  ·  IK-MES-ARCH-001  ·  2026-05-12  ·  대외비 ──",
    size_pt=8, color=RGBColor(0x80, 0x80, 0x80))

# ── 저장 ──
doc.save(OUTPUT_PATH)
print(f"[완료] 아키텍처설계서가 저장되었습니다:\n  {OUTPUT_PATH}")
