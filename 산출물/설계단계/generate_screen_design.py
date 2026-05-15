# -*- coding: utf-8 -*-
"""
임진강김치 MES 화면설계서 (시스템관리·KPI·설비관리) 생성 스크립트
"""
import sys
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy
except ImportError:
    print("python-docx not found. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

OUTPUT_PATH = r"C:\gerardo\01 SmallSF\Imjingang-Kimchi\산출물\설계단계\06_화면설계서_시스템관리.docx"

# ─────────────────────────────────────────────────────────────
# Helper utilities
# ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), kwargs.get(side, 'single'))
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), kwargs.get('color', '4472C4'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def add_para(doc, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Malgun Gothic'
    if color:
        run.font.color.rgb = color
    return p

def add_mockup(doc, lines):
    """Add ASCII mockup in Courier New monospace font."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Light gray background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    run = p.add_run(lines)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    return p

def add_table_header(doc, headers, col_widths=None):
    """Create a styled table with header row."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, (cell, hdr) in enumerate(zip(hdr_cells, headers)):
        set_cell_bg(cell, '2E74B5')
        cell.text = hdr
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Malgun Gothic'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_widths:
            cell.width = Cm(col_widths[i])
    return table

def add_table_row(table, values, alt=False):
    row = table.add_row()
    bg = 'EBF3FB' if alt else 'FFFFFF'
    for i, (cell, val) in enumerate(zip(row.cells, values)):
        set_cell_bg(cell, bg)
        cell.text = str(val)
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(val))
        run.font.size = Pt(9)
        run.font.name = 'Malgun Gothic'
    return row

def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{label}: ')
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.name = 'Malgun Gothic'
    r1.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    r2 = p.add_run(value)
    r2.font.size = Pt(9)
    r2.font.name = 'Malgun Gothic'
    return p

def page_break(doc):
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────
# Cover Page
# ─────────────────────────────────────────────────────────────
def build_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("임진강김치 MES 시스템 구축")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.name = 'Malgun Gothic'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("화면설계서\n(시스템관리 · KPI 모니터링 · 설비관리)")
    r2.font.size = Pt(26)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    r2.font.name = 'Malgun Gothic'

    doc.add_paragraph()
    doc.add_paragraph()

    meta = [
        ("문서번호", "IK-MES-DES-006"),
        ("버전", "V1.0"),
        ("작성일", "2026-05-12"),
        ("작성자", "디자이너3"),
        ("검토자", "PM / 시스템분석가"),
        ("승인자", "프로젝트 관리자"),
        ("프로젝트", "임진강김치 MES 구축 프로젝트"),
        ("대상 화면", "시스템관리(4), KPI(3), 설비관리(5), 공통팝업(4) — 총 16화면"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        row = table.rows[i]
        set_cell_bg(row.cells[0], 'D6E4F0')
        row.cells[0].text = k
        row.cells[1].text = v
        for cell in row.cells:
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            run.font.name = 'Malgun Gothic'
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(10)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("본 문서는 화면설계서(코어기능)의 보완 자료입니다.")
    r3.font.size = Pt(9)
    r3.font.italic = True
    r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    r3.font.name = 'Malgun Gothic'
    page_break(doc)

# ─────────────────────────────────────────────────────────────
# Section 1: Overview
# ─────────────────────────────────────────────────────────────
def build_overview(doc):
    add_heading(doc, "1. 개요", 1)
    add_heading(doc, "1.1 화면 목록", 2)

    screen_list = [
        # (ID, 화면명, 분류, 비고)
        ("SCR-SYS-001", "사용자 계정 관리", "시스템관리", "등록/조회/수정/삭제"),
        ("SCR-SYS-002", "권한 역할 관리", "시스템관리", "역할 정의 및 메뉴 권한 설정"),
        ("SCR-SYS-003", "시스템 로그 조회", "시스템관리", "감사 로그 조회"),
        ("SCR-SYS-004", "데이터 백업 이력 조회", "시스템관리", "백업 실행 및 이력 조회"),
        ("SCR-KPI-001", "생산성 KPI 조회", "KPI 모니터링", "시간당 생산량, 리드타임"),
        ("SCR-KPI-002", "품질 KPI 조회", "KPI 모니터링", "불량률, 검식 적합률"),
        ("SCR-KPI-003", "KPI 지표 관리", "KPI 모니터링", "KPI 코드/계산식 정의"),
        ("SCR-EQP-001", "설비 가동 현황 조회", "설비관리", "실시간 가동 상태"),
        ("SCR-EQP-002", "설비 점검 등록", "설비관리", "정기/비정기 점검 등록"),
        ("SCR-EQP-003", "설비 점검 이력 조회", "설비관리", "점검 결과 이력"),
        ("SCR-EQP-004", "설비 이상 등록", "설비관리", "이상 발생 즉시 등록"),
        ("SCR-EQP-005", "설비 고장 이력 조회", "설비관리", "고장/수리 이력"),
        ("SCR-CMN-001", "날짜 범위 선택 팝업", "공통", "달력 UI 팝업"),
        ("SCR-CMN-002", "엑셀 다운로드 처리", "공통", "진행 상태 표시"),
        ("SCR-CMN-003", "알람 상세 팝업", "공통", "알람 내용 및 조치 안내"),
        ("SCR-CMN-004", "확인/취소 공통 팝업", "공통", "사용자 확인 다이얼로그"),
    ]

    headers = ["화면 ID", "화면명", "분류", "비고"]
    widths = [3.5, 4.5, 3.0, 5.0]
    table = add_table_header(doc, headers, widths)
    for i, row in enumerate(screen_list):
        add_table_row(table, row, alt=(i % 2 == 1))

    doc.add_paragraph()
    add_heading(doc, "1.2 접근 권한 매트릭스", 2)
    add_para(doc, "※ O = 접근 가능, X = 접근 불가, R = 읽기 전용", size=9)

    perm_headers = ["화면 ID", "화면명", "시스템관리자", "현장관리자", "작업자"]
    perm_widths = [3.2, 4.5, 2.5, 2.8, 2.5]
    perm_data = [
        ("SCR-SYS-001", "사용자 계정 관리", "O", "X", "X"),
        ("SCR-SYS-002", "권한 역할 관리", "O", "X", "X"),
        ("SCR-SYS-003", "시스템 로그 조회", "O", "R", "X"),
        ("SCR-SYS-004", "데이터 백업 이력 조회", "O", "X", "X"),
        ("SCR-KPI-001", "생산성 KPI 조회", "O", "O", "R"),
        ("SCR-KPI-002", "품질 KPI 조회", "O", "O", "R"),
        ("SCR-KPI-003", "KPI 지표 관리", "O", "X", "X"),
        ("SCR-EQP-001", "설비 가동 현황 조회", "O", "O", "O"),
        ("SCR-EQP-002", "설비 점검 등록", "O", "O", "X"),
        ("SCR-EQP-003", "설비 점검 이력 조회", "O", "O", "R"),
        ("SCR-EQP-004", "설비 이상 등록", "O", "O", "O"),
        ("SCR-EQP-005", "설비 고장 이력 조회", "O", "O", "R"),
        ("SCR-CMN-001~004", "공통 팝업/컴포넌트", "O", "O", "O"),
    ]
    pt = add_table_header(doc, perm_headers, perm_widths)
    for i, row in enumerate(perm_data):
        tr = add_table_row(pt, row, alt=(i % 2 == 1))
        # Color O/X/R cells
        for ci in range(2, 5):
            cell = tr.cells[ci]
            val = row[ci]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if val == "O":
                set_cell_bg(cell, 'D9EAD3')
            elif val == "X":
                set_cell_bg(cell, 'FCE8E8')
            elif val == "R":
                set_cell_bg(cell, 'FFF2CC')

    page_break(doc)

# ─────────────────────────────────────────────────────────────
# Section 2: System Management Screens
# ─────────────────────────────────────────────────────────────
def build_sys_screens(doc):
    add_heading(doc, "2. 시스템관리 화면 (SCR-SYS-001 ~ 004)", 1)

    # ── SCR-SYS-001 ──────────────────────────────────────────
    add_heading(doc, "2.1 사용자 계정 관리 (SCR-SYS-001)", 2)
    add_label_value(doc, "화면 ID", "SCR-SYS-001")
    add_label_value(doc, "화면명", "사용자 계정 관리")
    add_label_value(doc, "접근 권한", "시스템관리자")
    add_label_value(doc, "화면 유형", "조회/등록/수정/삭제 (CRUD)")
    add_label_value(doc, "관련 메뉴", "시스템관리 > 사용자 계정 관리")

    add_heading(doc, "목업", 3)
    mockup_sys001 = """\
[사용자 계정 관리]  SCR-SYS-001                    [관리자: 시스템관리자]
+----------------------------------------------------------------------+
| 메뉴 > 시스템관리 > 사용자 계정 관리                                  |
+----------------------------------------------------------------------+
| [검색 조건]                                                           |
|  사용자명 : [____________]  부서 : [전체        v]                    |
|  역할     : [전체        v]  상태 : [전체        v]  [조회]  [초기화] |
+----------------------------------------------------------------------+
| [검색 결과]  총 24건                        [+ 신규 등록]  [엑셀 다운] |
| +--------+----------+--------+----------+--------+------+-----------+ |
| |선택    |사용자 ID |이름    |부서      |역할    |상태  |최근 로그인| |
| +--------+----------+--------+----------+--------+------+-----------+ |
| |[ ]     |USR001   |김현장  |생산1팀   |작업자  |활성  |2026-05-12 | |
| |[ ]     |USR002   |이관리  |생산관리팀|현장관리|활성  |2026-05-11 | |
| |[ ]     |USR003   |박품질  |품질관리팀|현장관리|활성  |2026-05-10 | |
| |[ ]     |USR004   |최설비  |설비팀    |작업자  |활성  |2026-05-09 | |
| |[ ]     |USR005   |정창고  |물류팀    |작업자  |비활성|2026-04-30 | |
| |[ ]     |USR006   |한시스  |IT팀      |시스템관|활성  |2026-05-12 | |
| +--------+----------+--------+----------+--------+------+-----------+ |
|                                           [선택삭제]  [< 1 2 3 4 >]  |
+----------------------------------------------------------------------+
|                                                                       |
| [사용자 상세 / 등록 패널]  (행 클릭 시 우측 또는 하단에 펼침)         |
| +------------------------------------------------------------------+  |
| |  사용자 ID  : [USR001        ]   이름  : [김현장      ]         |  |
| |  부서       : [생산1팀       v]  직책  : [반장        ]         |  |
| |  역할       : [작업자        v]  이메일: [kim@imjin.co.kr     ] |  |
| |  휴대폰     : [010-1234-5678 ]   상태  : ( )활성  ( )비활성    |  |
| |  비밀번호   : [초기화]           비고  : [                    ] |  |
| |                                  [저장]    [취소]    [삭제]     |  |
| +------------------------------------------------------------------+  |
+----------------------------------------------------------------------+
"""
    add_mockup(doc, mockup_sys001)

    add_heading(doc, "컴포넌트 설명", 3)
    comp_headers = ["컴포넌트", "유형", "설명", "유효성 검사"]
    comp_widths = [3.5, 2.5, 5.0, 4.5]
    ct = add_table_header(doc, comp_headers, comp_widths)
    comp_data = [
        ("사용자명 검색", "TextInput", "부분 일치 검색 (LIKE 쿼리)", "최대 50자"),
        ("부서 드롭다운", "Select", "조직 마스터에서 동적 로드", "필수 선택"),
        ("역할 드롭다운", "Select", "시스템 역할 코드 로드", "-"),
        ("상태 드롭다운", "Select", "활성/비활성/잠금", "-"),
        ("조회 버튼", "Button", "검색 조건으로 목록 조회", "-"),
        ("신규 등록 버튼", "Button", "하단 패널 초기화 후 표시", "-"),
        ("목록 그리드", "DataGrid", "페이지당 10건, 행 클릭 상세", "다중 선택 체크박스"),
        ("엑셀 다운 버튼", "Button", "현재 조회 결과 전체 다운로드", "SCR-CMN-002 호출"),
        ("사용자 ID", "TextInput (R)", "신규 시 자동 채번 (USR+순번)", "읽기 전용"),
        ("역할 드롭다운(상세)", "Select", "작업자/현장관리자/시스템관리자", "필수"),
        ("상태 라디오", "RadioButton", "활성/비활성 토글", "필수"),
        ("비밀번호 초기화", "Button", "임시 비밀번호 이메일 발송", "확인 팝업 필요"),
        ("저장 버튼", "Button", "입력값 유효성 검사 후 저장", "필수항목 검증"),
        ("삭제 버튼", "Button", "소프트 삭제 (비활성화 처리)", "확인 팝업 필요"),
    ]
    for i, row in enumerate(comp_data):
        add_table_row(ct, row, alt=(i % 2 == 1))

    doc.add_paragraph()
    add_heading(doc, "이벤트 및 기능 설명", 3)
    evt_headers = ["이벤트", "동작", "비고"]
    et = add_table_header(doc, evt_headers, [4.0, 8.0, 3.5])
    evt_data = [
        ("조회 버튼 클릭", "검색 조건 파라미터로 API 호출 → 목록 갱신", "GET /api/users"),
        ("신규 등록 버튼", "상세 패널 초기화 (ID 자동 채번)", "POST /api/users"),
        ("목록 행 클릭", "해당 사용자 상세 정보 하단 패널 표시", "GET /api/users/{id}"),
        ("비밀번호 초기화", "SCR-CMN-004 확인 팝업 → 이메일 발송", "POST /api/users/{id}/reset-pw"),
        ("저장 버튼", "입력값 유효성 검사 → 신규/수정 분기 저장", "PUT /api/users/{id}"),
        ("삭제 버튼", "SCR-CMN-004 확인 팝업 → 소프트 삭제", "DELETE /api/users/{id}"),
        ("선택 삭제 버튼", "체크박스 선택 행 일괄 소프트 삭제", "DELETE /api/users (batch)"),
        ("엑셀 다운 버튼", "SCR-CMN-002 다운로드 팝업 호출", "GET /api/users/export"),
    ]
    for i, row in enumerate(evt_data):
        add_table_row(et, row, alt=(i % 2 == 1))

    doc.add_paragraph()

    # ── SCR-SYS-002 ──────────────────────────────────────────
    add_heading(doc, "2.2 권한 역할 관리 (SCR-SYS-002)", 2)
    add_label_value(doc, "화면 ID", "SCR-SYS-002")
    add_label_value(doc, "화면명", "권한 역할 관리")
    add_label_value(doc, "접근 권한", "시스템관리자")
    add_label_value(doc, "화면 유형", "역할별 메뉴 권한 매트릭스 설정")
    add_label_value(doc, "관련 메뉴", "시스템관리 > 권한 역할 관리")

    add_heading(doc, "목업", 3)
    mockup_sys002 = """\
[권한 역할 관리]  SCR-SYS-002                       [관리자: 시스템관리자]
+-----------------------------------------------------------------------+
| 메뉴 > 시스템관리 > 권한 역할 관리                                     |
+-----------------------------------------------------------------------+
| [역할 목록]                                [+ 역할 추가]              |
| +----------+------------------+----------+----------+               |
| |역할 코드 |역할명            |사용자 수 |상태      |               |
| +----------+------------------+----------+----------+               |
| |ROLE_ADM  |시스템관리자      |2명       |활성      |  ← 선택됨     |
| |ROLE_MGR  |현장관리자        |8명       |활성      |               |
| |ROLE_OPR  |작업자            |14명      |활성      |               |
| +----------+------------------+----------+----------+               |
|                                                                       |
| [메뉴 권한 설정]  역할: 시스템관리자                                   |
| +-----+------------------+------+------+------+------+------+       |
| |No   |메뉴명            |조회  |등록  |수정  |삭제  |엑셀  |       |
| +-----+------------------+------+------+------+------+------+       |
| |1    |생산계획          |[V]   |[V]   |[V]   |[V]   |[V]   |       |
| |2    |작업지시          |[V]   |[V]   |[V]   |[V]   |[V]   |       |
| |3    |실적 입력         |[V]   |[V]   |[V]   |[ ]   |[V]   |       |
| |4    |품질 검사         |[V]   |[V]   |[V]   |[V]   |[V]   |       |
| |5    |KPI 모니터링      |[V]   |[V]   |[V]   |[V]   |[V]   |       |
| |6    |설비관리          |[V]   |[V]   |[V]   |[V]   |[V]   |       |
| |7    |시스템관리        |[V]   |[V]   |[V]   |[V]   |[V]   |       |
| +-----+------------------+------+------+------+------+------+       |
|                              [전체 선택]  [전체 해제]  [저장]         |
+-----------------------------------------------------------------------+
"""
    add_mockup(doc, mockup_sys002)

    add_heading(doc, "컴포넌트 설명", 3)
    ct2 = add_table_header(doc, ["컴포넌트", "유형", "설명", "비고"], [3.5, 2.5, 5.5, 4.0])
    comp2 = [
        ("역할 목록 그리드", "DataGrid", "등록된 역할 코드 및 사용자 수 표시", "행 클릭 시 우측 권한 로드"),
        ("역할 추가 버튼", "Button", "새 역할 코드/명 입력 인라인 폼", "-"),
        ("메뉴 권한 체크박스", "Checkbox", "조회/등록/수정/삭제/엑셀 5개 권한", "독립 토글"),
        ("전체 선택/해제", "Button", "현재 역할의 모든 권한 일괄 선택/해제", "-"),
        ("저장 버튼", "Button", "변경된 권한 매트릭스 일괄 저장", "SCR-CMN-004 확인"),
    ]
    for i, row in enumerate(comp2):
        add_table_row(ct2, row, alt=(i % 2 == 1))

    doc.add_paragraph()

    # ── SCR-SYS-003 ──────────────────────────────────────────
    add_heading(doc, "2.3 시스템 로그 조회 (SCR-SYS-003)", 2)
    add_label_value(doc, "화면 ID", "SCR-SYS-003")
    add_label_value(doc, "화면명", "시스템 로그 조회")
    add_label_value(doc, "접근 권한", "시스템관리자(전체), 현장관리자(읽기 전용)")
    add_label_value(doc, "화면 유형", "조회 전용")

    add_heading(doc, "목업", 3)
    mockup_sys003 = """\
[시스템 로그 조회]  SCR-SYS-003
+-----------------------------------------------------------------------+
| 메뉴 > 시스템관리 > 시스템 로그 조회                                   |
+-----------------------------------------------------------------------+
| [검색 조건]                                                            |
|  기간   : [2026-05-01] ~ [2026-05-12]  [SCR-CMN-001 날짜선택팝업]    |
|  사용자 : [____________]  메뉴명  : [____________]                    |
|  액션   : [전체        v]  결과   : [전체        v]  [조회]  [초기화] |
+-----------------------------------------------------------------------+
| [검색 결과]  총 1,284건                                 [엑셀 다운]   |
| +-------+----------+----------+------------------+------+-----------+ |
| |No     |일시      |사용자 ID |메뉴/기능         |액션  |결과       | |
| +-------+----------+----------+------------------+------+-----------+ |
| |1      |05-12 14:32|USR006  |사용자계정관리    |삭제  |성공       | |
| |2      |05-12 14:20|USR002  |작업지시 등록     |등록  |성공       | |
| |3      |05-12 13:55|USR003  |품질검사 결과입력 |수정  |성공       | |
| |4      |05-12 13:10|USR001  |설비이상 등록     |등록  |성공       | |
| |5      |05-12 12:44|USR005  |로그인            |로그인|실패(비번) | |
| |6      |05-12 12:30|USR002  |생산계획 조회     |조회  |성공       | |
| +-------+----------+----------+------------------+------+-----------+ |
|                                               [< 1 2 3 ... 129 >]    |
+-----------------------------------------------------------------------+
| [로그 상세]  No.5 선택 시                                              |
|  시각: 2026-05-12 12:44:30  /  IP: 192.168.1.105  /  브라우저: Chrome|
|  요청 URL: POST /api/auth/login                                        |
|  실패 사유: 비밀번호 불일치 (5회 연속 실패 → 계정 잠금 처리)          |
+-----------------------------------------------------------------------+
"""
    add_mockup(doc, mockup_sys003)

    add_heading(doc, "컴포넌트 설명", 3)
    ct3 = add_table_header(doc, ["컴포넌트", "유형", "설명", "비고"], [3.5, 2.5, 5.5, 4.0])
    comp3 = [
        ("기간 조건", "DateRangePicker", "SCR-CMN-001 날짜 범위 팝업 연동", "최대 90일"),
        ("액션 드롭다운", "Select", "로그인/조회/등록/수정/삭제/다운로드", "-"),
        ("결과 드롭다운", "Select", "성공/실패/오류", "-"),
        ("목록 그리드", "DataGrid", "페이지당 20건, 행 클릭 시 상세", "읽기 전용"),
        ("실패 행 강조", "ConditionalStyle", "결과=실패 시 배경색 연빨강 표시", "FFF0F0"),
        ("로그 상세 패널", "DetailPanel", "IP, URL, 실패 사유 등 표시", "행 클릭 연동"),
        ("엑셀 다운 버튼", "Button", "조회 결과 전체 엑셀 다운로드", "SCR-CMN-002 호출"),
    ]
    for i, row in enumerate(comp3):
        add_table_row(ct3, row, alt=(i % 2 == 1))

    doc.add_paragraph()

    # ── SCR-SYS-004 ──────────────────────────────────────────
    add_heading(doc, "2.4 데이터 백업 이력 조회 (SCR-SYS-004)", 2)
    add_label_value(doc, "화면 ID", "SCR-SYS-004")
    add_label_value(doc, "화면명", "데이터 백업 이력 조회")
    add_label_value(doc, "접근 권한", "시스템관리자")
    add_label_value(doc, "화면 유형", "조회 및 수동 백업 실행")

    add_heading(doc, "목업", 3)
    mockup_sys004 = """\
[데이터 백업 이력 조회]  SCR-SYS-004
+-----------------------------------------------------------------------+
| 메뉴 > 시스템관리 > 데이터 백업 이력 조회                              |
+-----------------------------------------------------------------------+
| [백업 현황 요약]                                                       |
|  마지막 자동 백업: 2026-05-12 02:00  /  상태: 성공  /  크기: 2.3 GB   |
|  백업 보관 주기: 30일  /  보관 경로: /backup/mes/                      |
|                                              [수동 백업 실행]          |
+-----------------------------------------------------------------------+
| [검색 조건]                                                            |
|  기간: [2026-05-01] ~ [2026-05-12]  유형: [전체 v]  결과: [전체 v]   |
|                                                    [조회]  [초기화]   |
+-----------------------------------------------------------------------+
| [백업 이력]  총 36건                                    [엑셀 다운]   |
| +----+------------------+--------+--------+--------+--------------+  |
| |No  |백업 일시         |유형    |크기    |결과    |보관 경로     |  |
| +----+------------------+--------+--------+--------+--------------+  |
| |1   |2026-05-12 02:00  |자동    |2.3 GB  |성공    |/backup/0512  |  |
| |2   |2026-05-11 02:00  |자동    |2.2 GB  |성공    |/backup/0511  |  |
| |3   |2026-05-10 09:15  |수동    |2.2 GB  |성공    |/backup/0510m |  |
| |4   |2026-05-10 02:00  |자동    |2.1 GB  |실패    |—             |  |
| |5   |2026-05-09 02:00  |자동    |2.1 GB  |성공    |/backup/0509  |  |
| +----+------------------+--------+--------+--------+--------------+  |
|                                               [< 1 2 3 4 >]         |
+-----------------------------------------------------------------------+
| [오류 상세]  No.4 선택 시                                              |
|  오류 사유: 디스크 용량 부족 (가용 공간 < 3 GB)  /  조치: 수동 재실행  |
+-----------------------------------------------------------------------+
"""
    add_mockup(doc, mockup_sys004)

    add_heading(doc, "컴포넌트 설명", 3)
    ct4 = add_table_header(doc, ["컴포넌트", "유형", "설명", "비고"], [3.5, 2.5, 5.5, 4.0])
    comp4 = [
        ("백업 현황 요약", "InfoPanel", "마지막 성공 백업 정보 요약 표시", "실시간 갱신"),
        ("수동 백업 실행 버튼", "Button", "즉시 백업 트리거, 진행 상태 표시", "SCR-CMN-004 확인"),
        ("유형 드롭다운", "Select", "자동/수동 필터", "-"),
        ("결과 드롭다운", "Select", "성공/실패 필터", "-"),
        ("목록 그리드", "DataGrid", "페이지당 10건, 실패 행 빨강 강조", "읽기 전용"),
        ("오류 상세 패널", "DetailPanel", "실패 원인 및 조치 방법 표시", "행 클릭 연동"),
    ]
    for i, row in enumerate(comp4):
        add_table_row(ct4, row, alt=(i % 2 == 1))

    page_break(doc)

# ─────────────────────────────────────────────────────────────
# Section 3: KPI Screens
# ─────────────────────────────────────────────────────────────
def build_kpi_screens(doc):
    add_heading(doc, "3. KPI 모니터링 화면 (SCR-KPI-001 ~ 003)", 1)

    # ── SCR-KPI-001 ──────────────────────────────────────────
    add_heading(doc, "3.1 생산성 KPI 조회 (SCR-KPI-001)", 2)
    add_label_value(doc, "화면 ID", "SCR-KPI-001")
    add_label_value(doc, "화면명", "생산성 KPI 조회")
    add_label_value(doc, "접근 권한", "전체 (작업자: 읽기 전용)")
    add_label_value(doc, "화면 유형", "대시보드형 KPI 조회")
    add_label_value(doc, "주요 지표", "시간당 생산량(UPH), 리드타임(Lead Time), 가동률")

    add_heading(doc, "목업", 3)
    mockup_kpi001 = """\
[생산성 KPI 조회]  SCR-KPI-001
+------------------------------------------------------------------------+
| 메뉴 > KPI 모니터링 > 생산성 KPI 조회                                   |
+------------------------------------------------------------------------+
| [조회 조건]                                                             |
|  기간  : [2026-05-01] ~ [2026-05-12]   집계 단위: ( )일 (*)주 ( )월   |
|  라인  : [전체 v]   제품군: [전체 v]                   [조회]           |
+------------------------------------------------------------------------+
| [KPI 요약 카드]                                                         |
|  +------------------+ +------------------+ +------------------+        |
|  | 시간당 생산량(UPH)|  | 리드타임(평균)   |  | 라인 가동률      |       |
|  |   ★ 1,240 박스   |  |   ★  4.2 시간   |  |   ★  94.3 %     |       |
|  | 목표: 1,200 | +3% |  | 목표:  4.0 | -5%|  | 목표: 95% | -1% |       |
|  +------------------+  +------------------+  +------------------+       |
|  +------------------+                                                   |
|  | 설비 OEE          |                                                   |
|  |   ★  87.2 %      |                                                   |
|  | 목표: 90% | -3%   |                                                   |
|  +------------------+                                                   |
+------------------------------------------------------------------------+
| [추이 차트]  시간당 생산량 (UPH) — 주간 추이                            |
|  1,400 |              *                                                  |
|  1,300 |        *         *    *                                         |
|  1,200 |----목표선--------------------------                             |
|  1,100 |  *                        *                                     |
|  1,000 |                                  *                              |
|        +--+--+--+--+--+--+--+--+--+--+--+                              |
|           5/1 5/2 5/3 5/4 5/5 5/6 5/7 5/8 5/9 5/10 5/11 5/12          |
+------------------------------------------------------------------------+
| [라인별 상세 테이블]                                                     |
| +-----------+--------+--------+--------+--------+---------------------+ |
| |라인       |UPH 실적|UPH 목표|달성률  |리드타임|OEE                  | |
| +-----------+--------+--------+--------+--------+---------------------+ |
| |생산1라인  |1,280   |1,200   |106.7%  |4.0h    |89.5%                | |
| |생산2라인  |1,190   |1,200   | 99.2%  |4.3h    |86.1%                | |
| |생산3라인  |1,250   |1,200   |104.2%  |4.1h    |88.7%                | |
| |포장라인   |  980   |1,000   | 98.0%  |2.5h    |84.3%                | |
| +-----------+--------+--------+--------+--------+---------------------+ |
|                                                         [엑셀 다운]     |
+------------------------------------------------------------------------+
"""
    add_mockup(doc, mockup_kpi001)

    add_heading(doc, "컴포넌트 설명", 3)
    ct_k1 = add_table_header(doc, ["컴포넌트", "유형", "설명", "비고"], [3.5, 2.5, 5.5, 4.0])
    comp_k1 = [
        ("기간 조건", "DateRangePicker", "SCR-CMN-001 팝업 연동", "최대 3개월"),
        ("집계 단위 라디오", "RadioGroup", "일/주/월 단위 집계 전환", "차트 단위 변경"),
        ("KPI 요약 카드", "KPICard", "실적/목표/달성률 색상 표시", "목표 초과: 파랑, 미달: 빨강"),
        ("추이 차트", "LineChart", "기간별 UPH 추이, 목표선 표시", "Chart.js / ECharts"),
        ("라인별 상세 테이블", "DataGrid", "라인별 KPI 수치 비교", "읽기 전용"),
        ("달성률 색상", "ConditionalStyle", "100% 이상: 파랑, 미만: 빨강", "-"),
        ("엑셀 다운 버튼", "Button", "조회 결과 엑셀 다운로드", "SCR-CMN-002 호출"),
    ]
    for i, row in enumerate(comp_k1):
        add_table_row(ct_k1, row, alt=(i % 2 == 1))

    doc.add_paragraph()
    add_heading(doc, "이벤트 및 기능 설명", 3)
    et_k1 = add_table_header(doc, ["이벤트", "동작", "API"], [4.0, 7.5, 4.0])
    for i, row in enumerate([
        ("조회 버튼", "기간/라인/집계단위로 KPI 데이터 호출", "GET /api/kpi/productivity"),
        ("집계 단위 변경", "차트 및 테이블 데이터 재집계 렌더링", "쿼리 파라미터 변경"),
        ("KPI 카드 클릭", "해당 KPI 상세 드릴다운 화면 이동", "내부 필터 변경"),
        ("엑셀 다운", "SCR-CMN-002 팝업 후 다운로드", "GET /api/kpi/productivity/export"),
    ]):
        add_table_row(et_k1, row, alt=(i % 2 == 1))

    doc.add_paragraph()

    # ── SCR-KPI-002 ──────────────────────────────────────────
    add_heading(doc, "3.2 품질 KPI 조회 (SCR-KPI-002)", 2)
    add_label_value(doc, "화면 ID", "SCR-KPI-002")
    add_label_value(doc, "화면명", "품질 KPI 조회")
    add_label_value(doc, "접근 권한", "전체 (작업자: 읽기 전용)")
    add_label_value(doc, "주요 지표", "불량률, 검식 적합률, 반품률, 재작업률")

    add_heading(doc, "목업", 3)
    mockup_kpi002 = """\
[품질 KPI 조회]  SCR-KPI-002
+------------------------------------------------------------------------+
| 메뉴 > KPI 모니터링 > 품질 KPI 조회                                     |
+------------------------------------------------------------------------+
| [조회 조건]                                                             |
|  기간  : [2026-05-01] ~ [2026-05-12]   집계 단위: ( )일 (*)주 ( )월   |
|  라인  : [전체 v]   제품군: [전체 v]   불량 유형: [전체 v]  [조회]      |
+------------------------------------------------------------------------+
| [KPI 요약 카드]                                                         |
|  +------------------+ +------------------+ +------------------+        |
|  | 불량률            |  | 검식 적합률      |  | 재작업률         |       |
|  |   ★  1.8 %       |  |   ★  98.2 %     |  |   ★  2.1 %      |       |
|  | 목표: 2.0% | 양호 |  | 목표: 98% | 양호|  | 목표: 2.0% | 주의|       |
|  +------------------+  +------------------+  +------------------+       |
|  +------------------+                                                   |
|  | 반품률            |                                                   |
|  |   ★  0.3 %       |                                                   |
|  | 목표: 0.5% | 양호 |                                                   |
|  +------------------+                                                   |
+------------------------------------------------------------------------+
| [불량 유형별 파레토 차트]                                                |
|  100%  |                                                                |
|   80%  |            ~~~누적불량률 곡선~~~                                |
|   60%  |                                                                |
|   40%  |  ████                                                           |
|   20%  |  ████  ████                                                     |
|    0%  |  ████  ████  ████  ████  ████                                  |
|        +--이물--중량--포장--이취--기타--                                  |
|         35%    25%    18%    13%   9%                                    |
+------------------------------------------------------------------------+
| [라인별 불량률 상세]                                                     |
| +-----------+-------+-------+--------+-------+-----------------------+ |
| |라인       |검사수 |불량수 |불량률  |적합률 |주요 불량 유형         | |
| +-----------+-------+-------+--------+-------+-----------------------+ |
| |생산1라인  |8,500  |145    |1.71%   |98.29% |이물(52), 중량(43)     | |
| |생산2라인  |8,200  |168    |2.05%   |97.95% |포장(61), 이물(55)     | |
| |생산3라인  |8,350  |141    |1.69%   |98.31% |이취(38), 중량(35)     | |
| |포장라인   |5,100  | 62    |1.22%   |98.78% |포장(42), 기타(20)     | |
| +-----------+-------+-------+--------+-------+-----------------------+ |
|                                                         [엑셀 다운]     |
+------------------------------------------------------------------------+
"""
    add_mockup(doc, mockup_kpi002)

    add_heading(doc, "컴포넌트 설명", 3)
    ct_k2 = add_table_header(doc, ["컴포넌트", "유형", "설명", "비고"], [3.5, 2.5, 5.5, 4.0])
    comp_k2 = [
        ("불량 유형 드롭다운", "Select", "이물/중량/포장/이취/기타/전체", "-"),
        ("KPI 요약 카드", "KPICard", "목표 대비 상태 표시 (양호/주의/위험)", "색상 코딩"),
        ("파레토 차트", "BarChart+LineChart", "불량 유형별 건수 + 누적 불량률 겹침", "복합 차트"),
        ("라인별 상세 테이블", "DataGrid", "검사수/불량수/불량률/적합률 표시", "읽기 전용"),
        ("불량률 색상 표시", "ConditionalStyle", "목표 초과시 빨강, 이하시 파랑", "-"),
        ("엑셀 다운 버튼", "Button", "SCR-CMN-002 호출", "-"),
    ]
    for i, row in enumerate(comp_k2):
        add_table_row(ct_k2, row, alt=(i % 2 == 1))

    doc.add_paragraph()

    # ── SCR-KPI-003 ──────────────────────────────────────────
    add_heading(doc, "3.3 KPI 지표 관리 (SCR-KPI-003)", 2)
    add_label_value(doc, "화면 ID", "SCR-KPI-003")
    add_label_value(doc, "화면명", "KPI 지표 관리")
    add_label_value(doc, "접근 권한", "시스템관리자")
    add_label_value(doc, "화면 유형", "KPI 코드 및 계산식 CRUD")

    add_heading(doc, "목업", 3)
    mockup_kpi003 = """\
[KPI 지표 관리]  SCR-KPI-003
+------------------------------------------------------------------------+
| 메뉴 > KPI 모니터링 > KPI 지표 관리                                     |
+------------------------------------------------------------------------+
| [KPI 목록]                              [+ 신규 등록]  [엑셀 다운]      |
| +--------+------------------+--------+--------+------+---------------+ |
| |KPI 코드|KPI명             |단위    |목표값  |상태  |계산식 유형    | |
| +--------+------------------+--------+--------+------+---------------+ |
| |KPI-001 |시간당 생산량(UPH)|박스/시간|1,200  |활성  |수식           | |
| |KPI-002 |리드타임          |시간    |4.0     |활성  |수식           | |
| |KPI-003 |불량률            |%       |2.0     |활성  |수식           | |
| |KPI-004 |검식 적합률       |%       |98.0    |활성  |수식           | |
| |KPI-005 |설비 가동률       |%       |95.0    |활성  |수식           | |
| |KPI-006 |OEE               |%       |90.0    |활성  |복합(A*B*C)    | |
| |KPI-007 |재작업률          |%       |2.0     |활성  |수식           | |
| |KPI-008 |반품률            |%       |0.5     |활성  |수식           | |
| +--------+------------------+--------+--------+------+---------------+ |
+------------------------------------------------------------------------+
| [KPI 상세 / 등록 패널]  KPI-001 선택 시                                 |
| +------------------------------------------------------------------+   |
| |  KPI 코드    : [KPI-001       ]   KPI명 : [시간당 생산량(UPH) ] |   |
| |  단위        : [박스/시간     ]   목표값: [1,200              ] |   |
| |  계산식 유형 : ( )수식  ( )복합식  ( )외부연동                  |   |
| |                                                                  |   |
| |  계산식 정의 :                                                   |   |
| |  +------------------------------------------------------------+  |   |
| |  | (총생산량_박스) / (가동시간_시간)                           |  |   |
| |  | 변수1: 총생산량_박스  → 테이블: T_PROD_RESULT.QTY           |  |   |
| |  | 변수2: 가동시간_시간  → 테이블: T_EQP_STATUS.OPER_TIME     |  |   |
| |  +------------------------------------------------------------+  |   |
| |                                                                  |   |
| |  경고 임계값 : [목표의 95% 이하]  위험 임계값: [목표의 90% 이하]|   |
| |  집계 주기   : ( )실시간  (*)시간별  ( )일별                    |   |
| |  상태        : (*)활성    ( )비활성                              |   |
| |  비고        : [                                               ] |   |
| |                                   [저장]    [취소]    [삭제]    |   |
| +------------------------------------------------------------------+   |
+------------------------------------------------------------------------+
"""
    add_mockup(doc, mockup_kpi003)

    add_heading(doc, "컴포넌트 설명", 3)
    ct_k3 = add_table_header(doc, ["컴포넌트", "유형", "설명", "비고"], [3.5, 2.5, 5.5, 4.0])
    comp_k3 = [
        ("KPI 목록 그리드", "DataGrid", "전체 KPI 코드 목록", "행 클릭 시 상세 패널"),
        ("신규 등록 버튼", "Button", "새 KPI 코드 패널 초기화", "-"),
        ("KPI 코드", "TextInput", "KPI-NNN 형식 자동 채번", "읽기 전용"),
        ("계산식 유형 라디오", "RadioGroup", "수식/복합식/외부연동 선택", "유형별 입력 폼 변경"),
        ("계산식 편집기", "CodeEditor", "수식 입력, 변수 → DB 컬럼 매핑", "문법 검증 기능"),
        ("임계값 설정", "NumberInput", "경고/위험 임계값 % 설정", "KPI 카드 색상 연동"),
        ("집계 주기 라디오", "RadioGroup", "실시간/시간별/일별 선택", "스케줄러 연동"),
        ("저장 버튼", "Button", "계산식 유효성 검증 후 저장", "SCR-CMN-004 확인"),
    ]
    for i, row in enumerate(comp_k3):
        add_table_row(ct_k3, row, alt=(i % 2 == 1))

    page_break(doc)

# ─────────────────────────────────────────────────────────────
# Section 4: Equipment Management Screens
# ─────────────────────────────────────────────────────────────
def build_eqp_screens(doc):
    add_heading(doc, "4. 설비관리 화면 (SCR-EQP-001 ~ 005)", 1)

    # ── SCR-EQP-001 ──────────────────────────────────────────
    add_heading(doc, "4.1 설비 가동 현황 조회 (SCR-EQP-001)", 2)
    add_label_value(doc, "화면 ID", "SCR-EQP-001")
    add_label_value(doc, "화면명", "설비 가동 현황 조회")
    add_label_value(doc, "접근 권한", "전체")
    add_label_value(doc, "화면 유형", "실시간 모니터링 대시보드")
    add_label_value(doc, "갱신 주기", "30초 자동 갱신")

    add_heading(doc, "목업", 3)
    mockup_eqp001 = """\
[설비 가동 현황 조회]  SCR-EQP-001                  [자동갱신: 30초]
+------------------------------------------------------------------------+
| 메뉴 > 설비관리 > 설비 가동 현황 조회                                   |
+------------------------------------------------------------------------+
| [조회 조건]                                                             |
|  공정라인: [전체 v]   설비 유형: [전체 v]   상태: [전체 v]   [조회]    |
+------------------------------------------------------------------------+
| [현황 요약]                                                             |
|  총 설비: 42대   가동 중: 36대   점검 중: 3대   고장: 2대   대기: 1대  |
|  전체 가동률: 85.7%                                           [알람 3건]|
+------------------------------------------------------------------------+
| [설비 상태 맵]  (카드 형태 표시)                                         |
|                                                                         |
|  [생산1라인]                                                             |
|  +----------+ +----------+ +----------+ +----------+ +----------+      |
|  |EQP-001   | |EQP-002   | |EQP-003   | |EQP-004   | |EQP-005   |     |
|  |절임기 #1 | |세척기 #1 | |혼합기 #1 | |충진기 #1 | |포장기 #1 |     |
|  |● 가동중  | |● 가동중  | |⚠ 점검중  | |● 가동중  | |● 가동중  |     |
|  |OEE: 91%  | |OEE: 88%  | |—         | |OEE: 93%  | |OEE: 86%  |     |
|  |온도: 4°C | |RPM: 120  | |점검자:김 | |속도: 95% | |속도: 90% |     |
|  +----------+ +----------+ +----------+ +----------+ +----------+      |
|                                                                         |
|  [생산2라인]                                                             |
|  +----------+ +----------+ +----------+ +----------+                   |
|  |EQP-011   | |EQP-012   | |EQP-013   | |EQP-014   |                  |
|  |절임기 #2 | |세척기 #2 | |혼합기 #2 | |충진기 #2 |                  |
|  |● 가동중  | |✖ 고장    | |● 가동중  | |● 가동중  |                  |
|  |OEE: 89%  | |고장시각: | |OEE: 87%  | |OEE: 90%  |                  |
|  |온도: 4°C | |14:20     | |온도: 5°C | |속도: 88% |                  |
|  +----------+ +----------+ +----------+ +----------+                   |
+------------------------------------------------------------------------+
| [설비 상세]  EQP-012 클릭 시                                             |
|  설비명: 세척기 #2  /  고장 시각: 2026-05-12 14:20  /  담당자: 최설비   |
|  고장 내용: 구동부 과부하 트립  /  조치 진행: 부품 교체 중              |
|  예상 복구: 2026-05-12 16:00  [이상 등록 바로가기 →]                    |
+------------------------------------------------------------------------+
"""
    add_mockup(doc, mockup_eqp001)

    add_heading(doc, "컴포넌트 설명", 3)
    ct_e1 = add_table_header(doc, ["컴포넌트", "유형", "설명", "비고"], [3.5, 2.5, 5.5, 4.0])
    comp_e1 = [
        ("현황 요약 바", "InfoPanel", "설비 총수/상태별 수량/가동률 표시", "30초 자동 갱신"),
        ("알람 버튼", "AlertButton", "현재 미조치 알람 건수 표시", "SCR-CMN-003 팝업"),
        ("설비 상태 카드", "StatusCard", "설비별 상태 아이콘+주요 파라미터", "클릭 시 상세"),
        ("상태 아이콘", "Icon", "● 가동(파랑), ⚠ 점검(노랑), ✖ 고장(빨강)", "색상 코딩"),
        ("자동 갱신", "Timer", "30초 주기 API 재호출 및 카드 갱신", "토글 가능"),
        ("설비 상세 패널", "DetailPanel", "고장 정보, 담당자, 예상 복구 시간", "카드 클릭 연동"),
        ("이상 등록 링크", "Link", "SCR-EQP-004 이상 등록 화면 이동", "설비 ID 전달"),
    ]
    for i, row in enumerate(comp_e1):
        add_table_row(ct_e1, row, alt=(i % 2 == 1))

    doc.add_paragraph()

    # ── SCR-EQP-002 ──────────────────────────────────────────
    add_heading(doc, "4.2 설비 점검 등록 (SCR-EQP-002)", 2)
    add_label_value(doc, "화면 ID", "SCR-EQP-002")
    add_label_value(doc, "화면명", "설비 점검 등록")
    add_label_value(doc, "접근 권한", "시스템관리자, 현장관리자")
    add_label_value(doc, "화면 유형", "점검 체크리스트 기반 등록")

    add_heading(doc, "목업", 3)
    mockup_eqp002 = """\
[설비 점검 등록]  SCR-EQP-002
+------------------------------------------------------------------------+
| 메뉴 > 설비관리 > 설비 점검 등록                                        |
+------------------------------------------------------------------------+
| [점검 기본 정보]                                                        |
|  설비 선택 : [EQP-001 절임기 #1     v]                                  |
|  점검 유형 : ( )일상점검  (*)정기점검  ( )특별점검                      |
|  점검 일시 : [2026-05-12 15:00     ]  점검자: [김설비 (USR004) v]       |
|  점검 주기 : 주간  /  다음 점검 예정: 2026-05-19                         |
+------------------------------------------------------------------------+
| [점검 체크리스트]  (설비 유형: 절임기)                                   |
| +----+---------------------------+----------+-------+----------------+ |
| |No  |점검 항목                  |점검 방법 |기준   |결과            | |
| +----+---------------------------+----------+-------+----------------+ |
| |1   |구동 모터 작동 상태        |육안/청각 |정상   |( )정상 ( )이상 | |
| |2   |절임조 온도 유지 (목표: 4°C)|온도계    |±1°C   |실측: [4.2] °C | |
| |3   |구동 벨트 장력             |육안      |정상   |( )정상 ( )이상 | |
| |4   |베어링 윤활유 수준         |육안      |적정   |( )정상 ( )이상 | |
| |5   |제어판 표시등 상태         |육안      |정상   |( )정상 ( )이상 | |
| |6   |안전 커버 체결 상태        |육안      |완전체결|( )정상 ( )이상 | |
| |7   |배수구 막힘 여부           |육안      |정상   |( )정상 ( )이상 | |
| |8   |전원 케이블 피복 상태      |육안      |손상없음|( )정상 ( )이상 | |
| +----+---------------------------+----------+-------+----------------+ |
|                                                         [항목 추가]     |
+------------------------------------------------------------------------+
| [이상 발견 시]   이상 항목: 4번 (베어링 윤활유)                          |
|  이상 내용 : [베어링 윤활유 수준 낮음 — 보충 필요                     ]  |
|  조치 방법 : ( )즉시조치  (*)차기점검  ( )설비정지 후 수리               |
|  조치 내용 : [다음 일상점검 시 윤활유 보충 예정                        ]  |
|  [설비 이상 등록 연계] ← 체크 시 SCR-EQP-004 자동 생성                  |
+------------------------------------------------------------------------+
| 종합 의견 : [전반적으로 양호. 베어링 윤활유 소량 보충 필요.             ]  |
| 점검 결과 : (*)이상없음  ( )이상발견(조치완료)  ( )이상발견(조치필요)     |
|                                                   [임시저장]  [완료저장] |
+------------------------------------------------------------------------+
"""
    add_mockup(doc, mockup_eqp002)

    add_heading(doc, "컴포넌트 설명", 3)
    ct_e2 = add_table_header(doc, ["컴포넌트", "유형", "설명", "비고"], [3.5, 2.5, 5.5, 4.0])
    comp_e2 = [
        ("설비 선택 드롭다운", "Select", "설비 마스터에서 코드+명칭 로드", "필수"),
        ("점검 유형 라디오", "RadioGroup", "일상/정기/특별점검 선택", "체크리스트 로드 연동"),
        ("점검 일시", "DateTimePicker", "기본값: 현재 시각", "과거 입력 가능"),
        ("체크리스트 그리드", "EditableGrid", "점검 항목별 결과 입력", "설비 유형별 동적 로드"),
        ("결과 라디오(행별)", "RadioGroup", "정상/이상 선택", "이상 선택 시 하단 패널 표시"),
        ("실측값 입력", "NumberInput", "수치 점검 항목의 측정값 입력", "기준값 범위 검증"),
        ("이상 패널", "ConditionalPanel", "이상 선택 시 표시 (이상 내용/조치)", "조건부 표시"),
        ("설비 이상 연계 체크", "Checkbox", "SCR-EQP-004 이상 등록 자동 생성", "체크 시 연계"),
        ("임시저장 버튼", "Button", "미완료 상태로 저장 (나중에 이어서)", "상태: DRAFT"),
        ("완료저장 버튼", "Button", "완료 상태로 저장", "필수항목 검증 후 저장"),
    ]
    for i, row in enumerate(comp_e2):
        add_table_row(ct_e2, row, alt=(i % 2 == 1))

    doc.add_paragraph()

    # ── SCR-EQP-003 ──────────────────────────────────────────
    add_heading(doc, "4.3 설비 점검 이력 조회 (SCR-EQP-003)", 2)
    add_label_value(doc, "화면 ID", "SCR-EQP-003")
    add_label_value(doc, "화면명", "설비 점검 이력 조회")
    add_label_value(doc, "접근 권한", "전체 (작업자: 읽기 전용)")
    add_label_value(doc, "화면 유형", "이력 조회")

    add_heading(doc, "목업", 3)
    mockup_eqp003 = """\
[설비 점검 이력 조회]  SCR-EQP-003
+------------------------------------------------------------------------+
| 메뉴 > 설비관리 > 설비 점검 이력 조회                                   |
+------------------------------------------------------------------------+
| [검색 조건]                                                             |
|  기간   : [2026-04-01] ~ [2026-05-12]   설비  : [전체 v]              |
|  점검유형: [전체 v]   결과: [전체 v]   점검자: [____________]  [조회]  |
+------------------------------------------------------------------------+
| [검색 결과]  총 87건                                     [엑셀 다운]    |
| +------+--------+----------+------+--------+--------+--------------+   |
| |점검ID|점검일시|설비      |유형  |점검자  |결과    |이상 연계 여부|   |
| +------+--------+----------+------+--------+--------+--------------+   |
| |INS087|05-12   |절임기 #1 |정기  |김설비  |이상발견|[이상#ABI023] |   |
| |      |15:00   |EQP-001   |      |        |(조치필)|              |   |
| +------+--------+----------+------+--------+--------+--------------+   |
| |INS086|05-11   |충진기 #1 |일상  |이점검  |이상없음|—             |   |
| |      |09:00   |EQP-004   |      |        |        |              |   |
| +------+--------+----------+------+--------+--------+--------------+   |
| |INS085|05-10   |세척기 #2 |특별  |박담당  |이상발견|[이상#ABI022] |   |
| |      |14:30   |EQP-012   |      |        |(조치완)|              |   |
| +------+--------+----------+------+--------+--------+--------------+   |
|                                               [< 1 2 3 ... 9 >]       |
+------------------------------------------------------------------------+
| [점검 상세]  INS087 클릭 시                                              |
|  설비: 절임기 #1 (EQP-001)  /  점검일: 2026-05-12 15:00               |
|  체크리스트: 8항목 중 이상 1건 (4번: 베어링 윤활유)                     |
|  종합의견: 전반적 양호. 베어링 윤활유 소량 보충 필요.                   |
|  연계 이상: ABI023  [이상 상세 보기 →]                                  |
+------------------------------------------------------------------------+
"""
    add_mockup(doc, mockup_eqp003)

    add_heading(doc, "컴포넌트 설명", 3)
    ct_e3 = add_table_header(doc, ["컴포넌트", "유형", "설명", "비고"], [3.5, 2.5, 5.5, 4.0])
    comp_e3 = [
        ("목록 그리드", "DataGrid", "페이지당 10건, 점검 이력 표시", "읽기 전용"),
        ("결과 색상", "ConditionalStyle", "이상발견(조치필): 빨강, 조치완: 노랑", "-"),
        ("이상 연계 링크", "Link", "ABI코드 클릭 시 SCR-EQP-005 이동", "이상 ID 전달"),
        ("점검 상세 패널", "DetailPanel", "체크리스트 결과 및 종합의견", "행 클릭 연동"),
        ("엑셀 다운 버튼", "Button", "SCR-CMN-002 호출", "-"),
    ]
    for i, row in enumerate(comp_e3):
        add_table_row(ct_e3, row, alt=(i % 2 == 1))

    doc.add_paragraph()

    # ── SCR-EQP-004 ──────────────────────────────────────────
    add_heading(doc, "4.4 설비 이상 등록 (SCR-EQP-004)", 2)
    add_label_value(doc, "화면 ID", "SCR-EQP-004")
    add_label_value(doc, "화면명", "설비 이상 등록")
    add_label_value(doc, "접근 권한", "전체 (이상 발생 즉시 누구나 등록 가능)")
    add_label_value(doc, "화면 유형", "이상 발생 즉시 등록 폼")

    add_heading(doc, "목업", 3)
    mockup_eqp004 = """\
[설비 이상 등록]  SCR-EQP-004
+------------------------------------------------------------------------+
| 메뉴 > 설비관리 > 설비 이상 등록                                        |
+------------------------------------------------------------------------+
| [이상 기본 정보]                                            [! 긴급등록] |
|  이상 번호  : [ABI-자동채번      ]   등록 일시: [2026-05-12 14:20]     |
|  설비 선택  : [EQP-012 세척기 #2 v]  공정라인: 생산2라인 (자동)         |
|  이상 유형  : [고장              v]  긴급도   : (*)긴급  ( )일반        |
|  발견자     : [USR001 김현장     v]  발견 경위: ( )육안  (*)이상음 ( )경보|
+------------------------------------------------------------------------+
| [이상 내용]                                                             |
|  이상 제목  : [세척기 #2 구동부 과부하 트립 발생                      ] |
|  이상 상세  :                                                           |
|  +------------------------------------------------------------------+   |
|  | 14:20경 생산2라인 세척기 #2에서 과부하 경보 발생 후 자동 트립.    |   |
|  | 구동 벨트 슬립 또는 이물 끼임으로 추정.                           |   |
|  +------------------------------------------------------------------+   |
|  현재 설비 상태: ( )가동중  (*)정지  ( )부분가동                        |
|  영향 공정     : [생산2라인 세척 공정 전면 중단                       ] |
+------------------------------------------------------------------------+
| [즉각 조치 내용]                                                        |
|  즉각 조치   : [전원 차단 후 안전 확인, 구동부 육안 점검 중           ]  |
|  조치자      : [USR004 최설비     v]   조치 시각: [14:25]              |
|  예상 복구   : [2026-05-12 16:00 ]   연락처    : [010-2345-6789    ]   |
+------------------------------------------------------------------------+
| [첨부 파일]                                                             |
|  [파일 선택]  이상 발생 사진 또는 동영상 첨부 (최대 5개, 10MB/개)       |
|  첨부됨: [EQP012_fault_20260512.jpg  ×]                                 |
+------------------------------------------------------------------------+
|  [알람 발송]  담당자 자동 알람: ☑ 설비팀장  ☑ 현장관리자  □ 생산팀장    |
|                                               [임시저장]  [이상 등록]   |
+------------------------------------------------------------------------+
"""
    add_mockup(doc, mockup_eqp004)

    add_heading(doc, "컴포넌트 설명", 3)
    ct_e4 = add_table_header(doc, ["컴포넌트", "유형", "설명", "비고"], [3.5, 2.5, 5.5, 4.0])
    comp_e4 = [
        ("이상 번호", "TextInput (R)", "ABI-YYYYMMDD-NNN 자동 채번", "읽기 전용"),
        ("설비 선택", "Select", "설비 마스터 로드, 선택 시 라인 자동", "필수"),
        ("이상 유형 드롭다운", "Select", "고장/이상음/과열/누출/기타", "필수"),
        ("긴급도 라디오", "RadioGroup", "긴급: 즉시 알람, 일반: 업무 시간 알람", "알람 정책 연동"),
        ("이상 상세 텍스트", "Textarea", "5W1H 형식 자유 서술", "최소 20자"),
        ("현재 설비 상태", "RadioGroup", "가동중/정지/부분가동", "필수"),
        ("즉각 조치 내용", "Textarea", "취한 조치 내용 기술", "-"),
        ("조치자 드롭다운", "Select", "활성 사용자 목록", "-"),
        ("예상 복구 일시", "DateTimePicker", "복구 예정 시각 입력", "-"),
        ("파일 첨부", "FileUpload", "사진/동영상 최대 5개, 10MB/개", "JPG/PNG/MP4"),
        ("알람 발송 체크", "CheckboxGroup", "수신자 그룹 선택, 등록 후 자동 발송", "-"),
        ("이상 등록 버튼", "Button", "유효성 검증 → 저장 → 알람 발송", "POST /api/eqp/anomaly"),
    ]
    for i, row in enumerate(comp_e4):
        add_table_row(ct_e4, row, alt=(i % 2 == 1))

    doc.add_paragraph()

    # ── SCR-EQP-005 ──────────────────────────────────────────
    add_heading(doc, "4.5 설비 고장 이력 조회 (SCR-EQP-005)", 2)
    add_label_value(doc, "화면 ID", "SCR-EQP-005")
    add_label_value(doc, "화면명", "설비 고장 이력 조회")
    add_label_value(doc, "접근 권한", "전체 (작업자: 읽기 전용)")
    add_label_value(doc, "화면 유형", "고장/이상 이력 조회 및 분석")

    add_heading(doc, "목업", 3)
    mockup_eqp005 = """\
[설비 고장 이력 조회]  SCR-EQP-005
+------------------------------------------------------------------------+
| 메뉴 > 설비관리 > 설비 고장 이력 조회                                   |
+------------------------------------------------------------------------+
| [검색 조건]                                                             |
|  기간   : [2026-04-01] ~ [2026-05-12]   설비  : [전체 v]              |
|  이상유형: [전체 v]   긴급도: [전체 v]   결과  : [전체 v]  [조회]      |
+------------------------------------------------------------------------+
| [이력 요약]                                                             |
|  조회 기간 총 고장: 12건  /  평균 복구 시간(MTTR): 2.8시간              |
|  최다 고장 설비: EQP-012 세척기 #2 (3건)                               |
+------------------------------------------------------------------------+
| [고장 이력 목록]  총 12건                               [엑셀 다운]     |
| +------+--------+----------+------+------+------+------+------------+ |
| |이상ID|발생일시|설비      |유형  |긴급도|발견자|복구일|복구시간(h) | |
| +------+--------+----------+------+------+------+------+------------+ |
| |ABI023|05-12   |세척기 #2 |고장  |긴급  |김현장|진행중|—           | |
| |      |14:20   |EQP-012   |      |      |      |      |            | |
| +------+--------+----------+------+------+------+------+------------+ |
| |ABI022|05-10   |세척기 #2 |이상음|일반  |박담당|05-10 |3.5h        | |
| |      |09:30   |EQP-012   |      |      |      |13:00 |            | |
| +------+--------+----------+------+------+------+------+------------+ |
| |ABI021|05-08   |충진기 #3 |과열  |긴급  |이관리|05-08 |1.5h        | |
| |      |11:00   |EQP-033   |      |      |      |12:30 |            | |
| +------+--------+----------+------+------+------+------+------------+ |
|                                               [< 1 2 >]               |
+------------------------------------------------------------------------+
| [이상 상세]  ABI023 클릭 시                                              |
|  발생: 2026-05-12 14:20  /  설비: EQP-012 세척기 #2  /  긴급도: 긴급   |
|  이상 제목: 세척기 #2 구동부 과부하 트립 발생                           |
|  이상 내용: 14:20경 과부하 경보 발생 후 자동 트립. 구동 벨트 슬립 추정. |
|  즉각 조치: 전원 차단 후 안전 확인, 구동부 육안 점검 중                 |
|  예상 복구: 2026-05-12 16:00  /  조치자: 최설비                         |
|  첨부: [EQP012_fault_20260512.jpg]  [수리 완료 처리 →]  [이상 수정 →]   |
+------------------------------------------------------------------------+
| [고장 통계 차트]                                                         |
|  설비별 고장 빈도 (최근 30일)                                            |
|  EQP-012: ████████████ 5건                                              |
|  EQP-033: ████████ 3건                                                  |
|  EQP-004: ████ 2건                                                      |
|  기타    : ████ 2건                                                      |
+------------------------------------------------------------------------+
"""
    add_mockup(doc, mockup_eqp005)

    add_heading(doc, "컴포넌트 설명", 3)
    ct_e5 = add_table_header(doc, ["컴포넌트", "유형", "설명", "비고"], [3.5, 2.5, 5.5, 4.0])
    comp_e5 = [
        ("이력 요약 패널", "InfoPanel", "총 건수/MTTR/최다 고장 설비 표시", "조회 결과 연동"),
        ("목록 그리드", "DataGrid", "페이지당 10건, 진행중 행 강조", "읽기 전용"),
        ("긴급도 뱃지", "Badge", "긴급: 빨강 뱃지, 일반: 파랑 뱃지", "-"),
        ("이상 상세 패널", "DetailPanel", "이상 내용, 조치, 첨부파일 표시", "행 클릭 연동"),
        ("수리 완료 처리 버튼", "Button", "복구 완료 시각 입력 → 상태 변경", "현장관리자+"),
        ("첨부파일 다운로드", "Link", "등록된 사진/동영상 다운로드", "-"),
        ("고장 통계 차트", "BarChart", "설비별 고장 빈도 가로 막대 차트", "최근 30일 고정"),
        ("엑셀 다운 버튼", "Button", "SCR-CMN-002 호출", "-"),
    ]
    for i, row in enumerate(comp_e5):
        add_table_row(ct_e5, row, alt=(i % 2 == 1))

    page_break(doc)

# ─────────────────────────────────────────────────────────────
# Section 5: Common Components
# ─────────────────────────────────────────────────────────────
def build_common_screens(doc):
    add_heading(doc, "5. 공통 컴포넌트 화면 (SCR-CMN-001 ~ 004)", 1)

    # ── SCR-CMN-001 ──────────────────────────────────────────
    add_heading(doc, "5.1 날짜 범위 선택 팝업 (SCR-CMN-001)", 2)
    add_label_value(doc, "화면 ID", "SCR-CMN-001")
    add_label_value(doc, "화면명", "날짜 범위 선택 팝업")
    add_label_value(doc, "화면 유형", "Modal 팝업 (달력 UI)")
    add_label_value(doc, "호출 화면", "검색 조건이 있는 모든 화면")

    add_heading(doc, "목업", 3)
    mockup_cmn001 = """\
[날짜 범위 선택 팝업]  SCR-CMN-001
+------------------------------------------------+
|  날짜 범위 선택                            [X] |
+------------------------------------------------+
|  [빠른 선택]  [오늘] [이번주] [이번달] [최근3월]|
+------------------------+-----------------------+
|  ◀  2026년 05월  ▶     |  ◀  2026년 06월  ▶   |
| 일  월  화  수  목  금  토 | 일  월  화  수  목  금  토|
|                  1   2  3|                       1  2  3 |
|  4   5   6   7   8   9  10|  4   5   6   7   8   9  10  |
| 11  12  13  14  15  16  17| 11  12  13  14  15  16  17  |
| 18  19  20  21  22  23  24| 18  19  20  21  22  23  24  |
| 25  26  27  28  29  30  31| 25  26  27  28  29  30      |
|  *선택: 05-01 ~ 05-12*   |                              |
+------------------------+-----------------------+
|  시작일: [2026-05-01]   종료일: [2026-05-12]   |
|  ※ 최대 선택 기간: 90일                        |
|                            [취소]  [확인]       |
+------------------------------------------------+
"""
    add_mockup(doc, mockup_cmn001)

    add_heading(doc, "컴포넌트 설명", 3)
    ct_c1 = add_table_header(doc, ["컴포넌트", "유형", "설명", "비고"], [3.5, 2.5, 5.5, 4.0])
    comp_c1 = [
        ("빠른 선택 버튼", "ButtonGroup", "오늘/이번주/이번달/최근3월 단축 선택", "클릭 시 날짜 자동 설정"),
        ("달력 (2개월)", "Calendar", "시작월/종료월 동시 표시, 드래그 범위 선택", "범위 선택 하이라이트"),
        ("월 이동 버튼", "Button", "◀/▶ 클릭으로 월 변경", "양쪽 달력 연동"),
        ("시작일/종료일 입력", "DateInput", "직접 입력 또는 달력 선택 연동", "YYYY-MM-DD 형식"),
        ("최대 기간 제한", "Validation", "90일 초과 선택 불가 메시지", "화면별 제한 파라미터"),
        ("확인 버튼", "Button", "선택 날짜 호출 화면에 반환", "콜백 함수 호출"),
        ("취소 버튼", "Button", "팝업 닫기, 날짜 변경 없음", "-"),
    ]
    for i, row in enumerate(comp_c1):
        add_table_row(ct_c1, row, alt=(i % 2 == 1))

    doc.add_paragraph()

    # ── SCR-CMN-002 ──────────────────────────────────────────
    add_heading(doc, "5.2 엑셀 다운로드 처리 화면 (SCR-CMN-002)", 2)
    add_label_value(doc, "화면 ID", "SCR-CMN-002")
    add_label_value(doc, "화면명", "엑셀 다운로드 처리")
    add_label_value(doc, "화면 유형", "진행 상태 표시 팝업")

    add_heading(doc, "목업", 3)
    mockup_cmn002 = """\
[엑셀 다운로드 처리]  SCR-CMN-002
+------------------------------------------+
|  엑셀 다운로드                       [X]  |
+------------------------------------------+
|  [다운로드 설정]                          |
|  파일명: [사용자계정관리_20260512.xlsx  ] |
|  다운로드 범위:                           |
|   (*)현재 조회 결과 전체  ( )선택 행만    |
|                                           |
|  포함 항목:  ☑ 헤더행   ☑ 번호열         |
|              ☑ 합계행   □ 숨김 컬럼       |
|                                [다운로드] |
+------------------------------------------+
|  [진행 상태]  (다운로드 클릭 후)          |
|                                           |
|  데이터 조회 중...                        |
|  ████████████████░░░░░░░░  65%  (650/1000)|
|                                           |
|  예상 완료: 약 3초                        |
|                          [취소]           |
+------------------------------------------+
|  [완료]  ✓ 다운로드 완료                 |
|  파일: 사용자계정관리_20260512.xlsx       |
|  크기: 245 KB  /  행수: 1,000건          |
|                     [다시 다운로드] [닫기]|
+------------------------------------------+
"""
    add_mockup(doc, mockup_cmn002)

    add_heading(doc, "컴포넌트 설명", 3)
    ct_c2 = add_table_header(doc, ["컴포넌트", "유형", "설명", "비고"], [3.5, 2.5, 5.5, 4.0])
    comp_c2 = [
        ("파일명 입력", "TextInput", "기본값: 화면명_날짜, 수정 가능", "특수문자 제한"),
        ("범위 라디오", "RadioGroup", "전체/선택행 선택", "-"),
        ("포함 항목 체크", "CheckboxGroup", "헤더/번호/합계/숨김컬럼 선택", "-"),
        ("다운로드 버튼", "Button", "서버 API 호출 → 스트리밍 다운로드", "-"),
        ("진행률 바", "ProgressBar", "0~100% 실시간 표시", "SSE 또는 폴링"),
        ("취소 버튼", "Button", "진행 중 다운로드 취소", "서버 작업 중단"),
        ("완료 상태", "InfoPanel", "파일명/크기/행수 표시", "브라우저 자동 저장"),
    ]
    for i, row in enumerate(comp_c2):
        add_table_row(ct_c2, row, alt=(i % 2 == 1))

    doc.add_paragraph()

    # ── SCR-CMN-003 ──────────────────────────────────────────
    add_heading(doc, "5.3 알람 상세 팝업 (SCR-CMN-003)", 2)
    add_label_value(doc, "화면 ID", "SCR-CMN-003")
    add_label_value(doc, "화면명", "알람 상세 팝업")
    add_label_value(doc, "화면 유형", "알람 상세 정보 및 조치 안내 팝업")

    add_heading(doc, "목업", 3)
    mockup_cmn003 = """\
[알람 상세 팝업]  SCR-CMN-003
+--------------------------------------------------+
|  ⚠ 알람 상세                              [X]   |
+--------------------------------------------------+
|  [알람 목록]  미조치 3건                          |
|  +---+-------+------------------+-------+------+ |
|  |No |유형   |내용              |발생시각|우선순위| |
|  +---+-------+------------------+-------+------+ |
|  |1  |설비고장|세척기#2 트립    |14:20  |긴급  | |
|  |2  |품질이상|불량률 임계 초과  |13:45  |경고  | |
|  |3  |설비점검|충진기#1 점검예정 |09:00  |일반  | |
|  +---+-------+------------------+-------+------+ |
+--------------------------------------------------+
|  [알람 상세]  No.1 선택 시                        |
|  알람 ID   : ALM-2026051214200001                 |
|  유형      : 설비 고장 (긴급)                     |
|  발생 시각 : 2026-05-12 14:20:35                  |
|  대상 설비 : EQP-012 세척기 #2 (생산2라인)        |
|  내용      : 구동부 과부하 트립 발생으로           |
|              자동 정지. 즉시 확인 필요.            |
|  권장 조치 : 1. 전원 차단 확인                    |
|              2. 구동부 육안 점검                   |
|              3. 이상 등록 및 담당자 통보            |
|  관련 화면 : [설비 이상 등록 →]  [가동 현황 →]    |
|                                                   |
|  조치 처리 :                                      |
|   조치자   : [USR004 최설비    v]                 |
|   조치 내용: [즉시 현장 출동 및 점검 중          ] |
|              [알람 조치 완료]  [나중에 처리]       |
+--------------------------------------------------+
"""
    add_mockup(doc, mockup_cmn003)

    add_heading(doc, "컴포넌트 설명", 3)
    ct_c3 = add_table_header(doc, ["컴포넌트", "유형", "설명", "비고"], [3.5, 2.5, 5.5, 4.0])
    comp_c3 = [
        ("알람 목록 그리드", "DataGrid", "미조치 알람 우선순위 정렬 표시", "긴급: 빨강 강조"),
        ("알람 상세 패널", "DetailPanel", "알람 내용 및 권장 조치 표시", "행 클릭 연동"),
        ("관련 화면 링크", "Link", "이상 등록/가동 현황 화면 바로 이동", "화면 ID 파라미터"),
        ("조치자 드롭다운", "Select", "활성 사용자 목록", "-"),
        ("조치 내용 입력", "Textarea", "조치 내용 기록", "-"),
        ("알람 조치 완료 버튼", "Button", "알람 상태 → 조치완료 변경", "조치내용 필수"),
        ("나중에 처리 버튼", "Button", "팝업 닫기, 알람 유지", "-"),
    ]
    for i, row in enumerate(comp_c3):
        add_table_row(ct_c3, row, alt=(i % 2 == 1))

    doc.add_paragraph()

    # ── SCR-CMN-004 ──────────────────────────────────────────
    add_heading(doc, "5.4 확인/취소 공통 팝업 (SCR-CMN-004)", 2)
    add_label_value(doc, "화면 ID", "SCR-CMN-004")
    add_label_value(doc, "화면명", "확인/취소 공통 팝업")
    add_label_value(doc, "화면 유형", "사용자 확인 다이얼로그 (공통 모달)")

    add_heading(doc, "목업", 3)
    mockup_cmn004 = """\
[확인/취소 공통 팝업]  SCR-CMN-004

사용 예시 A: 삭제 확인
+--------------------------------------+
|  ⚠ 삭제 확인                   [X] |
+--------------------------------------+
|                                      |
|  선택한 사용자 계정 1건을            |
|  삭제(비활성화)하시겠습니까?         |
|                                      |
|  대상: USR005 정창고                 |
|                                      |
|  이 작업은 취소할 수 없습니다.       |
|                                      |
|               [취소]    [확인]       |
+--------------------------------------+

사용 예시 B: 저장 확인
+--------------------------------------+
|  ✔ 저장 확인                   [X] |
+--------------------------------------+
|                                      |
|  권한 설정을 저장하시겠습니까?       |
|  변경 역할: 현장관리자               |
|  변경 항목: 시스템관리 > 삭제 권한   |
|                                      |
|               [취소]    [저장]       |
+--------------------------------------+

사용 예시 C: 비밀번호 초기화
+--------------------------------------+
|  ℹ 비밀번호 초기화              [X] |
+--------------------------------------+
|                                      |
|  USR001 김현장의 비밀번호를          |
|  초기화하고 임시 비밀번호를          |
|  이메일로 발송하시겠습니까?          |
|                                      |
|  발송 이메일: kim@imjin.co.kr        |
|                                      |
|               [취소]    [초기화]     |
+--------------------------------------+
"""
    add_mockup(doc, mockup_cmn004)

    add_heading(doc, "컴포넌트 설명 및 파라미터", 3)
    ct_c4 = add_table_header(doc, ["파라미터", "타입", "설명", "예시"], [3.5, 2.5, 5.5, 4.0])
    comp_c4 = [
        ("type", "string", "팝업 유형: warning/success/info/danger", "'warning'"),
        ("title", "string", "팝업 제목", "'삭제 확인'"),
        ("message", "string", "본문 메시지 (HTML 허용)", "'삭제하시겠습니까?'"),
        ("detail", "string", "대상/변경 내용 등 상세 정보", "'대상: USR005'"),
        ("confirmText", "string", "확인 버튼 레이블", "'삭제', '저장', '초기화'"),
        ("cancelText", "string", "취소 버튼 레이블", "'취소'"),
        ("onConfirm", "Function", "확인 버튼 클릭 시 콜백", "() => deleteUser()"),
        ("onCancel", "Function", "취소 버튼 클릭 시 콜백", "() => closeModal()"),
        ("preventClose", "boolean", "X 버튼 비활성화 여부", "false (기본)"),
    ]
    for i, row in enumerate(comp_c4):
        add_table_row(ct_c4, row, alt=(i % 2 == 1))

    page_break(doc)

# ─────────────────────────────────────────────────────────────
# Section 6: Screen Flow Diagrams
# ─────────────────────────────────────────────────────────────
def build_flow_diagrams(doc):
    add_heading(doc, "6. 화면 전환 흐름도", 1)

    add_heading(doc, "6.1 로그인 후 메뉴 접근 흐름", 2)
    flow1 = """\
[로그인 화면]
    |
    | ① 아이디/비밀번호 입력 → 인증 성공
    ▼
[메인 대시보드]
    |
    +── 시스템관리 메뉴 (관리자만 노출) ─────────────────────────┐
    |       |                                                      |
    |       ├─ 사용자 계정 관리 → SCR-SYS-001                     |
    |       ├─ 권한 역할 관리  → SCR-SYS-002                     |
    |       ├─ 시스템 로그 조회 → SCR-SYS-003                    |
    |       └─ 데이터 백업 이력 → SCR-SYS-004                    |
    |                                                              |
    +── KPI 모니터링 메뉴 ────────────────────────────────────────┤
    |       |                                                      |
    |       ├─ 생산성 KPI 조회 → SCR-KPI-001                      |
    |       ├─ 품질 KPI 조회  → SCR-KPI-002                      |
    |       └─ KPI 지표 관리  → SCR-KPI-003 (관리자만)           |
    |                                                              |
    +── 설비관리 메뉴 ───────────────────────────────────────────┘
            |
            ├─ 설비 가동 현황 → SCR-EQP-001
            ├─ 설비 점검 등록 → SCR-EQP-002
            ├─ 설비 점검 이력 → SCR-EQP-003
            ├─ 설비 이상 등록 → SCR-EQP-004
            └─ 설비 고장 이력 → SCR-EQP-005
"""
    add_mockup(doc, flow1)

    add_heading(doc, "6.2 KPI 화면 간 이동 흐름", 2)
    flow2 = """\
[메인 대시보드]  KPI 알림 배너 클릭
    |
    ▼
[생산성 KPI 조회: SCR-KPI-001]
    |
    ├─ 라인별 상세 테이블 행 클릭 → 해당 라인 필터 유지 재조회
    |
    ├─ KPI 카드 (불량률) 클릭 ──────────────────────────────────┐
    |                                                             ▼
    |                                               [품질 KPI 조회: SCR-KPI-002]
    |                                                    |
    |                                                    └─ KPI 코드 수정 필요
    |                                                         ▼
    └─ (관리자) KPI 지표 관리 버튼 ──────────────────► [KPI 지표 관리: SCR-KPI-003]
                                                              |
                                                              └─ 저장 후 → SCR-KPI-001 재조회
"""
    add_mockup(doc, flow2)

    add_heading(doc, "6.3 설비 이상 발생시 화면 흐름", 2)
    flow3 = """\
이상 발생 트리거 (설비 이상 발생)
    |
    ├─ 방법 A: 현장 작업자 직접 등록
    |     ▼
    |   [설비 이상 등록: SCR-EQP-004]
    |       |  ① 설비 선택, 이상 내용 입력
    |       |  ② 알람 발송 대상 선택
    |       |  ③ [이상 등록] 버튼 클릭
    |       ▼
    |   알람 자동 발송 → 설비팀장/현장관리자 수신
    |       ▼
    |   [알람 상세 팝업: SCR-CMN-003]  (수신자 화면)
    |       |  조치 내용 입력 → [알람 조치 완료]
    |       ▼
    |   [설비 가동 현황: SCR-EQP-001]  상태 자동 갱신 (고장 → 복구)
    |       ▼
    |   [설비 고장 이력: SCR-EQP-005]  이력 자동 등록
    |
    └─ 방법 B: 설비 점검 중 이상 발견
          ▼
        [설비 점검 등록: SCR-EQP-002]
            |  이상 발견 항목 선택
            |  [설비 이상 등록 연계] 체크박스 선택
            |  [완료저장] 클릭
            ▼
        SCR-EQP-004 이상 등록 자동 생성 (연계)
            |
            ▼ (방법 A 흐름과 합류)
        알람 발송 → SCR-CMN-003 → SCR-EQP-001 갱신 → SCR-EQP-005 이력

공통 팝업 연동:
  ─ 날짜 조건 입력 시 → SCR-CMN-001 (날짜 범위 선택)
  ─ 저장/삭제/초기화 → SCR-CMN-004 (확인/취소 팝업)
  ─ 목록 엑셀 다운 → SCR-CMN-002 (엑셀 다운로드 처리)
"""
    add_mockup(doc, flow3)

    doc.add_paragraph()
    add_heading(doc, "6.4 화면 ID 및 API 엔드포인트 매핑", 2)
    api_headers = ["화면 ID", "주요 API 엔드포인트", "HTTP 메서드", "비고"]
    api_widths = [3.0, 5.5, 2.5, 4.5]
    at = add_table_header(doc, api_headers, api_widths)
    api_data = [
        ("SCR-SYS-001", "/api/users", "GET, POST, PUT, DELETE", "페이징 지원"),
        ("SCR-SYS-002", "/api/roles, /api/roles/{id}/permissions", "GET, PUT", "-"),
        ("SCR-SYS-003", "/api/logs, /api/logs/export", "GET", "읽기 전용"),
        ("SCR-SYS-004", "/api/backups, /api/backups/run", "GET, POST", "-"),
        ("SCR-KPI-001", "/api/kpi/productivity, /api/kpi/productivity/export", "GET", "집계 파라미터"),
        ("SCR-KPI-002", "/api/kpi/quality, /api/kpi/quality/export", "GET", "집계 파라미터"),
        ("SCR-KPI-003", "/api/kpi/indicators", "GET, POST, PUT, DELETE", "-"),
        ("SCR-EQP-001", "/api/equipment/status", "GET", "30초 폴링"),
        ("SCR-EQP-002", "/api/inspections", "GET, POST, PUT", "임시저장 포함"),
        ("SCR-EQP-003", "/api/inspections, /api/inspections/export", "GET", "-"),
        ("SCR-EQP-004", "/api/anomalies", "POST, PUT", "알람 발송 연동"),
        ("SCR-EQP-005", "/api/anomalies, /api/anomalies/export", "GET, PUT", "복구완료 처리"),
        ("SCR-CMN-001", "—", "—", "프론트엔드 전용"),
        ("SCR-CMN-002", "/api/*/export", "GET (Streaming)", "SSE 기반"),
        ("SCR-CMN-003", "/api/alarms, /api/alarms/{id}/resolve", "GET, PUT", "-"),
        ("SCR-CMN-004", "—", "—", "프론트엔드 전용"),
    ]
    for i, row in enumerate(api_data):
        add_table_row(at, row, alt=(i % 2 == 1))

# ─────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────
def main():
    doc = Document()

    # Page margins
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.orientation = 1  # landscape
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)

    # Default font
    from docx.oxml.ns import qn
    doc.styles['Normal'].font.name = 'Malgun Gothic'
    doc.styles['Normal'].font.size = Pt(10)

    print("표지 작성 중...")
    build_cover(doc)

    print("1장 개요 작성 중...")
    build_overview(doc)

    print("2장 시스템관리 화면 작성 중...")
    build_sys_screens(doc)

    print("3장 KPI 모니터링 화면 작성 중...")
    build_kpi_screens(doc)

    print("4장 설비관리 화면 작성 중...")
    build_eqp_screens(doc)

    print("5장 공통 컴포넌트 작성 중...")
    build_common_screens(doc)

    print("6장 화면 전환 흐름도 작성 중...")
    build_flow_diagrams(doc)

    print(f"저장 중: {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print("완료!")

if __name__ == "__main__":
    main()
