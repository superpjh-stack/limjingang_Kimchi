"""
임진강김치 MES 시스템 구축 - 프로그램설계서1 생성 스크립트
대상 모듈: 기준정보관리(BM), 수주생산계획관리(PM), 자재재고관리(INV), 입고전처리관리(RCV), 세척공정관리(WSH), 절임공정관리(PCK)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"C:\gerardo\01 SmallSF\Imjingang-Kimchi\산출물\설계단계\07_프로그램설계서1.docx"

# ──────────────────────────────────────────────
# 헬퍼 함수
# ──────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_col_width(table, col_idx: int, width_cm: float):
    """특정 열 너비 설정"""
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def add_cover_page(doc: Document):
    """표지 작성"""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('프로그램설계서 1')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('임진강김치 MES 시스템 구축')
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x26, 0x26, 0x26)

    doc.add_paragraph()
    doc.add_paragraph()

    info_lines = [
        ('작성일', '2026-05-12'),
        ('버전', 'V1.0'),
        ('작성자', '설계자2'),
        ('대상 모듈', 'BM / PM / INV / RCV / WSH / PCK'),
    ]
    tbl = doc.add_table(rows=len(info_lines), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    for i, (k, v) in enumerate(info_lines):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl.rows[i].cells[0], 'D9E1F2')
        tbl.rows[i].cells[1].text = v
    for row in tbl.rows:
        for cell in row.cells:
            cell.width = Cm(7)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(10)
    return p


def add_api_table(doc: Document, apis: list):
    """
    apis: list of dict {method, url, request, response, desc}
    """
    headers = ['Method', 'URL', 'Request (주요 파라미터)', 'Response (주요 필드)', '설명']
    tbl = doc.add_table(rows=1 + len(apis), cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 헤더
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr[i], '1F497D')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 데이터
    widths = [2.0, 5.5, 5.5, 5.5, 3.5]
    for r_idx, api in enumerate(apis, 1):
        row = tbl.rows[r_idx].cells
        row[0].text = api.get('method', '')
        row[1].text = api.get('url', '')
        row[2].text = api.get('request', '-')
        row[3].text = api.get('response', '-')
        row[4].text = api.get('desc', '')
        # Method 색
        method_colors = {'GET': 'E2EFDA', 'POST': 'FFF2CC', 'PUT': 'DDEBF7', 'DELETE': 'FCE4D6', 'PATCH': 'EAD1DC'}
        bg = method_colors.get(api.get('method', ''), 'FFFFFF')
        set_cell_bg(row[0], bg)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if r_idx % 2 == 0:
            set_cell_bg(row[1], 'F2F2F2')
            set_cell_bg(row[2], 'F2F2F2')
            set_cell_bg(row[3], 'F2F2F2')
            set_cell_bg(row[4], 'F2F2F2')

    for col_idx, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[col_idx].width = Cm(w)

    doc.add_paragraph()


def add_program_section(doc: Document, prog_id: str, prog_name: str, desc: str, flow: list, apis: list, logic: str, exception: str):
    """개별 프로그램 설계 블록"""
    # 프로그램 정보 테이블
    info = [
        ('프로그램 ID', prog_id),
        ('프로그램명', prog_name),
        ('기능 설명', desc),
    ]
    tbl = doc.add_table(rows=len(info), cols=2)
    tbl.style = 'Table Grid'
    for i, (k, v) in enumerate(info):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl.rows[i].cells[0], 'D9E1F2')
        tbl.rows[i].cells[1].text = v
        tbl.rows[i].cells[0].width = Cm(3.5)
        tbl.rows[i].cells[1].width = Cm(13.5)
    doc.add_paragraph()

    # 처리 흐름
    p = doc.add_paragraph()
    r = p.add_run('■ 처리 흐름')
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    for step in flow:
        bp = doc.add_paragraph(step, style='List Number')
        bp.paragraph_format.left_indent = Cm(0.5)

    doc.add_paragraph()

    # API 설계
    p = doc.add_paragraph()
    r = p.add_run('■ API 설계')
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    add_api_table(doc, apis)

    # 주요 로직
    p = doc.add_paragraph()
    r = p.add_run('■ 주요 로직 설명')
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    doc.add_paragraph(logic)
    doc.add_paragraph()

    # 예외 처리
    p = doc.add_paragraph()
    r = p.add_run('■ 예외 처리')
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    doc.add_paragraph(exception)
    doc.add_paragraph()

    # 구분선
    hr = doc.add_paragraph('─' * 80)
    hr.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()


# ──────────────────────────────────────────────
# 1장. 개요
# ──────────────────────────────────────────────

def add_chapter1_overview(doc: Document):
    add_heading(doc, '1. 개요', 1)

    add_heading(doc, '1.1 문서 목적 및 범위', 2)
    add_body(doc,
        '본 문서는 임진강김치 MES(Manufacturing Execution System) 시스템 구축 프로젝트의 프로그램설계서 1편으로, '
        '기준정보관리(BM), 수주생산계획관리(PM), 자재재고관리(INV), 입고전처리관리(RCV), '
        '세척공정관리(WSH), 절임공정관리(PCK) 모듈에 대한 프로그램별 상세 설계를 기술한다.')
    add_body(doc,
        '본 문서는 개발자가 실제 코딩을 수행하기 위한 기준 문서로 활용되며, '
        'RESTful API 명세, 처리 흐름, 주요 로직, 예외 처리 방안을 포함한다.')

    add_heading(doc, '1.2 프로그램 설계 규칙', 2)

    rules = [
        ('패키지 구조', 'com.imjingang.mes.{module}.{layer}\n예) com.imjingang.mes.bm.controller, com.imjingang.mes.bm.service, com.imjingang.mes.bm.repository'),
        ('컨트롤러 네이밍', '{ModuleName}Controller.java\n예) ProductController.java, RecipeController.java'),
        ('서비스 네이밍', '{ModuleName}Service.java / {ModuleName}ServiceImpl.java'),
        ('Repository', '{ModuleName}Repository.java (JPA 기반 인터페이스)'),
        ('DTO', '{ModuleName}RequestDto.java / {ModuleName}ResponseDto.java'),
        ('Entity', '{ModuleName}.java (JPA @Entity)'),
        ('API URL 규칙', '/api/v1/{module}/{resource} (소문자, 복수형)\n예) /api/v1/bm/products, /api/v1/pm/orders'),
        ('HTTP 메서드', 'GET(조회), POST(등록), PUT(전체수정), PATCH(부분수정), DELETE(삭제)'),
        ('응답 형식', '{"success": true/false, "data": {...}, "message": "...", "code": "SUCCESS|ERROR_CODE"}'),
        ('공통 헤더', 'Authorization: Bearer {JWT_TOKEN}, Content-Type: application/json'),
    ]
    tbl = doc.add_table(rows=1 + len(rules), cols=2)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    hdr[0].text = '항목'
    hdr[1].text = '규칙'
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_bg(cell, '1F497D')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, (k, v) in enumerate(rules, 1):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl.rows[i].cells[0], 'D9E1F2')
        tbl.rows[i].cells[1].text = v
        tbl.rows[i].cells[0].width = Cm(4.0)
        tbl.rows[i].cells[1].width = Cm(13.0)
    doc.add_paragraph()

    add_heading(doc, '1.3 공통 설계 사항', 2)

    add_body(doc, '■ 에러 처리 방침')
    errors = [
        ('400 Bad Request', '입력값 유효성 검사 실패 - @Valid + BindingResult 활용'),
        ('401 Unauthorized', 'JWT 토큰 만료/미제출 - JwtAuthenticationFilter에서 처리'),
        ('403 Forbidden', '권한 부족 - @PreAuthorize + Spring Security'),
        ('404 Not Found', '리소스 미존재 - GlobalExceptionHandler에서 일괄 처리'),
        ('409 Conflict', '중복 데이터 - 비즈니스 로직에서 DuplicateException 발생'),
        ('500 Internal Server Error', '예상치 못한 오류 - GlobalExceptionHandler + Slack 알림'),
    ]
    tbl2 = doc.add_table(rows=1 + len(errors), cols=2)
    tbl2.style = 'Table Grid'
    h2 = tbl2.rows[0].cells
    h2[0].text = 'HTTP 상태코드'
    h2[1].text = '처리 방침'
    for cell in h2:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_bg(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, (k, v) in enumerate(errors, 1):
        tbl2.rows[i].cells[0].text = k
        tbl2.rows[i].cells[1].text = v
        tbl2.rows[i].cells[0].width = Cm(4.5)
        tbl2.rows[i].cells[1].width = Cm(12.5)
    doc.add_paragraph()

    add_body(doc, '■ 로깅 전략')
    logging_text = (
        '- 프레임워크: SLF4J + Logback\n'
        '- 레벨: ERROR(운영), WARN(주의), INFO(업무 처리), DEBUG(개발)\n'
        '- AOP 기반 메서드 실행 시간 측정 (@Around 어드바이스)\n'
        '- MDC(Mapped Diagnostic Context)로 요청 추적 ID 삽입\n'
        '- 로그 파일: /logs/mes-{module}-{date}.log (30일 보관)'
    )
    add_body(doc, logging_text)
    doc.add_paragraph()

    add_body(doc, '■ 트랜잭션 처리')
    tx_text = (
        '- @Transactional(readOnly = true): 조회 메서드 기본 적용\n'
        '- @Transactional: 등록/수정/삭제 처리\n'
        '- 분산 트랜잭션(재고↔생산): @Transactional(propagation = REQUIRES_NEW) 활용\n'
        '- 배치 처리: ChunkOrientedTasklet, 5,000건 단위 커밋'
    )
    add_body(doc, tx_text)
    doc.add_paragraph()
    doc.add_page_break()


# ──────────────────────────────────────────────
# 2장. 기준정보관리 (BM)
# ──────────────────────────────────────────────

def add_chapter2_bm(doc: Document):
    add_heading(doc, '2. 기준정보관리 (BM) 프로그램 설계', 1)

    add_body(doc,
        '기준정보관리(BM, Base Master) 모듈은 MES 시스템 전체에서 참조하는 마스터 데이터를 등록·관리한다. '
        '제품 품목, 레시피 BOM, 공정 CCP, 설비/탱크, 거래처, 작업자, 공통코드 등 7개 기능으로 구성된다.')
    doc.add_paragraph()

    # 2.1 제품품목기준등록조회
    add_heading(doc, '2.1 제품품목기준등록조회', 2)
    add_program_section(
        doc,
        prog_id='BM-001',
        prog_name='제품품목기준등록조회',
        desc='김치 제품의 품목 마스터를 등록·수정·삭제·조회한다. 제품코드, 제품명, 규격, 단위, 유통기한, 보관온도, 포장형태 등을 관리한다.',
        flow=[
            'React 화면에서 조회 조건(제품코드, 제품명, 제품유형) 입력',
            'GET /api/v1/bm/products 호출 → ProductController.getProducts()',
            'ProductService.findProducts() → ProductRepository.findByCondition() 실행',
            'Page<Product> 결과를 ProductResponseDto로 변환 후 JSON 반환',
            '등록 시: POST /api/v1/bm/products → @Valid 유효성 검사 → 중복코드 확인 → save()',
            '수정 시: PUT /api/v1/bm/products/{id} → 존재 여부 확인 → update() → 변경 이력 기록',
            '삭제 시: DELETE /api/v1/bm/products/{id} → 참조 무결성 확인(레시피BOM 연관) → soft delete 처리',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/bm/products', 'request': 'productCode(선택), productName(선택), productType(선택), page, size', 'response': 'productId, productCode, productName, unit, spec, shelfLifeDays, storageTemp, packageType, useYn', 'desc': '제품 목록 조회 (페이징)'},
            {'method': 'GET', 'url': '/api/v1/bm/products/{id}', 'request': 'id (Path)', 'response': '제품 전체 상세 정보 + 연관 레시피 목록', 'desc': '제품 단건 조회'},
            {'method': 'POST', 'url': '/api/v1/bm/products', 'request': 'productCode*, productName*, unit*, spec, shelfLifeDays*, storageTemp, packageType*, useYn', 'response': 'productId, productCode, createdAt', 'desc': '제품 등록'},
            {'method': 'PUT', 'url': '/api/v1/bm/products/{id}', 'request': 'productName, unit, spec, shelfLifeDays, storageTemp, packageType, useYn', 'response': 'productId, updatedAt', 'desc': '제품 수정'},
            {'method': 'DELETE', 'url': '/api/v1/bm/products/{id}', 'request': 'id (Path)', 'response': 'success: true', 'desc': '제품 삭제 (소프트삭제)'},
        ],
        logic=(
            '- 제품코드는 시스템 자동생성 규칙 적용: {제품유형코드}-{일련번호4자리} (예: KIM-0001)\n'
            '- 유통기한(shelfLifeDays)은 양수만 허용, 최대 365일 제한\n'
            '- 보관온도(storageTemp)는 -25 ~ 10°C 범위 유효성 검사\n'
            '- 삭제 시 레시피BOM, 생산계획에서 참조 중인 경우 삭제 불가 예외 발생\n'
            '- 변경 이력은 product_change_history 테이블에 자동 기록 (Hibernate Envers 활용)'
        ),
        exception=(
            '- DuplicateProductCodeException: 동일 제품코드 등록 시도 시 409 반환\n'
            '- ProductInUseException: 참조 중인 제품 삭제 시도 시 409 반환\n'
            '- InvalidTemperatureRangeException: 보관온도 범위 초과 시 400 반환\n'
            '- ProductNotFoundException: 존재하지 않는 제품 수정/삭제 시 404 반환'
        )
    )

    # 2.2 레시피BOM관리
    add_heading(doc, '2.2 레시피BOM관리', 2)
    add_program_section(
        doc,
        prog_id='BM-002',
        prog_name='레시피BOM관리',
        desc='제품별 레시피(원·부자재 구성 및 배합비율)를 등록·관리한다. 다단계 BOM 구조를 지원하며, 버전 관리를 통해 레시피 변경 이력을 추적한다.',
        flow=[
            '제품 선택 후 레시피 트리 구조 조회 요청',
            'GET /api/v1/bm/recipes?productId={id} → RecipeController',
            'RecipeService.getRecipeTree() → 재귀 BOM 조회 및 계층 구조 변환',
            '레시피 등록: 상위 레시피 코드, 자재 목록, 배합비율 입력',
            'POST /api/v1/bm/recipes → 총 배합비율 합계 100% 검증',
            '레시피 버전 변경: 기존 레시피 복사 후 신규 버전 생성, 이전 버전 비활성화',
            '소요량 계산 시 레시피 기준 자동 연산하여 자재재고관리 모듈에 연동',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/bm/recipes', 'request': 'productId(필수), version(선택)', 'response': 'recipeId, productId, version, status, recipeItems[{materialId, materialName, quantity, unit, ratio}]', 'desc': '레시피 BOM 조회'},
            {'method': 'GET', 'url': '/api/v1/bm/recipes/{id}/versions', 'request': 'id (Path)', 'response': '버전 목록 및 변경 이력', 'desc': '레시피 버전 이력 조회'},
            {'method': 'POST', 'url': '/api/v1/bm/recipes', 'request': 'productId*, version*, items[{materialId, quantity, unit, ratio}]*', 'response': 'recipeId, version, createdAt', 'desc': '레시피 등록'},
            {'method': 'PUT', 'url': '/api/v1/bm/recipes/{id}', 'request': 'items[{materialId, quantity, unit, ratio}]', 'response': 'recipeId, updatedAt', 'desc': '레시피 수정'},
            {'method': 'POST', 'url': '/api/v1/bm/recipes/{id}/copy', 'request': 'newVersion', 'response': 'newRecipeId, version', 'desc': '레시피 버전 복사'},
        ],
        logic=(
            '- BOM은 최대 5단계 depth 지원, 재귀 CTE 쿼리로 조회 성능 최적화\n'
            '- 배합비율 합계 유효성: 소수점 처리 오차 허용 범위 ±0.01% 이내\n'
            '- 레시피 버전은 ACTIVE/INACTIVE/DRAFT 상태로 관리\n'
            '- 생산계획 수립 시 ACTIVE 상태의 최신 버전 자동 적용\n'
            '- Redis Cache (@Cacheable) 적용: "recipe::{productId}" 키로 5분 TTL 캐싱'
        ),
        exception=(
            '- RecipeRatioSumException: 배합비율 합계 != 100% 시 400 반환\n'
            '- RecipeVersionConflictException: 동일 버전 중복 등록 시 409 반환\n'
            '- CircularReferenceException: BOM 순환 참조 감지 시 400 반환\n'
            '- RecipeActiveException: ACTIVE 레시피 직접 삭제 시도 시 409 반환'
        )
    )

    # 2.3 공정CCP기준관리
    add_heading(doc, '2.3 공정CCP기준관리', 2)
    add_program_section(
        doc,
        prog_id='BM-003',
        prog_name='공정CCP기준관리',
        desc='HACCP 기반 공정별 CCP(Critical Control Point) 관리 기준값을 등록·관리한다. 세척, 절임, 버무림 등 각 공정의 온도, 농도, 시간 기준과 허용 범위를 설정한다.',
        flow=[
            '공정 선택 후 CCP 기준 목록 조회',
            'GET /api/v1/bm/ccps?processCode={code} → CcpController',
            'CCP 항목별 기준값(최솟값, 최댓값, 목표값) 설정',
            'POST /api/v1/bm/ccps → 유효성 검사 (min <= target <= max)',
            '공정 실적 데이터 수집 시 CCP 기준 초과 여부 자동 판별',
            '기준값 변경 시 변경 이력 기록 및 품질이상관리 모듈 연동',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/bm/ccps', 'request': 'processCode(선택), ccpType(선택)', 'response': 'ccpId, processCode, ccpName, ccpType, minValue, maxValue, targetValue, unit, actionLevel', 'desc': 'CCP 기준 목록 조회'},
            {'method': 'GET', 'url': '/api/v1/bm/ccps/{id}', 'request': 'id (Path)', 'response': 'CCP 상세 정보 + 측정 이력', 'desc': 'CCP 기준 단건 조회'},
            {'method': 'POST', 'url': '/api/v1/bm/ccps', 'request': 'processCode*, ccpName*, ccpType*, minValue*, maxValue*, targetValue*, unit*, actionLevel', 'response': 'ccpId, createdAt', 'desc': 'CCP 기준 등록'},
            {'method': 'PUT', 'url': '/api/v1/bm/ccps/{id}', 'request': 'minValue, maxValue, targetValue, actionLevel', 'response': 'ccpId, updatedAt', 'desc': 'CCP 기준 수정'},
        ],
        logic=(
            '- CCP 유형: TEMPERATURE(온도), CONCENTRATION(농도), TIME(시간), WEIGHT(중량)\n'
            '- 관리 기준: minValue ≤ targetValue ≤ maxValue 제약 조건 검증\n'
            '- 경보 기준(actionLevel): CCP 이탈 시 알람 발생 임계값 별도 관리\n'
            '- 설비 데이터(MQTT) 수신 시 Redis 저장 기준값과 실시간 비교하여 이상 감지\n'
            '- CCP 이탈 기록은 품질이상관리(QC) 모듈의 이슈이력에 자동 연동'
        ),
        exception=(
            '- InvalidCcpRangeException: min > max 또는 target 범위 이탈 시 400 반환\n'
            '- ProcessNotFoundException: 존재하지 않는 공정코드 참조 시 404 반환\n'
            '- CcpDuplicateException: 동일 공정 내 동일 CCP명 중복 시 409 반환'
        )
    )

    # 2.4 설비탱크기준관리
    add_heading(doc, '2.4 설비탱크기준관리', 2)
    add_program_section(
        doc,
        prog_id='BM-004',
        prog_name='설비탱크기준관리',
        desc='생산 현장의 설비(세척기, 충진기, 테이핑기 등) 및 탱크(절임통, 양념통 등) 마스터 정보를 등록·관리한다.',
        flow=[
            'GET /api/v1/bm/equipments - 설비 목록 조회 (공정별, 라인별 필터)',
            '설비 상세: 설비코드, 설비명, 설비유형, 제조사, 설치일, 최대용량 등 조회',
            'POST /api/v1/bm/equipments - 신규 설비 등록',
            'MQTT topic: mes/equipment/{eqpCode}/status - 설비 상태 실시간 수신',
            '탱크 관리: 탱크코드, 탱크명, 최대용량(kg), 현재용량, 온도센서ID 등록',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/bm/equipments', 'request': 'processCode(선택), lineNo(선택), eqpType(선택)', 'response': 'eqpId, eqpCode, eqpName, eqpType, lineNo, capacity, status, installDate', 'desc': '설비 목록 조회'},
            {'method': 'POST', 'url': '/api/v1/bm/equipments', 'request': 'eqpCode*, eqpName*, eqpType*, lineNo*, capacity, manufacturer, installDate, sensorTopicId', 'response': 'eqpId, eqpCode', 'desc': '설비 등록'},
            {'method': 'PUT', 'url': '/api/v1/bm/equipments/{id}', 'request': '설비 수정 필드', 'response': 'eqpId, updatedAt', 'desc': '설비 수정'},
            {'method': 'GET', 'url': '/api/v1/bm/tanks', 'request': 'processCode(선택)', 'response': 'tankId, tankCode, tankName, maxCapacity, currentVolume, tempSensorId, status', 'desc': '탱크 목록 조회'},
            {'method': 'POST', 'url': '/api/v1/bm/tanks', 'request': 'tankCode*, tankName*, maxCapacity*, processCode*, tempSensorId', 'response': 'tankId, tankCode', 'desc': '탱크 등록'},
        ],
        logic=(
            '- 설비 유형 분류: WASHING(세척), PICKLING(절임), MIXING(버무림), PACKING(포장), TAPING(테이핑)\n'
            '- MQTT 토픽 구조: mes/equipment/{eqpCode}/{dataType} (status, temperature, vibration)\n'
            '- 센서 토픽 ID는 MQTT 브로커에 등록된 실제 토픽명과 1:1 매핑\n'
            '- 탱크 현재용량은 입고/출고 이벤트 발생 시 자동 업데이트\n'
            '- 설비 이상 감지 시 설비관리(EQP) 모듈의 설비이상등록에 자동 연동'
        ),
        exception=(
            '- DuplicateEquipmentCodeException: 설비코드 중복 시 409 반환\n'
            '- InvalidCapacityException: 현재용량 > 최대용량 시 400 반환\n'
            '- SensorTopicNotFoundException: MQTT 토픽 미등록 상태에서 연결 시도 시 404 반환'
        )
    )

    # 2.5 거래처관리
    add_heading(doc, '2.5 거래처관리', 2)
    add_program_section(
        doc,
        prog_id='BM-005',
        prog_name='거래처관리',
        desc='원자재 공급업체, 포장재 공급업체, 판매처(납품처) 등 거래처 마스터를 등록·관리한다.',
        flow=[
            'GET /api/v1/bm/partners - 거래처 목록 조회 (유형별, 지역별 필터)',
            '거래처 등록: 사업자번호 중복 확인 → POST /api/v1/bm/partners',
            '거래처 유형: SUPPLIER(공급업체), CUSTOMER(판매처), BOTH(겸용)',
            '담당자 정보(이름, 전화번호, 이메일) 함께 관리',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/bm/partners', 'request': 'partnerType(선택), partnerName(선택), bizNo(선택)', 'response': 'partnerId, partnerCode, partnerName, partnerType, bizNo, ceoName, address, tel, email', 'desc': '거래처 목록 조회'},
            {'method': 'POST', 'url': '/api/v1/bm/partners', 'request': 'partnerName*, partnerType*, bizNo*, ceoName, address*, tel*, email, bankInfo', 'response': 'partnerId, partnerCode', 'desc': '거래처 등록'},
            {'method': 'PUT', 'url': '/api/v1/bm/partners/{id}', 'request': '수정 필드', 'response': 'partnerId, updatedAt', 'desc': '거래처 수정'},
            {'method': 'DELETE', 'url': '/api/v1/bm/partners/{id}', 'request': 'id (Path)', 'response': 'success: true', 'desc': '거래처 삭제 (소프트삭제)'},
        ],
        logic=(
            '- 사업자번호 형식 검증: 10자리 숫자, 체크디지트 알고리즘 적용\n'
            '- 거래처코드 자동생성: {유형코드}-{일련번호4자리} (예: SUP-0001, CUS-0001)\n'
            '- 수주 또는 입고 이력이 있는 거래처는 삭제 불가 (소프트삭제만 허용)\n'
            '- 이메일 주소 유효성 검사: RFC 5322 표준 준수'
        ),
        exception=(
            '- DuplicateBizNoException: 사업자번호 중복 등록 시 409 반환\n'
            '- PartnerInUseException: 거래 이력 있는 거래처 삭제 시도 시 409 반환\n'
            '- InvalidBizNoFormatException: 사업자번호 형식 오류 시 400 반환'
        )
    )

    # 2.6 작업자관리
    add_heading(doc, '2.6 작업자관리', 2)
    add_program_section(
        doc,
        prog_id='BM-006',
        prog_name='작업자관리',
        desc='현장 작업자 마스터를 관리한다. 작업자 기본 정보, 담당 공정, 자격증, 재직 상태 등을 등록·조회한다.',
        flow=[
            'GET /api/v1/bm/workers - 작업자 목록 조회',
            '신규 작업자 등록: 개인정보 입력 → 담당공정 선택 → POST /api/v1/bm/workers',
            '퇴직 처리: PATCH /api/v1/bm/workers/{id}/status → 재직상태 INACTIVE 변경',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/bm/workers', 'request': 'workerName(선택), processCode(선택), status(선택: ACTIVE/INACTIVE)', 'response': 'workerId, workerCode, workerName, processCode, hireDate, status', 'desc': '작업자 목록 조회'},
            {'method': 'POST', 'url': '/api/v1/bm/workers', 'request': 'workerName*, gender, birthDate, hireDate*, processCode*, certifications', 'response': 'workerId, workerCode', 'desc': '작업자 등록'},
            {'method': 'PUT', 'url': '/api/v1/bm/workers/{id}', 'request': '수정 필드', 'response': 'workerId, updatedAt', 'desc': '작업자 수정'},
            {'method': 'PATCH', 'url': '/api/v1/bm/workers/{id}/status', 'request': 'status (ACTIVE/INACTIVE), retireDate', 'response': 'workerId, status', 'desc': '작업자 재직상태 변경'},
        ],
        logic=(
            '- 작업자코드 자동생성: WRK-{일련번호4자리}\n'
            '- 개인정보(생년월일) 암호화 저장: AES-256-CBC\n'
            '- 담당공정은 다중 선택 가능 (worker_process 연결 테이블)\n'
            '- POP 작업 화면에서 작업자 인식: QR코드 또는 바코드 스캔 방식'
        ),
        exception=(
            '- WorkerNotFoundException: 존재하지 않는 작업자 조회 시 404 반환\n'
            '- AlreadyInactiveException: 이미 퇴직 처리된 작업자 재퇴직 시도 시 409 반환'
        )
    )

    # 2.7 공통코드관리
    add_heading(doc, '2.7 공통코드관리', 2)
    add_program_section(
        doc,
        prog_id='BM-007',
        prog_name='공통코드관리',
        desc='시스템 전반에서 사용하는 공통코드(그룹코드 + 상세코드) 체계를 관리한다. 드롭다운, 셀렉트박스 등 UI 요소에 동적으로 적용된다.',
        flow=[
            'GET /api/v1/bm/codes?groupCode={code} - 그룹코드별 상세코드 목록 조회',
            '코드 그룹 등록: POST /api/v1/bm/code-groups',
            '상세 코드 등록: POST /api/v1/bm/codes',
            'Redis 캐싱: 공통코드 변경 시 캐시 즉시 갱신 (Cache Evict)',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/bm/code-groups', 'request': '-', 'response': 'groupCode, groupName, useYn, codes[{codeId, codeValue, codeName, sortOrder}]', 'desc': '코드 그룹 전체 조회'},
            {'method': 'GET', 'url': '/api/v1/bm/codes', 'request': 'groupCode(필수)', 'response': 'codeId, groupCode, codeValue, codeName, sortOrder, useYn', 'desc': '그룹별 상세코드 조회'},
            {'method': 'POST', 'url': '/api/v1/bm/code-groups', 'request': 'groupCode*, groupName*, useYn', 'response': 'groupCode, createdAt', 'desc': '코드 그룹 등록'},
            {'method': 'POST', 'url': '/api/v1/bm/codes', 'request': 'groupCode*, codeValue*, codeName*, sortOrder, useYn', 'response': 'codeId, groupCode, codeValue', 'desc': '상세코드 등록'},
            {'method': 'PUT', 'url': '/api/v1/bm/codes/{id}', 'request': 'codeName, sortOrder, useYn', 'response': 'codeId, updatedAt', 'desc': '상세코드 수정'},
        ],
        logic=(
            '- Redis 캐시 키: "commonCode::{groupCode}", TTL 1시간\n'
            '- 코드값(codeValue)은 영문 대문자+언더스코어만 허용 (예: PRODUCT_TYPE)\n'
            '- 사용중인 코드(useYn=N) 비활성화 시 연관 데이터 영향도 사전 확인 팝업\n'
            '- 정렬순서(sortOrder) 미입력 시 최대값+1 자동 부여'
        ),
        exception=(
            '- DuplicateCodeException: 동일 그룹 내 중복 코드값 등록 시 409 반환\n'
            '- SystemCodeModifyException: 시스템 예약 코드 수정/삭제 시도 시 403 반환'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# 3장. 수주생산계획관리 (PM)
# ──────────────────────────────────────────────

def add_chapter3_pm(doc: Document):
    add_heading(doc, '3. 수주생산계획관리 (PM) 프로그램 설계', 1)
    add_body(doc,
        '수주생산계획관리(PM, Production Planning Management) 모듈은 고객 수주 정보를 기반으로 '
        '생산계획을 수립하고, 작업지시서를 생성하여 현장에 배포하는 기능을 제공한다.')
    doc.add_paragraph()

    # 3.1 수주등록조회
    add_heading(doc, '3.1 수주등록조회', 2)
    add_program_section(
        doc,
        prog_id='PM-001',
        prog_name='수주등록조회',
        desc='고객(판매처)으로부터의 수주 정보를 등록하고 조회한다. 수주번호, 거래처, 제품, 수량, 납기일, 납품지 등을 관리한다.',
        flow=[
            '수주 목록 조회: GET /api/v1/pm/orders (기간, 거래처, 상태 필터)',
            '수주 등록: 거래처 선택 → 제품/수량/납기일 입력 → POST /api/v1/pm/orders',
            '수주번호 자동생성: ORD-{YYYYMMDD}-{일련번호4자리}',
            '수주 확정: PATCH /api/v1/pm/orders/{id}/confirm → 생산계획 연동 가능 상태로 전환',
            '수주 취소: PATCH /api/v1/pm/orders/{id}/cancel → 취소 사유 필수 입력',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/pm/orders', 'request': 'startDate, endDate, partnerId, status, page, size', 'response': 'orderId, orderNo, partnerName, productName, orderQty, deliveryDate, status', 'desc': '수주 목록 조회'},
            {'method': 'GET', 'url': '/api/v1/pm/orders/{id}', 'request': 'id (Path)', 'response': '수주 전체 상세 + 수주품목 목록 + 변경이력', 'desc': '수주 단건 조회'},
            {'method': 'POST', 'url': '/api/v1/pm/orders', 'request': 'partnerId*, deliveryDate*, orderItems[{productId, qty, unitPrice}]*', 'response': 'orderId, orderNo, totalAmount', 'desc': '수주 등록'},
            {'method': 'PATCH', 'url': '/api/v1/pm/orders/{id}/confirm', 'request': 'id (Path)', 'response': 'orderId, status: CONFIRMED', 'desc': '수주 확정'},
            {'method': 'PATCH', 'url': '/api/v1/pm/orders/{id}/cancel', 'request': 'cancelReason*', 'response': 'orderId, status: CANCELLED', 'desc': '수주 취소'},
        ],
        logic=(
            '- 수주 상태 흐름: DRAFT → CONFIRMED → PLANNED → PRODUCING → SHIPPED → COMPLETED\n'
            '- 납기일 기준 리드타임 역산: 생산계획 시작일 = 납기일 - 기준리드타임(일)\n'
            '- 수주 확정 시 재고 가용량 자동 확인 (현재고 - 예약재고 >= 필요수량)\n'
            '- 수주 금액 자동계산: Σ(수량 × 단가) → 세금계산서 금액 기준'
        ),
        exception=(
            '- OrderStatusTransitionException: 허용되지 않는 상태 전환 시 400 반환\n'
            '- DeliveryDatePastException: 납기일이 오늘 이전인 경우 400 반환\n'
            '- PartnerInactiveException: 비활성 거래처로의 수주 등록 시 400 반환'
        )
    )

    # 3.2 수주변경이력관리
    add_heading(doc, '3.2 수주변경이력관리', 2)
    add_program_section(
        doc,
        prog_id='PM-002',
        prog_name='수주변경이력관리',
        desc='수주 정보 변경(수량 조정, 납기 변경, 제품 교체) 이력을 추적하고 조회한다.',
        flow=[
            '수주 수정 요청: PUT /api/v1/pm/orders/{id} → 변경 전/후 자동 diff 기록',
            '변경이력 조회: GET /api/v1/pm/orders/{id}/histories',
            '변경 유형: QUANTITY_CHANGE, DELIVERY_DATE_CHANGE, PRODUCT_CHANGE, CANCEL',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/pm/orders/{id}/histories', 'request': 'id (Path)', 'response': 'historyId, changeType, beforeValue, afterValue, changedBy, changedAt', 'desc': '수주 변경이력 조회'},
            {'method': 'PUT', 'url': '/api/v1/pm/orders/{id}', 'request': 'deliveryDate, orderItems[{productId, qty}], changeReason*', 'response': 'orderId, updatedAt', 'desc': '수주 수정 (이력 자동 기록)'},
        ],
        logic=(
            '- 수주 수정 시 @EventListener로 OrderChangedEvent 발행 → OrderHistoryService 비동기 처리\n'
            '- 변경 전/후 값은 JSON 형태로 직렬화하여 before_value, after_value 컬럼에 저장\n'
            '- 생산 진행 중(PRODUCING) 상태의 수주 수정은 관리자 승인 필수\n'
            '- 이력 조회는 최대 1년치 데이터 제공'
        ),
        exception=(
            '- OrderChangeNotAllowedException: COMPLETED/CANCELLED 수주 수정 시도 시 409 반환\n'
            '- ChangeReasonRequiredException: CONFIRMED 이후 상태 수정 시 변경 사유 미입력 시 400 반환'
        )
    )

    # 3.3 납기캘린더조회
    add_heading(doc, '3.3 납기캘린더조회', 2)
    add_program_section(
        doc,
        prog_id='PM-003',
        prog_name='납기캘린더조회',
        desc='월별/주별 납기 일정을 캘린더 뷰로 조회한다. 납기 건수, 납품 수량, 생산 진행 상황을 시각화한다.',
        flow=[
            'GET /api/v1/pm/delivery-calendar?year={yyyy}&month={mm} 요청',
            '해당 월의 수주 납기 데이터를 날짜별로 그룹핑',
            'React FullCalendar 컴포넌트에 이벤트 데이터 바인딩',
            '날짜 클릭 시 해당일 수주 상세 목록 조회',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/pm/delivery-calendar', 'request': 'year*(4자리), month*(1~12)', 'response': 'calendarItems[{date, orderCount, totalQty, status, orders[{orderId, orderNo, partnerName, qty}]}]', 'desc': '납기 캘린더 조회'},
            {'method': 'GET', 'url': '/api/v1/pm/delivery-calendar/summary', 'request': 'startDate, endDate', 'response': '기간별 납기 요약 통계', 'desc': '납기 기간 요약 조회'},
        ],
        logic=(
            '- 캘린더 데이터는 Redis 캐시 적용 (키: "deliveryCalendar::{yyyy}-{mm}", TTL 10분)\n'
            '- 수주 상태별 색상 구분: CONFIRMED(파란색), PRODUCING(노란색), SHIPPED(녹색), DELAYED(빨간색)\n'
            '- 납기 임박(D-3 이내) 건은 대시보드 알림 및 이메일 발송'
        ),
        exception=(
            '- InvalidYearMonthException: 유효하지 않은 연월 입력 시 400 반환\n'
            '- FutureMonthLimitException: 현재 월 기준 6개월 초과 조회 시 400 반환'
        )
    )

    # 3.4 생산계획수립
    add_heading(doc, '3.4 생산계획수립', 2)
    add_program_section(
        doc,
        prog_id='PM-004',
        prog_name='생산계획수립',
        desc='확정된 수주를 기반으로 일별/주별 생산계획을 수립한다. 라인별 생산 능력을 고려하여 작업 할당을 최적화한다.',
        flow=[
            '확정 수주 목록 조회 → 미계획 수주 선택',
            '생산라인 선택 → 라인 가용 능력 자동 조회',
            'POST /api/v1/pm/plans → 생산계획 생성',
            '레시피 기준 원자재 소요량 자동 계산 → 재고 충족 여부 확인',
            '계획 확정 → 작업지시서 자동 생성 (PM-006)',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/pm/plans', 'request': 'planDate, lineNo, status', 'response': 'planId, planNo, orderId, productName, planQty, startTime, endTime, lineNo, status', 'desc': '생산계획 목록 조회'},
            {'method': 'POST', 'url': '/api/v1/pm/plans', 'request': 'orderId*, lineNo*, planDate*, planQty*, startTime*, endTime*, recipeId*', 'response': 'planId, planNo, materialRequirements[{materialId, requiredQty, availableStock}]', 'desc': '생산계획 등록'},
            {'method': 'PUT', 'url': '/api/v1/pm/plans/{id}', 'request': 'planDate, planQty, startTime, endTime, lineNo', 'response': 'planId, updatedAt', 'desc': '생산계획 수정'},
            {'method': 'PATCH', 'url': '/api/v1/pm/plans/{id}/confirm', 'request': '-', 'response': 'planId, status: CONFIRMED, workOrderId', 'desc': '생산계획 확정 (작업지시서 자동생성)'},
            {'method': 'GET', 'url': '/api/v1/pm/plans/{id}/material-requirements', 'request': 'id (Path)', 'response': '자재 소요량 목록 (materialId, name, requiredQty, availableStock, shortage)', 'desc': '자재 소요량 조회'},
        ],
        logic=(
            '- 생산능력(Capacity) 계산: 라인별 시간당 생산량 × 계획 시간\n'
            '- 레시피 기반 소요량: recipeQty × planQty / standardBatchSize\n'
            '- 재고 부족 경고: 필요수량 > 가용재고 시 shortage 필드에 부족량 표시\n'
            '- 계획 겹침 방지: 동일 라인 동일 시간대 중복 계획 생성 불가 검증\n'
            '- 생산계획 확정 시 자동으로 작업지시서 생성 (PM-006 연동)'
        ),
        exception=(
            '- LineCapacityExceededException: 라인 생산능력 초과 계획 시 400 반환\n'
            '- PlanTimeConflictException: 동일 라인 시간 충돌 시 409 반환\n'
            '- MaterialShortageException: 원자재 재고 부족 시 경고(200 + warning 플래그) 반환\n'
            '- OrderNotConfirmedException: 미확정 수주에 대한 계획 수립 시 400 반환'
        )
    )

    # 3.5 작업지시서관리
    add_heading(doc, '3.5 작업지시서관리', 2)
    add_program_section(
        doc,
        prog_id='PM-005',
        prog_name='작업지시서관리',
        desc='생산계획 확정 시 자동 생성되는 작업지시서를 조회·관리한다. 현장 POP 화면과 연동하여 실시간 생산 진행 현황을 추적한다.',
        flow=[
            '생산계획 확정 → 작업지시서 자동 생성 (planId → workOrderId)',
            'GET /api/v1/pm/work-orders - 작업지시서 목록 조회',
            'POP 화면에서 작업지시서 수신 → 생산 시작/종료 실적 입력',
            '작업지시서 완료: PATCH /api/v1/pm/work-orders/{id}/complete → 실적 집계',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/pm/work-orders', 'request': 'planDate, lineNo, status, page, size', 'response': 'workOrderId, workOrderNo, planId, productName, planQty, actualQty, status, startTime, endTime', 'desc': '작업지시서 목록 조회'},
            {'method': 'GET', 'url': '/api/v1/pm/work-orders/{id}', 'request': 'id (Path)', 'response': '작업지시서 상세 + 공정별 진행현황 + CCP 측정값', 'desc': '작업지시서 상세 조회'},
            {'method': 'PATCH', 'url': '/api/v1/pm/work-orders/{id}/start', 'request': 'workerId, startTime', 'response': 'workOrderId, status: IN_PROGRESS', 'desc': '작업 시작'},
            {'method': 'PATCH', 'url': '/api/v1/pm/work-orders/{id}/complete', 'request': 'actualQty*, defectQty, completedAt', 'response': 'workOrderId, status: COMPLETED, yieldRate', 'desc': '작업 완료'},
            {'method': 'GET', 'url': '/api/v1/pm/work-orders/{id}/progress', 'request': 'id (Path)', 'response': '공정별 진행률, 현재 공정 상태, CCP 실측값', 'desc': '작업 진행현황 조회'},
        ],
        logic=(
            '- 작업지시서 번호: WO-{YYYYMMDD}-{일련번호4자리}\n'
            '- 수율 계산: yieldRate = (actualQty - defectQty) / planQty × 100\n'
            '- 완료 시 KPI 집계 이벤트 발행 → KPI 모듈에서 비동기 처리\n'
            '- POP 화면과 WebSocket 연결로 실시간 진행률 Push 알림'
        ),
        exception=(
            '- WorkOrderAlreadyStartedException: 이미 시작된 작업지시서 재시작 시 409 반환\n'
            '- ActualQtyExceededException: 실적수량 > 계획수량 × 1.1 초과 시 경고 반환\n'
            '- WorkerNotAssignedException: 담당자 없이 작업 시작 시 400 반환'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# 4장. 자재재고관리 (INV)
# ──────────────────────────────────────────────

def add_chapter4_inv(doc: Document):
    add_heading(doc, '4. 자재재고관리 (INV) 프로그램 설계', 1)
    add_body(doc,
        '자재재고관리(INV, Inventory Management) 모듈은 원·부자재의 입고부터 소비까지 '
        '전체 재고 이동을 실시간으로 추적하고 관리한다.')
    doc.add_paragraph()

    add_heading(doc, '4.1 원부자재입고등록', 2)
    add_program_section(
        doc,
        prog_id='INV-001',
        prog_name='원부자재입고등록',
        desc='원자재(배추, 고추, 마늘 등) 및 부자재(포장재, 용기 등) 입고 정보를 등록한다. 입고검사 결과와 연동하여 검사 합격 후 재고에 반영한다.',
        flow=[
            '발주서(또는 거래처 납품서) 기반 입고 등록',
            'POST /api/v1/inv/receipts → 입고번호 자동생성',
            '입고 수량, 단위, LOT번호, 유통기한 입력',
            '입고검사 요청 이벤트 발행 → RCV 모듈 입고검사관리(RCV-001) 연동',
            '검사 합격 시: 재고 자동 증가 → 재고 이동 이력 기록',
            '검사 불합격 시: 반품 처리 → 재고 미반영',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/inv/receipts', 'request': 'startDate, endDate, materialId, status, page, size', 'response': 'receiptId, receiptNo, materialName, qty, unit, lotNo, receiptDate, inspectionStatus', 'desc': '입고 목록 조회'},
            {'method': 'POST', 'url': '/api/v1/inv/receipts', 'request': 'partnerId*, materialId*, qty*, unit*, lotNo*, receiptDate*, expiryDate, vehicleNo, notes', 'response': 'receiptId, receiptNo, inspectionRequestId', 'desc': '입고 등록'},
            {'method': 'GET', 'url': '/api/v1/inv/receipts/{id}', 'request': 'id (Path)', 'response': '입고 상세 + 검사 결과 + 재고 반영 이력', 'desc': '입고 단건 조회'},
            {'method': 'PATCH', 'url': '/api/v1/inv/receipts/{id}/approve', 'request': 'inspectionResult (PASS/FAIL)', 'response': 'receiptId, inventoryUpdated: true/false', 'desc': '입고 검사 결과 반영'},
        ],
        logic=(
            '- 입고번호: RCP-{YYYYMMDD}-{일련번호4자리}\n'
            '- LOT 추적: 입고 LOT → 생산 LOT → 출하 LOT 전체 이력 추적 (Lot Traceability)\n'
            '- 유통기한 관리: FIFO(선입선출) 원칙 적용, 유통기한 임박(7일 이내) 알림\n'
            '- 재고 반영 트랜잭션: 입고검사 합격 이벤트 수신 후 @Transactional 보장'
        ),
        exception=(
            '- DuplicateLotNoException: 동일 자재의 LOT번호 중복 시 409 반환\n'
            '- ExpiryDatePastException: 유통기한이 오늘 이전인 자재 입고 시 400 반환\n'
            '- ReceiptInspectionPendingException: 검사 대기 중인 입고 건 재처리 시도 시 409 반환'
        )
    )

    add_heading(doc, '4.2 입고이력조회', 2)
    add_program_section(
        doc,
        prog_id='INV-002',
        prog_name='입고이력조회',
        desc='기간별, 자재별, 거래처별 입고 이력을 다양한 조건으로 조회하고 엑셀 다운로드를 제공한다.',
        flow=[
            'GET /api/v1/inv/receipts/history - 다중 조건 검색',
            '결과 데이터 엑셀 내보내기: GET /api/v1/inv/receipts/history/export',
            'POI 라이브러리를 사용하여 서버에서 Excel 파일 생성 후 스트림 전송',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/inv/receipts/history', 'request': 'startDate*, endDate*, materialCode, partnerId, inspectionStatus, page, size', 'response': '입고 이력 목록 (페이징)', 'desc': '입고 이력 조회'},
            {'method': 'GET', 'url': '/api/v1/inv/receipts/history/export', 'request': '조회 조건 동일', 'response': 'application/vnd.ms-excel (파일 스트림)', 'desc': '입고 이력 엑셀 내보내기'},
        ],
        logic=(
            '- 최대 조회 기간: 1년 (365일)\n'
            '- 엑셀 내보내기 최대 건수: 10,000건 (초과 시 분할 다운로드 안내)\n'
            '- Apache POI 사용: XSSFWorkbook, 헤더 스타일/데이터 스타일 분리'
        ),
        exception=(
            '- SearchPeriodExceededException: 1년 초과 기간 조회 시 400 반환\n'
            '- ExportLimitExceededException: 10,000건 초과 내보내기 시도 시 400 반환'
        )
    )

    add_heading(doc, '4.3 실시간재고관리', 2)
    add_program_section(
        doc,
        prog_id='INV-003',
        prog_name='실시간재고관리',
        desc='원·부자재의 현재 재고 수량을 실시간으로 조회하고, 재고 이동(입고/출고/조정) 이력을 관리한다. 안전재고 이하 시 자동 알림을 제공한다.',
        flow=[
            'GET /api/v1/inv/stocks - 현재 재고 현황 조회',
            'Redis 캐시에서 최신 재고 데이터 우선 조회 (캐시 미스 시 DB 조회)',
            '재고 조정: POST /api/v1/inv/stocks/adjustments (실사 기반 조정)',
            '안전재고 기준: 각 자재별 safetyStock 설정값과 비교 → 임계 시 알림',
            'Server-Sent Events(SSE)로 재고 변동 시 프론트 실시간 갱신',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/inv/stocks', 'request': 'materialCode(선택), warehouseCode(선택), belowSafety(boolean, 선택)', 'response': 'stockId, materialCode, materialName, currentQty, unit, safetyStock, warehouseCode, lastUpdated', 'desc': '현재 재고 현황 조회'},
            {'method': 'GET', 'url': '/api/v1/inv/stocks/{materialId}/movements', 'request': 'materialId (Path), startDate, endDate', 'response': '재고 이동 이력 (입고/출고/조정/반품)', 'desc': '재고 이동 이력 조회'},
            {'method': 'POST', 'url': '/api/v1/inv/stocks/adjustments', 'request': 'materialId*, adjustQty*, reason*, adjustType (INCREASE/DECREASE/SET)', 'response': 'adjustmentId, beforeQty, afterQty', 'desc': '재고 조정'},
            {'method': 'GET', 'url': '/api/v1/inv/stocks/sse', 'request': '-', 'response': 'text/event-stream (SSE 스트림)', 'desc': '재고 변동 실시간 스트림'},
        ],
        logic=(
            '- 재고 데이터 Redis 캐싱: "stock::{materialId}" 키, 실시간 갱신 (TTL 없음, 이벤트 기반 갱신)\n'
            '- 재고 이동 유형: RECEIPT(입고), CONSUMPTION(소비), ADJUSTMENT(조정), RETURN(반품), DISPOSAL(폐기)\n'
            '- 안전재고 알림: Spring @Scheduled(fixedRate=300000) - 5분마다 체크 → Redis Pub/Sub으로 알림\n'
            '- 재고 음수 방지: 출고 시 currentQty < 0 불가 제약 → DB 레벨 CHECK 제약조건 적용'
        ),
        exception=(
            '- StockNegativeException: 재고 음수 발생 가능 출고 시 400 반환\n'
            '- InvalidAdjustmentException: 조정 후 수량이 음수가 되는 경우 400 반환\n'
            '- StockNotFoundException: 미등록 자재의 재고 조회 시 404 반환'
        )
    )

    add_heading(doc, '4.4 레시피기반소요량계산', 2)
    add_program_section(
        doc,
        prog_id='INV-004',
        prog_name='레시피기반소요량계산',
        desc='생산계획 수량과 제품 레시피를 기반으로 필요한 원·부자재 소요량을 자동 계산한다. 현재 재고와 비교하여 부족 자재 및 발주 필요량을 산출한다.',
        flow=[
            '생산계획 선택 또는 수동 수량 입력',
            'GET /api/v1/inv/material-requirements?planId={id} 요청',
            '레시피 BOM 조회 → 배합비율 기반 소요량 계산',
            '현재 재고 조회 → 소요량 대비 재고 충족 여부 분석',
            '부족 자재 목록 및 발주 권고량 반환',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/inv/material-requirements', 'request': 'planId(planId 또는 productId+qty 중 택일), productId, qty', 'response': 'requirements[{materialId, materialName, requiredQty, availableStock, shortage, recommendOrderQty}]', 'desc': '소요량 계산 조회'},
            {'method': 'POST', 'url': '/api/v1/inv/material-requirements/reserve', 'request': 'planId*, requirements[{materialId, reserveQty}]', 'response': 'reservationId, reservedItems', 'desc': '자재 예약 (생산 전 재고 확보)'},
        ],
        logic=(
            '- 소요량 계산식: 각 자재 소요량 = (레시피 배합비율 × 계획수량) / 표준 배치 크기\n'
            '- 손실률(yieldLossRate) 반영: 실 소요량 = 이론 소요량 / (1 - 손실률)\n'
            '- 발주 권고량: 부족량 + 안전재고 여유분 (최소발주단위 반올림)\n'
            '- 자재 예약: reserved_stock 컬럼 업데이트로 가용재고에서 차감'
        ),
        exception=(
            '- RecipeNotFoundException: 해당 제품의 활성 레시피 없을 시 404 반환\n'
            '- PlanNotFoundException: 존재하지 않는 planId 참조 시 404 반환'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# 5장. 입고전처리관리 (RCV)
# ──────────────────────────────────────────────

def add_chapter5_rcv(doc: Document):
    add_heading(doc, '5. 입고전처리관리 (RCV) 프로그램 설계', 1)
    add_body(doc,
        '입고전처리관리(RCV, Receiving Pre-processing) 모듈은 원물(배추, 무 등)이 입고된 후 '
        '전처리(선별, 절단, 손질) 전까지의 품질 검사 및 중량 관리를 담당한다.')
    doc.add_paragraph()

    add_heading(doc, '5.1 원물입고검사관리', 2)
    add_program_section(
        doc,
        prog_id='RCV-001',
        prog_name='원물입고검사관리',
        desc='입고된 원물의 품질 검사(관능검사, 이물검사, 선도검사)를 수행하고 합격/불합격 판정을 기록한다.',
        flow=[
            '입고 등록(INV-001) 완료 → 검사 요청 이벤트 수신',
            'GET /api/v1/rcv/inspections?receiptId={id} - 검사 대상 조회',
            '검사자 지정 → 항목별 검사 수행 (관능, 이물, 선도, 중량)',
            'POST /api/v1/rcv/inspections/{id}/results - 검사 결과 입력',
            '합격 판정: PASS → INV 재고 반영 이벤트 발행',
            '불합격 판정: FAIL → 반품 처리 이벤트 발행 → 거래처 통보',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/rcv/inspections', 'request': 'inspectionDate, status (PENDING/PASS/FAIL), materialCode', 'response': 'inspectionId, receiptId, materialName, qty, unit, inspectionStatus, inspectorName', 'desc': '검사 목록 조회'},
            {'method': 'POST', 'url': '/api/v1/rcv/inspections/{id}/results', 'request': 'inspectorId*, sensoryScore, foreignMatterResult (PASS/FAIL), freshnessScore, actualWeight*, overallResult (PASS/FAIL), notes', 'response': 'inspectionId, overallResult, inventoryUpdated', 'desc': '검사 결과 등록'},
            {'method': 'GET', 'url': '/api/v1/rcv/inspections/{id}', 'request': 'id (Path)', 'response': '검사 상세 + 각 항목별 결과 + 검사자 정보', 'desc': '검사 단건 조회'},
        ],
        logic=(
            '- 합격 기준: 관능점수 ≥ 3점(5점 만점), 이물검사 PASS, 선도점수 ≥ 3점\n'
            '- 하나라도 불합격 기준 미달 시 전체 불합격 처리\n'
            '- 검사 결과 이미지 첨부 지원: AWS S3 업로드 후 URL 저장\n'
            '- 불합격 시 자동 SMS 발송: 거래처 담당자 연락처로 통보'
        ),
        exception=(
            '- InspectionAlreadyCompletedException: 이미 완료된 검사 결과 재입력 시 409 반환\n'
            '- InspectorNotFoundException: 미등록 검사자 ID 사용 시 404 반환\n'
            '- WeightDiscrepancyException: 실측중량이 입고중량 대비 ±10% 초과 시 경고 반환'
        )
    )

    add_heading(doc, '5.2 전처리중량관리', 2)
    add_program_section(
        doc,
        prog_id='RCV-002',
        prog_name='전처리중량관리',
        desc='원물 전처리(손질, 절단, 탈피) 전후의 중량을 측정하여 전처리 수율을 관리한다.',
        flow=[
            '전처리 시작: 원물 투입 중량 측정 및 입력',
            'POST /api/v1/rcv/pre-processing → 전처리 작업 등록',
            '전처리 완료: 완료 중량 입력 → 수율 자동 계산',
            '수율 이상(기준 대비 ±5% 초과) 시 알림 발생',
        ],
        apis=[
            {'method': 'POST', 'url': '/api/v1/rcv/pre-processing', 'request': 'receiptId*, materialId*, inputWeight*, processType (TRIMMING/CUTTING/PEELING), workerId*', 'response': 'preProcessId, startedAt', 'desc': '전처리 작업 시작'},
            {'method': 'PATCH', 'url': '/api/v1/rcv/pre-processing/{id}/complete', 'request': 'outputWeight*, notes', 'response': 'preProcessId, inputWeight, outputWeight, yieldRate, yieldStatus (NORMAL/WARNING)', 'desc': '전처리 완료 (수율 계산)'},
            {'method': 'GET', 'url': '/api/v1/rcv/pre-processing', 'request': 'processDate, materialId, yieldStatus', 'response': '전처리 목록 (수율 포함)', 'desc': '전처리 이력 조회'},
        ],
        logic=(
            '- 수율 계산: yieldRate = outputWeight / inputWeight × 100 (%)\n'
            '- 기준 수율: 자재별 표준 수율 BM 등록 데이터 참조\n'
            '- 수율 이상 기준: 기준수율 ± 5% 초과 시 WARN 상태 기록 + 알림\n'
            '- 전처리 완료 중량은 재고관리 모듈의 가공재고에 자동 반영'
        ),
        exception=(
            '- OutputWeightExceedInputException: 출력중량 > 입력중량 시 400 반환\n'
            '- PreProcessAlreadyCompletedException: 완료된 작업 재완료 시도 시 409 반환'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# 6장. 세척공정관리 (WSH)
# ──────────────────────────────────────────────

def add_chapter6_wsh(doc: Document):
    add_heading(doc, '6. 세척공정관리 (WSH) 프로그램 설계', 1)
    add_body(doc,
        '세척공정관리(WSH, Washing Process Management) 모듈은 배추, 무 등 원물의 세척 및 소독 공정을 관리한다. '
        'HACCP CCP 기준에 따른 소독수 농도와 세척 시간을 모니터링한다.')
    doc.add_paragraph()

    add_heading(doc, '6.1 세척조건관리', 2)
    add_program_section(
        doc,
        prog_id='WSH-001',
        prog_name='세척조건관리',
        desc='제품/원물 유형별 세척 조건(세척수 온도, 세척 시간, 압력)을 설정하고 관리한다.',
        flow=[
            'GET /api/v1/wsh/conditions - 세척 조건 목록 조회',
            '세척 조건 등록: 원물유형 선택 → 세척시간, 수온, 압력 기준값 입력',
            'POST /api/v1/wsh/conditions → CCP 기준(BM-003)과 연동하여 저장',
            '공정 시작 시 해당 조건 자동 로딩 → MQTT로 설비에 전송',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/wsh/conditions', 'request': 'materialType(선택)', 'response': 'conditionId, materialType, washingTime, waterTemp, pressure, sanitizerConc, ccpRefId', 'desc': '세척 조건 목록 조회'},
            {'method': 'POST', 'url': '/api/v1/wsh/conditions', 'request': 'materialType*, washingTime*, waterTemp*, pressure, sanitizerConc*, ccpRefId', 'response': 'conditionId, createdAt', 'desc': '세척 조건 등록'},
            {'method': 'PUT', 'url': '/api/v1/wsh/conditions/{id}', 'request': '세척 조건 수정 필드', 'response': 'conditionId, updatedAt', 'desc': '세척 조건 수정'},
        ],
        logic=(
            '- 세척 조건은 원물 유형별 1개 ACTIVE 조건만 유지 (중복 불가)\n'
            '- MQTT publish: mes/washing/condition/{lineNo} 토픽으로 설비 조건 전송\n'
            '- CCP 연동: sanitizerConc는 BM-003의 CONCENTRATION 유형 CCP와 연결'
        ),
        exception=(
            '- DuplicateWashingConditionException: 동일 원물유형 조건 중복 등록 시 409 반환\n'
            '- CcpRefNotFoundException: 참조 CCP 기준 미존재 시 404 반환'
        )
    )

    add_heading(doc, '6.2 소독수농도관리', 2)
    add_program_section(
        doc,
        prog_id='WSH-002',
        prog_name='소독수농도관리',
        desc='세척 공정 중 소독수 농도를 실시간 측정하고, CCP 기준 이탈 시 즉시 알림을 발생시킨다.',
        flow=[
            'MQTT 설비 센서에서 소독수 농도 데이터 수신',
            'MqttMessageHandler → SanitizerConcentrationService 처리',
            'CCP 기준값과 비교 → 이탈 여부 판단',
            'POST /api/v1/wsh/sanitizer-records → 측정 기록 저장',
            '이탈 시: 품질이상관리(QC) 이슈 자동 등록 + 담당자 알림',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/wsh/sanitizer-records', 'request': 'lineNo, startDatetime, endDatetime', 'response': 'recordId, lineNo, measuredValue, ccpMinValue, ccpMaxValue, status (NORMAL/WARNING/CRITICAL), measuredAt', 'desc': '소독수 농도 측정 이력 조회'},
            {'method': 'POST', 'url': '/api/v1/wsh/sanitizer-records', 'request': 'lineNo*, measuredValue*, measuredAt* (MQTT 자동 호출)', 'response': 'recordId, status', 'desc': '소독수 농도 기록 저장'},
            {'method': 'GET', 'url': '/api/v1/wsh/sanitizer-records/realtime', 'request': 'lineNo', 'response': 'text/event-stream (현재 농도 SSE)', 'desc': '소독수 농도 실시간 조회 (SSE)'},
        ],
        logic=(
            '- MQTT 수신 주기: 30초마다 측정값 수신\n'
            '- 상태 판정: NORMAL(기준 내), WARNING(±10% 이탈), CRITICAL(±20% 이탈)\n'
            '- CRITICAL 상태 연속 3회 발생 시 라인 자동 정지 신호 MQTT 발행\n'
            '- 모든 측정 데이터 TimeSeries DB(MySQL Partition by month) 저장'
        ),
        exception=(
            '- InvalidMeasuredValueException: 측정값 음수 또는 비현실적 값(>10,000ppm) 시 400 반환\n'
            '- LineNotFoundException: 미등록 라인 번호 데이터 수신 시 경고 로그 기록'
        )
    )

    add_heading(doc, '6.3 세척실적데이터수집', 2)
    add_program_section(
        doc,
        prog_id='WSH-003',
        prog_name='세척실적데이터수집',
        desc='세척 공정의 투입량, 완료량, 세척 시간 등 실적 데이터를 수집하고 집계한다.',
        flow=[
            '세척 공정 시작: POST /api/v1/wsh/results → 작업지시서 연동',
            'MQTT에서 설비 실적 데이터 자동 수신',
            '세척 완료: PATCH /api/v1/wsh/results/{id}/complete → 실적 집계',
            'KPI 모듈에 세척 공정 실적 데이터 이벤트 발행',
        ],
        apis=[
            {'method': 'POST', 'url': '/api/v1/wsh/results', 'request': 'workOrderId*, lineNo*, inputWeight*, workerId*', 'response': 'resultId, startedAt', 'desc': '세척 실적 시작'},
            {'method': 'PATCH', 'url': '/api/v1/wsh/results/{id}/complete', 'request': 'outputWeight*, defectWeight, washingDuration', 'response': 'resultId, yieldRate, completedAt', 'desc': '세척 실적 완료'},
            {'method': 'GET', 'url': '/api/v1/wsh/results', 'request': 'workOrderId, resultDate, lineNo', 'response': '세척 실적 목록', 'desc': '세척 실적 조회'},
        ],
        logic=(
            '- 세척 수율 계산: yieldRate = outputWeight / inputWeight × 100\n'
            '- MQTT 실적 수신: mes/washing/result/{lineNo} 토픽\n'
            '- 세척 완료 데이터 → 절임공정(PCK) 투입 준비 상태로 자동 전환'
        ),
        exception=(
            '- WorkOrderNotFoundException: 연결된 작업지시서 미존재 시 404 반환\n'
            '- DuplicateWashingResultException: 동일 작업지시서 세척 실적 중복 등록 시 409 반환'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# 7장. 절임공정관리 (PCK)
# ──────────────────────────────────────────────

def add_chapter7_pck(doc: Document):
    add_heading(doc, '7. 절임공정관리 (PCK) 프로그램 설계', 1)
    add_body(doc,
        '절임공정관리(PCK, Pickling Process Management) 모듈은 배추 등 원물의 염수 절임 공정을 관리한다. '
        '절임통별 염도, 온도, 절임 시간 등의 CCP를 모니터링하고 실적을 기록한다.')
    doc.add_paragraph()

    add_heading(doc, '7.1 절임조건설정', 2)
    add_program_section(
        doc,
        prog_id='PCK-001',
        prog_name='절임조건설정',
        desc='제품 유형별 절임 조건(염도, 절임 시간, 절임 온도, 염수 비율)을 설정한다.',
        flow=[
            'GET /api/v1/pck/conditions - 절임 조건 목록 조회',
            '제품 유형 선택 → 절임 조건(염도%, 시간, 온도) 입력',
            'POST /api/v1/pck/conditions → 저장',
            '절임통 운영 시작 시 해당 조건 자동 적용',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/pck/conditions', 'request': 'productType(선택)', 'response': 'conditionId, productType, saltConc, picklingTime, picklingTemp, brineRatio', 'desc': '절임 조건 목록 조회'},
            {'method': 'POST', 'url': '/api/v1/pck/conditions', 'request': 'productType*, saltConc*(%), picklingTime*(분), picklingTemp*(°C), brineRatio*(배추:염수 비율)', 'response': 'conditionId, createdAt', 'desc': '절임 조건 등록'},
            {'method': 'PUT', 'url': '/api/v1/pck/conditions/{id}', 'request': '절임 조건 수정 필드', 'response': 'conditionId, updatedAt', 'desc': '절임 조건 수정'},
        ],
        logic=(
            '- 염도 허용 범위: 1% ~ 15% (HACCP 기준)\n'
            '- 절임 시간: 최소 2시간 ~ 최대 24시간\n'
            '- 절임 온도: 0°C ~ 20°C 범위 내에서만 설정 가능\n'
            '- 조건 변경 이력 자동 기록 (BM 데이터 변경과 동일 메커니즘)'
        ),
        exception=(
            '- InvalidSaltConcException: 염도 범위 초과 시 400 반환\n'
            '- InvalidPicklingTimeException: 절임 시간 범위 초과 시 400 반환'
        )
    )

    add_heading(doc, '7.2 절임통운영관리', 2)
    add_program_section(
        doc,
        prog_id='PCK-002',
        prog_name='절임통운영관리',
        desc='절임통별 투입 배추 중량, 현재 염도, 절임 경과 시간, 완료 예정 시간을 실시간으로 관리한다.',
        flow=[
            '절임통 선택 → 절임 시작: POST /api/v1/pck/tank-operations',
            'MQTT에서 탱크 센서 데이터 수신 (염도, 온도, 중량)',
            '실시간 현황: GET /api/v1/pck/tank-operations/realtime (SSE)',
            '절임 완료 시간 도래 → 알림 발생 → PATCH /api/v1/pck/tank-operations/{id}/complete',
            '완료된 절임 배추 → 양념버무림(SZN) 공정 투입 준비',
        ],
        apis=[
            {'method': 'POST', 'url': '/api/v1/pck/tank-operations', 'request': 'tankId*, conditionId*, inputWeight*, workOrderId*, workerId*, startedAt', 'response': 'operationId, expectedCompletedAt', 'desc': '절임 작업 시작'},
            {'method': 'GET', 'url': '/api/v1/pck/tank-operations', 'request': 'status (IN_PROGRESS/COMPLETED), tankId', 'response': '절임통 운영 목록 (경과시간, 현재염도, 상태)', 'desc': '절임통 운영 현황 조회'},
            {'method': 'GET', 'url': '/api/v1/pck/tank-operations/realtime', 'request': '-', 'response': 'text/event-stream (탱크별 실시간 센서 데이터)', 'desc': '절임통 실시간 현황 (SSE)'},
            {'method': 'PATCH', 'url': '/api/v1/pck/tank-operations/{id}/complete', 'request': 'outputWeight*, finalSaltConc*, completedAt', 'response': 'operationId, status: COMPLETED, yieldRate', 'desc': '절임 완료 처리'},
            {'method': 'GET', 'url': '/api/v1/pck/tank-operations/{id}/sensor-history', 'request': 'id (Path)', 'response': '작업 기간 내 센서 측정 이력 (염도, 온도, 타임스탬프)', 'desc': '절임 중 센서 이력 조회'},
        ],
        logic=(
            '- MQTT 구독 토픽: mes/pickling/tank/{tankCode}/+ (염도/온도/중량)\n'
            '- 염도 이탈 감지: 기준값 ±2% 초과 시 WARNING, ±5% 초과 시 CRITICAL\n'
            '- CRITICAL 상태 지속 시 QC 모듈에 자동 이슈 등록\n'
            '- 절임 완료 예정 알림: 완료 30분 전 담당자에게 Push 알림\n'
            '- 동시 절임통 운영 현황: Redis Hash 구조로 실시간 상태 저장'
        ),
        exception=(
            '- TankAlreadyInUseException: 사용 중인 절임통 재사용 시도 시 409 반환\n'
            '- FinalSaltConcOutOfRangeException: 최종 염도가 조건 범위 이탈 시 경고 + 완료 처리\n'
            '- TankCapacityExceededException: 투입 중량이 탱크 최대 용량 초과 시 400 반환'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────

def main():
    doc = Document()

    # 페이지 여백 설정
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    add_cover_page(doc)
    add_chapter1_overview(doc)
    add_chapter2_bm(doc)
    add_chapter3_pm(doc)
    add_chapter4_inv(doc)
    add_chapter5_rcv(doc)
    add_chapter6_wsh(doc)
    add_chapter7_pck(doc)

    doc.save(OUTPUT_PATH)
    print(f"[완료] 프로그램설계서1 저장: {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
