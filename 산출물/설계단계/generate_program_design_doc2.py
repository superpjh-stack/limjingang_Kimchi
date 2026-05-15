"""
임진강김치 MES 시스템 구축 - 프로그램설계서2 생성 스크립트
대상 모듈: 양념버무림(SZN), 포장출하(PKG), 숙성냉장(AGE), 품질이상(QC),
          시스템관리(SYS), KPI모니터링(KPI), 설비관리(EQP), POP작업(POP), AI Agent(AI)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\gerardo\01 SmallSF\Imjingang-Kimchi\산출물\설계단계\07_프로그램설계서2.docx"


# ──────────────────────────────────────────────
# 헬퍼 함수 (doc1과 동일)
# ──────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int):
    return doc.add_heading(text, level=level)


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(10)
    return p


def add_api_table(doc: Document, apis: list):
    headers = ['Method', 'URL', 'Request (주요 파라미터)', 'Response (주요 필드)', '설명']
    tbl = doc.add_table(rows=1 + len(apis), cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr[i], '1F497D')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    method_colors = {'GET': 'E2EFDA', 'POST': 'FFF2CC', 'PUT': 'DDEBF7', 'DELETE': 'FCE4D6', 'PATCH': 'EAD1DC'}
    widths = [2.0, 5.5, 5.5, 5.5, 3.5]

    for r_idx, api in enumerate(apis, 1):
        row = tbl.rows[r_idx].cells
        row[0].text = api.get('method', '')
        row[1].text = api.get('url', '')
        row[2].text = api.get('request', '-')
        row[3].text = api.get('response', '-')
        row[4].text = api.get('desc', '')
        bg = method_colors.get(api.get('method', ''), 'FFFFFF')
        set_cell_bg(row[0], bg)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if r_idx % 2 == 0:
            for col in [1, 2, 3, 4]:
                set_cell_bg(row[col], 'F2F2F2')

    for col_idx, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[col_idx].width = Cm(w)

    doc.add_paragraph()


def add_program_section(doc: Document, prog_id: str, prog_name: str, desc: str, flow: list, apis: list, logic: str, exception: str):
    info = [
        ('프로그램 ID', prog_id),
        ('프로그램명', prog_name),
        ('기능 설명', desc),
    ]
    tbl = doc.add_table(rows=len(info), cols=2)
    tbl.style = 'Table Grid'
    for i, (k, v) in enumerate(info):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl.rows[i].cells[0], 'D9E1F2')
        tbl.rows[i].cells[1].text = v
        tbl.rows[i].cells[0].width = Cm(3.5)
        tbl.rows[i].cells[1].width = Cm(13.5)
    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run('■ 처리 흐름')
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    for step in flow:
        bp = doc.add_paragraph(step, style='List Number')
        bp.paragraph_format.left_indent = Cm(0.5)

    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run('■ API 설계')
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    add_api_table(doc, apis)

    p = doc.add_paragraph()
    r = p.add_run('■ 주요 로직 설명')
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    doc.add_paragraph(logic)
    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run('■ 예외 처리')
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    doc.add_paragraph(exception)
    doc.add_paragraph()

    hr = doc.add_paragraph('─' * 80)
    hr.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()


# ──────────────────────────────────────────────
# 표지
# ──────────────────────────────────────────────

def add_cover_page(doc: Document):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('프로그램설계서 2')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('임진강김치 MES 시스템 구축')
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x26, 0x26, 0x26)

    doc.add_paragraph()
    doc.add_paragraph()

    info_lines = [
        ('작성일', '2026-05-12'),
        ('버전', 'V1.0'),
        ('작성자', '설계자2'),
        ('대상 모듈', 'SZN / PKG / AGE / QC / SYS / KPI / EQP / POP / AI'),
    ]
    tbl = doc.add_table(rows=len(info_lines), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    for i, (k, v) in enumerate(info_lines):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl.rows[i].cells[0], 'D9E1F2')
        tbl.rows[i].cells[1].text = v
    for row in tbl.rows:
        for cell in row.cells:
            cell.width = Cm(7)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


# ──────────────────────────────────────────────
# 1장. 양념버무림공정관리 (SZN)
# ──────────────────────────────────────────────

def add_chapter1_szn(doc: Document):
    add_heading(doc, '1. 양념버무림공정관리 (SZN) 프로그램 설계', 1)
    add_body(doc,
        '양념버무림공정관리(SZN, Seasoning Process Management) 모듈은 절임 배추에 양념을 배합하고 버무리는 공정을 관리한다. '
        '양념 계량, 충진기 실적, CCP 모니터링을 포함한다.')
    doc.add_paragraph()

    add_heading(doc, '1.1 양념계량관리', 2)
    add_program_section(
        doc,
        prog_id='SZN-001',
        prog_name='양념계량관리',
        desc='레시피에 따른 각 양념 재료(고추가루, 마늘, 생강 등)의 계량 작업을 등록하고 계획 대비 실측값을 관리한다.',
        flow=[
            '작업지시서 선택 → 레시피 기반 양념 소요량 자동 계산',
            'GET /api/v1/szn/seasoning-plans?workOrderId={id} - 계획 소요량 조회',
            '계량 시작: 재료별 실측 중량 입력',
            'POST /api/v1/szn/seasoning-records → 계량 실적 저장',
            '계획 대비 편차 ±3% 초과 시 경고 표시 및 확인 요청',
            '전체 양념 계량 완료 → 버무림 공정 시작 가능 상태로 전환',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/szn/seasoning-plans', 'request': 'workOrderId(필수)', 'response': 'planItems[{materialId, materialName, plannedQty, unit}]', 'desc': '양념 계획 소요량 조회'},
            {'method': 'POST', 'url': '/api/v1/szn/seasoning-records', 'request': 'workOrderId*, seasoningItems[{materialId, actualQty, scaleId, workerId}]*', 'response': 'recordId, deviationItems[{materialId, deviation, status}]', 'desc': '양념 계량 실적 등록'},
            {'method': 'GET', 'url': '/api/v1/szn/seasoning-records', 'request': 'workOrderId, recordDate', 'response': '계량 실적 목록 (계획 대비 편차 포함)', 'desc': '양념 계량 이력 조회'},
        ],
        logic=(
            '- 편차 계산: deviation = (actualQty - plannedQty) / plannedQty × 100 (%)\n'
            '- 편차 기준: ±1% 이내 NORMAL, ±3% 이내 WARNING, ±3% 초과 ALERT\n'
            '- ALERT 발생 시 QC 모듈에 자동 이슈 등록\n'
            '- 계량 기기(저울) 연동: RS-232/USB 인터페이스로 자동 중량 입력 지원\n'
            '- 레시피 기반 소요량 = 절임배추 실입고량 × 배합비율 (배치 단위 계산)'
        ),
        exception=(
            '- WeightDeviationAlertException: 편차 ±5% 초과 시 관리자 승인 필요 응답 반환\n'
            '- ScaleCalibrationException: 저울 미교정(최종교정일 > 7일) 시 경고 반환\n'
            '- SeasoningRecordDuplicateException: 동일 작업지시서 계량 중복 등록 시 409 반환'
        )
    )

    add_heading(doc, '1.2 속넣기충진기데이터수집', 2)
    add_program_section(
        doc,
        prog_id='SZN-002',
        prog_name='속넣기충진기데이터수집',
        desc='충진기(속넣기 설비)에서 MQTT를 통해 실시간 생산 데이터(충진 횟수, 충진량, 가동시간)를 수집하고 작업지시서에 연결한다.',
        flow=[
            'MQTT 구독: mes/filling/{eqpCode}/+ 토픽으로 충진기 데이터 수신',
            'FillingDataHandler → 데이터 유효성 검사 → DB 저장',
            'GET /api/v1/szn/filling-results - 실적 조회',
            'SSE로 실시간 충진 속도, 누계 수량 프론트 전송',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/szn/filling-results', 'request': 'workOrderId, eqpCode, resultDate', 'response': 'resultId, eqpCode, totalFillingCount, totalFillingWeight, operatingTime, avgFillingSpeed', 'desc': '충진기 실적 조회'},
            {'method': 'GET', 'url': '/api/v1/szn/filling-results/realtime', 'request': 'eqpCode', 'response': 'text/event-stream (충진 횟수, 속도 실시간)', 'desc': '충진기 실시간 현황 (SSE)'},
            {'method': 'POST', 'url': '/api/v1/szn/filling-results/manual', 'request': 'workOrderId*, eqpCode*, fillingCount, fillingWeight, workerId', 'response': 'resultId', 'desc': '충진 실적 수동 입력 (MQTT 장애 시)'},
        ],
        logic=(
            '- MQTT 수신 데이터: {eqpCode, timestamp, fillingCount, fillingWeight, runStatus}\n'
            '- 충진 속도 계산: fillingCount / operatingTime (회/분)\n'
            '- 목표 속도 대비 80% 이하 시 설비 저성능 알림 발생\n'
            '- 데이터 무결성: MQTT 메시지 중복 수신 방지 (messageId 기반 idempotency)'
        ),
        exception=(
            '- FillingDataParseException: MQTT 메시지 파싱 오류 시 DLQ(Dead Letter Queue) 이동\n'
            '- EqpNotLinkedToWorkOrderException: 작업지시서 미연결 설비 데이터 수신 시 경고 로그'
        )
    )

    add_heading(doc, '1.3 버무림CCP관리', 2)
    add_program_section(
        doc,
        prog_id='SZN-003',
        prog_name='버무림CCP관리',
        desc='버무림 공정의 CCP(혼합 온도, 혼합 시간)를 실시간 모니터링하고 이탈 시 즉시 알림을 발생시킨다.',
        flow=[
            'MQTT에서 혼합기 온도/시간 데이터 수신',
            'CCP 기준값(BM-003)과 실시간 비교',
            'POST /api/v1/szn/ccp-records → CCP 측정값 저장',
            'CCP 이탈 시 QC 모듈 자동 이슈 등록 + 담당자 알림',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/szn/ccp-records', 'request': 'workOrderId, startDatetime, endDatetime', 'response': 'records[{recordId, ccpType, measuredValue, ccpStatus, measuredAt}]', 'desc': 'CCP 측정 이력 조회'},
            {'method': 'POST', 'url': '/api/v1/szn/ccp-records', 'request': 'workOrderId*, ccpId*, measuredValue*, measuredAt* (MQTT 자동 호출)', 'response': 'recordId, ccpStatus', 'desc': 'CCP 측정값 저장'},
            {'method': 'GET', 'url': '/api/v1/szn/ccp-records/realtime', 'request': 'workOrderId', 'response': '실시간 CCP 현황 SSE', 'desc': '버무림 CCP 실시간 모니터링'},
        ],
        logic=(
            '- 혼합 온도 CCP: 0°C ~ 10°C 유지 (김치 발효 억제)\n'
            '- 혼합 시간 CCP: 최소 5분 ~ 최대 20분\n'
            '- 이탈 횟수 카운팅: Redis Counter 활용, 3회 연속 이탈 시 라인 알람'
        ),
        exception=(
            '- CcpMeasurementMissingException: CCP 측정 누락(30초 이상 데이터 없음) 시 알림\n'
            '- InvalidCcpValueException: 측정값이 물리적으로 불가능한 경우 400 반환'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# 2장. 포장출하관리 (PKG)
# ──────────────────────────────────────────────

def add_chapter2_pkg(doc: Document):
    add_heading(doc, '2. 포장출하관리 (PKG) 프로그램 설계', 1)
    add_body(doc,
        '포장출하관리(PKG, Packaging & Shipping Management) 모듈은 완성된 김치의 포장, 중량 검사, 테이핑, '
        '출하 냉장 보관까지 전 과정을 관리한다.')
    doc.add_paragraph()

    add_heading(doc, '2.1 중량검사관리', 2)
    add_program_section(
        doc,
        prog_id='PKG-001',
        prog_name='중량검사관리',
        desc='포장된 제품의 중량이 기준값(표시중량 ±허용오차)을 만족하는지 검사하고 불량품을 선별한다.',
        flow=[
            '포장 완료 → 중량 선별기에서 중량 자동 측정',
            'MQTT로 중량 데이터 수신 → 중량 기준 비교',
            'POST /api/v1/pkg/weight-inspections → 측정 기록 저장',
            '기준 이탈(Under/Over Weight) 제품: 불량 처리 → QC 연동',
            '중량 검사 통계: 시간대별 불량률 조회',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/pkg/weight-inspections', 'request': 'workOrderId, startDatetime, endDatetime, result (PASS/FAIL)', 'response': '중량 검사 목록 (측정값, 결과, 불량유형)', 'desc': '중량 검사 이력 조회'},
            {'method': 'POST', 'url': '/api/v1/pkg/weight-inspections', 'request': 'workOrderId*, eqpCode*, measuredWeight*, measuredAt* (MQTT 자동 호출)', 'response': 'inspectionId, result (PASS/UNDER/OVER), deviation', 'desc': '중량 검사 기록'},
            {'method': 'GET', 'url': '/api/v1/pkg/weight-inspections/statistics', 'request': 'workOrderId', 'response': 'totalCount, passCount, failCount, passRate, avgWeight, stdDev', 'desc': '중량 검사 통계 조회'},
        ],
        logic=(
            '- 허용 오차: ±표시중량의 3% (식품표시법 기준)\n'
            '- UNDER WEIGHT: 측정중량 < (표시중량 × 0.97)\n'
            '- OVER WEIGHT: 측정중량 > (표시중량 × 1.05) — 내용물 과다\n'
            '- 불량률 ≥ 3% 시 라인 자동 알림 + 품질이상 이슈 등록\n'
            '- 실시간 X-bar/R 관리도 데이터 제공 (통계적 공정관리 SPC)'
        ),
        exception=(
            '- WeightStandardNotFoundException: 해당 제품 중량 기준 미등록 시 404 반환\n'
            '- WeightDataOutlierException: 측정값이 표시중량 ±30% 초과 시 센서 오류로 판단, 제외 처리'
        )
    )

    add_heading(doc, '2.2 자동테이핑기실적관리', 2)
    add_program_section(
        doc,
        prog_id='PKG-002',
        prog_name='자동테이핑기실적관리',
        desc='자동 테이핑기의 가동 실적(테이핑 횟수, 가동시간, 불량 횟수)을 수집하고 OEE(종합설비효율)를 계산한다.',
        flow=[
            'MQTT 구독: mes/taping/{eqpCode}/result 토픽',
            'TapingResultHandler → 데이터 파싱 및 저장',
            'GET /api/v1/pkg/taping-results - 실적 및 OEE 조회',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/pkg/taping-results', 'request': 'eqpCode, resultDate', 'response': 'resultId, eqpCode, tapingCount, defectCount, operatingTime, oee, availability, performance, quality', 'desc': '테이핑기 실적 및 OEE 조회'},
            {'method': 'GET', 'url': '/api/v1/pkg/taping-results/realtime', 'request': 'eqpCode', 'response': 'text/event-stream (실시간 테이핑 카운트, OEE)', 'desc': '테이핑기 실시간 현황'},
        ],
        logic=(
            '- OEE = 가용성(Availability) × 성능(Performance) × 품질(Quality)\n'
            '- 가용성: 실제가동시간 / 계획가동시간\n'
            '- 성능: 실제생산수 / (이론사이클타임 × 실제가동시간)\n'
            '- 품질: (생산수 - 불량수) / 생산수\n'
            '- OEE < 65% 시 설비 점검 권고 알림'
        ),
        exception=(
            '- OeeCalculationException: 계획가동시간 = 0 일 때 OEE 계산 불가 → null 반환\n'
            '- TapingDataGapException: 30분 이상 데이터 미수신 시 알림 발생'
        )
    )

    add_heading(doc, '2.3 출하냉장고관리', 2)
    add_program_section(
        doc,
        prog_id='PKG-003',
        prog_name='출하냉장고관리',
        desc='완성 제품의 출하 전 냉장 보관 현황을 관리한다. 냉장고별 재고, 온도, 입출고 이력을 추적한다.',
        flow=[
            '포장 완료 제품 → 냉장고 입고: POST /api/v1/pkg/cold-storage/in',
            '냉장고 온도 MQTT 실시간 모니터링',
            '출하 시: POST /api/v1/pkg/cold-storage/out → 수주 연동 출고 처리',
            '냉장고 재고 현황: GET /api/v1/pkg/cold-storage/stocks',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/pkg/cold-storage/stocks', 'request': 'coldStorageId(선택), productId(선택)', 'response': '냉장고별 재고 현황 (제품명, 수량, 입고일, 온도)', 'desc': '출하 냉장고 재고 조회'},
            {'method': 'POST', 'url': '/api/v1/pkg/cold-storage/in', 'request': 'coldStorageId*, productId*, qty*, workOrderId*, storedAt', 'response': 'storageId, coldStorageId, currentStock', 'desc': '냉장고 입고'},
            {'method': 'POST', 'url': '/api/v1/pkg/cold-storage/out', 'request': 'orderId*, storageItems[{storageId, qty}]*', 'response': 'shippingId, shippedItems, remainingStock', 'desc': '냉장고 출고 (출하)'},
            {'method': 'GET', 'url': '/api/v1/pkg/cold-storage/temperature', 'request': 'coldStorageId', 'response': 'text/event-stream (냉장고 온도 실시간)', 'desc': '냉장고 온도 실시간 모니터링'},
        ],
        logic=(
            '- 냉장 온도 CCP: 0°C ~ 5°C (김치 보관 기준)\n'
            '- FIFO 출고 원칙: 입고일 순 자동 정렬로 선입선출 보장\n'
            '- 온도 이탈(5°C 초과) 시 즉시 알림 + QC 이슈 등록\n'
            '- 출하 시 수주(PM)와 재고(INV) 자동 연동 처리'
        ),
        exception=(
            '- ColdStorageCapacityException: 냉장고 최대 용량 초과 입고 시 400 반환\n'
            '- StockShortageException: 출하 수량 > 재고 수량 시 400 반환\n'
            '- TemperatureAlarmException: 냉장 온도 10°C 초과 시 CRITICAL 알림 발생'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# 3장. 숙성냉장관리 (AGE)
# ──────────────────────────────────────────────

def add_chapter3_age(doc: Document):
    add_heading(doc, '3. 숙성냉장관리 (AGE) 프로그램 설계', 1)
    add_body(doc,
        '숙성냉장관리(AGE, Aging & Refrigeration Management) 모듈은 버무림 완료 후 숙성 과정의 '
        '재고 관리와 냉장고 온도·습도 모니터링을 담당한다.')
    doc.add_paragraph()

    add_heading(doc, '3.1 숙성재고관리', 2)
    add_program_section(
        doc,
        prog_id='AGE-001',
        prog_name='숙성재고관리',
        desc='숙성 중인 김치의 재고(냉장고별, 배치별)를 관리하고 숙성 완료 예정일을 추적한다.',
        flow=[
            '버무림 완료 → 숙성 냉장고 입고: POST /api/v1/age/stocks/in',
            '숙성 기간 자동 계산: 입고일 + 표준숙성일수 = 완료 예정일',
            '숙성 완료 D-1 알림 발송',
            '숙성 완료 → 출하 냉장고(PKG-003) 이관',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/age/stocks', 'request': 'coldStorageId, status (AGING/COMPLETED/SHIPPED)', 'response': 'stockId, batchNo, productName, qty, storedDate, expectedCompletedDate, status', 'desc': '숙성 재고 현황 조회'},
            {'method': 'POST', 'url': '/api/v1/age/stocks/in', 'request': 'coldStorageId*, workOrderId*, qty*, productId*, batchNo*, storedDate', 'response': 'stockId, batchNo, expectedCompletedDate', 'desc': '숙성 냉장고 입고'},
            {'method': 'PATCH', 'url': '/api/v1/age/stocks/{id}/complete', 'request': '-', 'response': 'stockId, status: COMPLETED', 'desc': '숙성 완료 처리'},
            {'method': 'GET', 'url': '/api/v1/age/stocks/upcoming', 'request': 'days (기본 3)', 'response': 'N일 이내 숙성 완료 예정 목록', 'desc': '숙성 완료 예정 조회'},
        ],
        logic=(
            '- 배치번호(batchNo): BAT-{YYYYMMDD}-{일련번호3자리}\n'
            '- 표준 숙성 기간: 제품유형별 기준 데이터(BM) 참조 (예: 포기김치 3일, 깍두기 1일)\n'
            '- 숙성 완료 알림: @Scheduled(cron="0 9 * * *") 매일 오전 9시 체크\n'
            '- LOT 추적: 입고 LOT → 숙성 배치 → 출하 LOT 전체 연결'
        ),
        exception=(
            '- BatchNoConflictException: 배치번호 중복 시 409 반환\n'
            '- AgingStockNotFoundException: 미존재 재고 완료 처리 시 404 반환'
        )
    )

    add_heading(doc, '3.2 냉장고온도습도모니터링', 2)
    add_program_section(
        doc,
        prog_id='AGE-002',
        prog_name='냉장고온도습도모니터링',
        desc='숙성 냉장고의 온도·습도를 실시간으로 모니터링하고 이탈 시 즉각적인 알림을 제공한다.',
        flow=[
            'MQTT 구독: mes/coldroom/{roomCode}/environment',
            '수신 데이터: 온도(temp), 습도(humidity), 타임스탬프',
            'CCP 기준과 비교 → 상태 판정 → DB 저장',
            'SSE로 대시보드 실시간 전송',
            '이탈 지속(10분 이상) 시 CRITICAL 알림 + 담당자 문자 발송',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/age/environment/records', 'request': 'roomCode*, startDatetime, endDatetime, interval (1m/5m/1h)', 'response': '시계열 온습도 데이터 (집계 포함)', 'desc': '온습도 측정 이력 조회'},
            {'method': 'GET', 'url': '/api/v1/age/environment/realtime', 'request': '-', 'response': 'text/event-stream (전체 냉장고 온습도)', 'desc': '온습도 실시간 모니터링 (SSE)'},
            {'method': 'GET', 'url': '/api/v1/age/environment/alarms', 'request': 'roomCode, startDate, endDate', 'response': '온습도 이탈 알람 이력', 'desc': '온습도 알람 이력 조회'},
        ],
        logic=(
            '- 온도 기준: -2°C ~ 5°C (숙성 최적 온도)\n'
            '- 습도 기준: 70% ~ 90%\n'
            '- 측정 주기: 60초 (MQTT)\n'
            '- 이탈 알람 등급: WARNING(기준 ±10%), CRITICAL(기준 ±20%)\n'
            '- 시계열 데이터는 파티셔닝된 MySQL 테이블에 저장 (월별 파티션)\n'
            '- 7일치 데이터: Redis ZSet으로 빠른 그래프 조회 지원'
        ),
        exception=(
            '- SensorOfflineException: 60초 이상 데이터 미수신 시 오프라인 알림\n'
            '- InvalidTemperatureReadingException: 물리적 불가능 값(< -50 또는 > 100°C) 시 센서 오류로 처리'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# 4장. 품질이상관리 (QC)
# ──────────────────────────────────────────────

def add_chapter4_qc(doc: Document):
    add_heading(doc, '4. 품질이상관리 (QC) 프로그램 설계', 1)
    add_body(doc,
        '품질이상관리(QC, Quality Control) 모듈은 생산 전 공정에서 발생하는 품질 이상을 등록·추적하고 '
        '시정조치를 관리한다. 금속 검출기 실적 및 공정별 불량 분석을 포함한다.')
    doc.add_paragraph()

    add_heading(doc, '4.1 공정별불량요소관리', 2)
    add_program_section(
        doc,
        prog_id='QC-001',
        prog_name='공정별불량요소관리',
        desc='각 공정(세척, 절임, 버무림, 포장)에서 발생하는 불량 유형과 원인을 등록하고, 불량 발생 추이를 분석한다.',
        flow=[
            '불량 발생 감지 (자동: CCP 이탈, MQTT / 수동: 작업자 입력)',
            'POST /api/v1/qc/defects → 불량 정보 등록',
            '불량 원인 분석: 파레토 차트 데이터 제공',
            '시정조치(CAPA) 등록: PATCH /api/v1/qc/defects/{id}/capa',
            '시정조치 완료 후 재발 방지 여부 추적',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/qc/defects', 'request': 'processCode, defectType, startDate, endDate, status, page, size', 'response': 'defectId, processCode, defectType, qty, workOrderId, detectedAt, status', 'desc': '불량 목록 조회'},
            {'method': 'POST', 'url': '/api/v1/qc/defects', 'request': 'workOrderId*, processCode*, defectType*, qty*, defectCause, detectedBy, detectedAt*', 'response': 'defectId, issueNo', 'desc': '불량 등록'},
            {'method': 'PATCH', 'url': '/api/v1/qc/defects/{id}/capa', 'request': 'capaDescription*, capaAssignee*, dueDate*', 'response': 'defectId, capaId, status: CAPA_IN_PROGRESS', 'desc': '시정조치(CAPA) 등록'},
            {'method': 'GET', 'url': '/api/v1/qc/defects/statistics', 'request': 'startDate, endDate, processCode', 'response': '공정별/유형별 불량 통계 (파레토 데이터)', 'desc': '불량 통계 분석'},
        ],
        logic=(
            '- 불량 유형: WEIGHT(중량불량), FOREIGN_MATTER(이물질), APPEARANCE(외관불량), EXPIRY(유통기한), CCP_DEVIATION(CCP이탈)\n'
            '- 이슈번호 자동생성: QC-{YYYYMMDD}-{일련번호4자리}\n'
            '- 동일 불량 유형이 1시간 내 3건 이상 발생 시 라인 전체 알림\n'
            '- CAPA 마감일 초과 시 담당자 + 관리자에게 에스컬레이션 알림'
        ),
        exception=(
            '- DefectWorkOrderNotFoundException: 연결 작업지시서 미존재 시 404 반환\n'
            '- CapaAlreadyAssignedException: 이미 CAPA 지정된 불량에 재지정 시 409 반환'
        )
    )

    add_heading(doc, '4.2 금속검출기관리', 2)
    add_program_section(
        doc,
        prog_id='QC-002',
        prog_name='금속검출기관리',
        desc='포장 라인의 금속 검출기 검사 실적을 수집하고, 금속 검출 시 즉각적인 격리 조치 프로세스를 관리한다.',
        flow=[
            'MQTT: mes/metaldetector/{eqpCode}/detection 구독',
            '금속 미검출(PASS): 실적 카운트만 누적',
            '금속 검출(DETECT): 라인 정지 신호 발행 + 즉시 격리 처리',
            'POST /api/v1/qc/metal-detections → 검출 기록 저장',
            'QC 불량(QC-001)과 자동 연동 → 이슈 생성',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/qc/metal-detections', 'request': 'eqpCode, startDatetime, endDatetime, result (PASS/DETECT)', 'response': '검출 이력 목록 (측정시각, 결과, 조치내용)', 'desc': '금속 검출 이력 조회'},
            {'method': 'POST', 'url': '/api/v1/qc/metal-detections', 'request': 'eqpCode*, result*, detectedAt*, workOrderId', 'response': 'detectionId, result, isolationRequired', 'desc': '금속 검출 기록'},
            {'method': 'PATCH', 'url': '/api/v1/qc/metal-detections/{id}/isolate', 'request': 'isolationNote, isolatedBy', 'response': 'detectionId, isolatedAt', 'desc': '격리 조치 완료 처리'},
        ],
        logic=(
            '- 금속 검출 즉시: Redis Pub/Sub으로 전체 시스템 알림 발행\n'
            '- 격리 조치 필수: DETECT 기록 생성 후 15분 내 isolate 처리 없으면 에스컬레이션\n'
            '- 금속 검출기 교정 주기: 4시간마다 표준 금속편(Fe, Non-Fe, SUS) 감도 확인\n'
            '- 교정 미수행 시 해당 라인 생산 차단 로직 적용'
        ),
        exception=(
            '- MetalDetectionIsolationTimeoutException: 15분 내 격리 미처리 시 자동 에스컬레이션\n'
            '- CalibrationOverdueException: 교정 시간 초과 시 생산 중단 경고'
        )
    )

    add_heading(doc, '4.3 이슈이력관리', 2)
    add_program_section(
        doc,
        prog_id='QC-003',
        prog_name='이슈이력관리',
        desc='전 공정에서 발생한 품질 이슈의 전체 이력(등록→조치→완료→재발확인)을 통합 조회하고 관리한다.',
        flow=[
            'GET /api/v1/qc/issues - 전체 이슈 통합 조회',
            '이슈 상세: GET /api/v1/qc/issues/{id} → 발생원인, 조치내역, 완료일 포함',
            '이슈 완료: PATCH /api/v1/qc/issues/{id}/close',
            '월별 이슈 리포트 자동 생성 (AI Agent 연동)',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/qc/issues', 'request': 'issueType, processCode, status (OPEN/IN_PROGRESS/CLOSED), startDate, endDate', 'response': '이슈 목록 (이슈번호, 유형, 공정, 상태, 발생일, 담당자)', 'desc': '이슈 통합 목록 조회'},
            {'method': 'GET', 'url': '/api/v1/qc/issues/{id}', 'request': 'id (Path)', 'response': '이슈 상세 + CAPA 이력 + 관련 공정 데이터', 'desc': '이슈 상세 조회'},
            {'method': 'PATCH', 'url': '/api/v1/qc/issues/{id}/close', 'request': 'resolution*, closedBy*, closedAt', 'response': 'issueId, status: CLOSED', 'desc': '이슈 완료 처리'},
            {'method': 'GET', 'url': '/api/v1/qc/issues/report', 'request': 'year, month', 'response': '월별 이슈 요약 리포트 (유형별 건수, 해결률, 재발률)', 'desc': '이슈 월별 리포트'},
        ],
        logic=(
            '- 이슈 소스: DEFECT(불량), METAL_DETECTION(금속검출), CCP_DEVIATION(CCP이탈), MANUAL(수동등록)\n'
            '- 이슈 상태 흐름: OPEN → IN_PROGRESS → CLOSED (재오픈 가능)\n'
            '- 재발률 계산: 동일 유형 이슈가 30일 내 재발 시 RECURRENCE 태그 자동 부여\n'
            '- 완료된 이슈 데이터 → AI Agent 분석 학습 데이터로 활용'
        ),
        exception=(
            '- IssueAlreadyClosedException: 완료된 이슈 재완료 시도 시 409 반환\n'
            '- ResolutionRequiredException: 조치 내용 없이 완료 처리 시 400 반환'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# 5장. 시스템관리 (SYS)
# ──────────────────────────────────────────────

def add_chapter5_sys(doc: Document):
    add_heading(doc, '5. 시스템관리 (SYS) 프로그램 설계', 1)
    add_body(doc, '시스템관리(SYS) 모듈은 사용자 계정, 권한, 시스템 로그, 데이터 백업 등 시스템 운영에 필요한 관리 기능을 제공한다.')
    doc.add_paragraph()

    add_heading(doc, '5.1 사용자계정등록/조회', 2)
    add_program_section(
        doc,
        prog_id='SYS-001',
        prog_name='사용자계정등록/조회',
        desc='MES 시스템 사용자 계정을 등록·수정·삭제·조회한다. 계정별 역할(Role) 및 접근 가능 메뉴를 관리한다.',
        flow=[
            'GET /api/v1/sys/users - 사용자 목록 조회 (관리자만 접근)',
            '신규 계정 등록: POST /api/v1/sys/users → 임시 비밀번호 이메일 발송',
            '계정 비활성화: PATCH /api/v1/sys/users/{id}/status',
            '비밀번호 초기화: POST /api/v1/sys/users/{id}/reset-password',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/sys/users', 'request': 'username, roleId, status, page, size', 'response': 'userId, username, name, email, roles, lastLoginAt, status', 'desc': '사용자 목록 조회'},
            {'method': 'POST', 'url': '/api/v1/sys/users', 'request': 'username*, name*, email*, roleIds*, department', 'response': 'userId, username, tempPassword', 'desc': '사용자 등록'},
            {'method': 'PUT', 'url': '/api/v1/sys/users/{id}', 'request': 'name, email, roleIds, department', 'response': 'userId, updatedAt', 'desc': '사용자 수정'},
            {'method': 'PATCH', 'url': '/api/v1/sys/users/{id}/status', 'request': 'status (ACTIVE/INACTIVE)', 'response': 'userId, status', 'desc': '계정 상태 변경'},
            {'method': 'POST', 'url': '/api/v1/sys/users/{id}/reset-password', 'request': '-', 'response': 'message: 임시비밀번호 이메일 발송 완료', 'desc': '비밀번호 초기화'},
        ],
        logic=(
            '- 비밀번호 정책: 최소 8자, 영문+숫자+특수문자 조합, 90일마다 변경 강제\n'
            '- 비밀번호 암호화: BCryptPasswordEncoder (strength 12)\n'
            '- 5회 연속 로그인 실패 시 계정 30분 잠금 (Redis Counter 활용)\n'
            '- 임시 비밀번호: 8자 랜덤 생성, 최초 로그인 시 변경 강제\n'
            '- JWT 토큰: Access Token(30분) + Refresh Token(7일) 구조'
        ),
        exception=(
            '- DuplicateUsernameException: 중복 사용자명 등록 시 409 반환\n'
            '- AccountLockedException: 계정 잠금 시 423 반환 (잠금 해제 시간 포함)\n'
            '- SelfDeactivationException: 자기 자신 계정 비활성화 시도 시 400 반환'
        )
    )

    add_heading(doc, '5.2 권한역할등록/조회', 2)
    add_program_section(
        doc,
        prog_id='SYS-002',
        prog_name='권한역할등록/조회',
        desc='RBAC(Role-Based Access Control) 기반 역할을 정의하고 메뉴별 접근 권한을 관리한다.',
        flow=[
            'GET /api/v1/sys/roles - 역할 목록 조회',
            '역할 등록: POST /api/v1/sys/roles → 메뉴 권한 매핑',
            '권한 매트릭스 조회: GET /api/v1/sys/roles/{id}/permissions',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/sys/roles', 'request': '-', 'response': 'roleId, roleName, roleCode, permissions[{menuId, menuName, read, write, delete}]', 'desc': '역할 목록 및 권한 조회'},
            {'method': 'POST', 'url': '/api/v1/sys/roles', 'request': 'roleName*, roleCode*, permissions[{menuId, read, write, delete}]', 'response': 'roleId, roleName', 'desc': '역할 등록'},
            {'method': 'PUT', 'url': '/api/v1/sys/roles/{id}/permissions', 'request': 'permissions[{menuId, read, write, delete}]', 'response': 'roleId, updatedPermissions', 'desc': '역할 권한 수정'},
        ],
        logic=(
            '- 기본 역할: ADMIN(전체), MANAGER(관리자), OPERATOR(현장작업자), VIEWER(조회전용)\n'
            '- Spring Security @PreAuthorize 어노테이션으로 API 레벨 권한 제어\n'
            '- 권한 데이터 Redis 캐싱: "role::{userId}" 키, 로그인 시 갱신\n'
            '- ADMIN 역할 삭제 불가 (시스템 예약 역할)'
        ),
        exception=(
            '- SystemRoleModifyException: 시스템 예약 역할 삭제 시 403 반환\n'
            '- RoleInUseException: 사용자가 배정된 역할 삭제 시도 시 409 반환'
        )
    )

    add_heading(doc, '5.3 시스템로그조회', 2)
    add_program_section(
        doc,
        prog_id='SYS-003',
        prog_name='시스템로그조회',
        desc='사용자 활동 로그(로그인/로그아웃, API 호출, 데이터 변경)를 조회하고 보안 감사에 활용한다.',
        flow=[
            'AOP 어드바이스로 모든 API 호출 자동 로그 기록',
            'GET /api/v1/sys/logs - 다중 조건 로그 조회',
            '이상 행위 패턴 감지: 동일 IP에서 단시간 다량 요청 탐지',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/sys/logs', 'request': 'userId, logType (LOGIN/API/DATA_CHANGE), startDatetime, endDatetime, page, size', 'response': 'logId, userId, username, logType, action, targetResource, ipAddress, requestAt, responseTime', 'desc': '시스템 로그 조회'},
            {'method': 'GET', 'url': '/api/v1/sys/logs/export', 'request': '조회 조건 동일', 'response': 'application/vnd.ms-excel', 'desc': '로그 엑셀 내보내기'},
        ],
        logic=(
            '- 로그 저장: AOP @Around 어드바이스로 요청/응답 자동 기록\n'
            '- 민감 데이터 마스킹: 비밀번호, 개인정보 필드 로그에서 자동 제거\n'
            '- 로그 보관: 1년 (법적 의무 보관 기간)\n'
            '- 로그 파티셔닝: MySQL 월별 파티션으로 성능 최적화'
        ),
        exception=(
            '- LogExportLimitException: 10만 건 초과 내보내기 시 400 반환'
        )
    )

    add_heading(doc, '5.4 데이터백업이력조회', 2)
    add_program_section(
        doc,
        prog_id='SYS-004',
        prog_name='데이터백업이력조회',
        desc='시스템 데이터의 자동 백업 실행 이력과 결과를 조회한다.',
        flow=[
            '자동 백업: Spring Batch Job, 매일 새벽 2시 실행',
            'mysqldump 기반 전체 백업 → AWS S3 업로드',
            'GET /api/v1/sys/backups - 백업 이력 조회',
            '백업 실패 시 관리자 이메일/Slack 알림',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/sys/backups', 'request': 'startDate, endDate, status (SUCCESS/FAIL)', 'response': 'backupId, backupType, status, fileSize, s3Path, startedAt, completedAt, duration', 'desc': '백업 이력 조회'},
            {'method': 'POST', 'url': '/api/v1/sys/backups/manual', 'request': '-', 'response': 'backupId, status: IN_PROGRESS', 'desc': '수동 백업 실행 (관리자 전용)'},
        ],
        logic=(
            '- 백업 주기: 전체 백업 1회/일 (02:00), 증분 백업 1회/6시간\n'
            '- 백업 파일 보관: S3 Standard 30일, S3 Glacier 1년\n'
            '- 백업 검증: 백업 완료 후 복원 가능 여부 MD5 체크섬 검증'
        ),
        exception=(
            '- BackupInProgressException: 백업 실행 중 수동 백업 중복 시도 시 409 반환\n'
            '- S3UploadException: S3 업로드 실패 시 재시도 3회 후 FAIL 상태 기록'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# 6장. KPI모니터링 (KPI)
# ──────────────────────────────────────────────

def add_chapter6_kpi(doc: Document):
    add_heading(doc, '6. KPI모니터링 (KPI) 프로그램 설계', 1)
    add_body(doc, 'KPI모니터링(KPI) 모듈은 생산성, 품질, 설비 가동률 등 핵심 성과 지표를 실시간으로 집계하고 대시보드에 시각화한다.')
    doc.add_paragraph()

    add_heading(doc, '6.1 생산성KPI조회', 2)
    add_program_section(
        doc,
        prog_id='KPI-001',
        prog_name='생산성KPI조회',
        desc='일별/주별/월별 생산량, 계획 대비 달성률, 공정별 사이클 타임, 라인 가동률 등 생산성 KPI를 조회한다.',
        flow=[
            '작업지시서 완료 이벤트 수신 → KPI 집계 Spring Batch 실행',
            'GET /api/v1/kpi/productivity - 생산성 KPI 조회',
            'Redis에 캐싱된 실시간 KPI 데이터 우선 조회',
            '대시보드: Chart.js 기반 시각화',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/kpi/productivity', 'request': 'period (DAILY/WEEKLY/MONTHLY), startDate, endDate, lineNo', 'response': 'planQty, actualQty, achievementRate, operatingTime, downtime, cycleTime, lineEfficiency', 'desc': '생산성 KPI 조회'},
            {'method': 'GET', 'url': '/api/v1/kpi/productivity/realtime', 'request': '-', 'response': '당일 실시간 생산 현황 (SSE)', 'desc': '생산성 실시간 현황'},
            {'method': 'GET', 'url': '/api/v1/kpi/productivity/trend', 'request': 'startDate, endDate, period', 'response': '기간별 생산성 트렌드 데이터 (차트용)', 'desc': '생산성 트렌드 조회'},
        ],
        logic=(
            '- 계획달성률 = actualQty / planQty × 100\n'
            '- 라인효율 = 실제가동시간 / 계획가동시간 × 100\n'
            '- 집계 배치: @Scheduled(fixedDelay=300000) 5분마다 Redis KPI 갱신\n'
            '- 일별 집계: Spring Batch Job, 매일 23:50 실행 → 일별 KPI 테이블 적재'
        ),
        exception=(
            '- KpiCalculationException: 분모가 0인 경우 해당 지표 null 반환\n'
            '- KpiPeriodException: 조회 기간 2년 초과 시 400 반환'
        )
    )

    add_heading(doc, '6.2 품질KPI조회', 2)
    add_program_section(
        doc,
        prog_id='KPI-002',
        prog_name='품질KPI조회',
        desc='불량률, CCP 이탈 건수, 검사 합격률, 금속 검출 건수 등 품질 KPI를 조회한다.',
        flow=[
            'GET /api/v1/kpi/quality - 품질 KPI 데이터 조회',
            '공정별 불량률 파레토 차트 데이터 제공',
            '월간 품질 트렌드 분석',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/kpi/quality', 'request': 'period, startDate, endDate', 'response': 'defectRate, ccpDeviationCount, inspectionPassRate, metalDetectionCount, firstPassYield', 'desc': '품질 KPI 조회'},
            {'method': 'GET', 'url': '/api/v1/kpi/quality/pareto', 'request': 'startDate, endDate', 'response': '불량 유형별 건수 및 누적 비율 (파레토)', 'desc': '불량 파레토 분석'},
        ],
        logic=(
            '- 불량률 = 불량수 / 생산수 × 100\n'
            '- First Pass Yield = 첫 번째 검사 합격수 / 전체 검사수 × 100\n'
            '- 품질 점수(Quality Score) = (100 - 불량률) × 합격률 / 100 (종합 품질 지수)\n'
            '- 목표 불량률: 0.5% 이하 (임계값 설정은 KPI-003에서 관리)'
        ),
        exception=(
            '- KpiDataNotFoundException: 해당 기간 데이터 없을 시 빈 배열 반환 (404 아님)'
        )
    )

    add_heading(doc, '6.3 KPI지표관리', 2)
    add_program_section(
        doc,
        prog_id='KPI-003',
        prog_name='KPI지표관리',
        desc='KPI 목표값, 임계값, 측정 주기 등 KPI 지표 정의를 관리한다.',
        flow=[
            'GET /api/v1/kpi/indicators - KPI 지표 목록 조회',
            '지표별 목표값, 경보 임계값 설정: PUT /api/v1/kpi/indicators/{id}',
            '지표 달성 여부 자동 판정 및 색상 표시 (RAG: Red/Amber/Green)',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/kpi/indicators', 'request': '-', 'response': '지표 목록 (indicatorId, name, targetValue, warningThreshold, criticalThreshold, unit, period)', 'desc': 'KPI 지표 목록 조회'},
            {'method': 'PUT', 'url': '/api/v1/kpi/indicators/{id}', 'request': 'targetValue, warningThreshold, criticalThreshold', 'response': 'indicatorId, updatedAt', 'desc': 'KPI 목표값 수정'},
        ],
        logic=(
            '- RAG 판정: GREEN(목표 달성), AMBER(Warning 임계값 이탈), RED(Critical 임계값 이탈)\n'
            '- KPI 지표 변경 이력 자동 관리\n'
            '- 지표 변경 시 Redis KPI 캐시 즉시 갱신 (Cache Evict)'
        ),
        exception=(
            '- InvalidThresholdException: warning > critical 설정 시 400 반환'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# 7장. 설비관리 (EQP)
# ──────────────────────────────────────────────

def add_chapter7_eqp(doc: Document):
    add_heading(doc, '7. 설비관리 (EQP) 프로그램 설계', 1)
    add_body(doc, '설비관리(EQP, Equipment Management) 모듈은 생산 현장 설비의 가동 현황, 점검, 고장 이력을 관리한다.')
    doc.add_paragraph()

    add_heading(doc, '7.1 설비가동현황조회', 2)
    add_program_section(
        doc,
        prog_id='EQP-001',
        prog_name='설비가동현황조회',
        desc='전체 설비의 실시간 가동 상태(운전/정지/점검/고장)와 가동 시간을 조회한다.',
        flow=[
            'MQTT에서 설비 상태 데이터 수신 → Redis 상태 갱신',
            'GET /api/v1/eqp/status - 전체 설비 가동 현황 조회',
            'SSE로 설비 상태 변경 시 대시보드 실시간 반영',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/eqp/status', 'request': 'lineNo(선택), processCode(선택)', 'response': 'eqpCode, eqpName, status (RUNNING/STOPPED/MAINTENANCE/FAULT), operatingHours, lastStatusChangedAt', 'desc': '설비 가동 현황 조회'},
            {'method': 'GET', 'url': '/api/v1/eqp/status/realtime', 'request': '-', 'response': 'text/event-stream (설비 상태 변경 이벤트)', 'desc': '설비 상태 실시간 모니터링'},
            {'method': 'GET', 'url': '/api/v1/eqp/status/oee', 'request': 'eqpCode, date', 'response': '일별 OEE 지표 (가용성, 성능, 품질)', 'desc': '설비별 OEE 조회'},
        ],
        logic=(
            '- 설비 상태 Redis Hash: "eqp:status:{eqpCode}" → {status, updatedAt, operatingHours}\n'
            '- 가동 시간 누적: RUNNING 상태 유지 시간을 분 단위로 적산\n'
            '- FAULT 상태 발생 시 설비이상등록(EQP-003) 자동 트리거\n'
            '- MQTT 토픽: mes/equipment/{eqpCode}/status'
        ),
        exception=(
            '- EquipmentNotFoundException: 미등록 설비코드 데이터 수신 시 경고 로그 기록'
        )
    )

    add_heading(doc, '7.2 설비점검등록 및 이력조회', 2)
    add_program_section(
        doc,
        prog_id='EQP-002',
        prog_name='설비점검등록/이력조회',
        desc='설비 정기 점검 계획을 수립하고, 점검 실시 결과를 등록하며 점검 이력을 관리한다.',
        flow=[
            '점검 계획 조회: GET /api/v1/eqp/inspections/schedule',
            '점검 실시: 점검자, 점검 항목별 결과 입력 → POST /api/v1/eqp/inspections',
            '점검 완료: PATCH /api/v1/eqp/inspections/{id}/complete',
            '정기점검 누락 시 알림 발생',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/eqp/inspections', 'request': 'eqpCode, startDate, endDate, status', 'response': '점검 이력 목록', 'desc': '설비 점검 이력 조회'},
            {'method': 'POST', 'url': '/api/v1/eqp/inspections', 'request': 'eqpCode*, inspectorId*, inspectionItems[{itemName, result, notes}]*, inspectedAt', 'response': 'inspectionId, status', 'desc': '설비 점검 등록'},
            {'method': 'GET', 'url': '/api/v1/eqp/inspections/schedule', 'request': 'eqpCode, month', 'response': '점검 계획 목록 (예정일, 점검유형, 담당자)', 'desc': '점검 계획 조회'},
        ],
        logic=(
            '- 점검 유형: DAILY(일상점검), WEEKLY(주간), MONTHLY(월간), ANNUAL(연간)\n'
            '- 점검 주기 초과 시: 해당 설비 MAINTENANCE 상태 자동 전환\n'
            '- 점검 항목 불합격 시 설비이상 자동 등록 (EQP-003 연동)'
        ),
        exception=(
            '- InspectionOverdueException: 점검 주기 초과(1일 이상) 시 담당자 알림\n'
            '- DuplicateInspectionException: 동일 설비 당일 동일 유형 점검 중복 등록 시 409 반환'
        )
    )

    add_heading(doc, '7.3 설비이상등록 및 고장이력조회', 2)
    add_program_section(
        doc,
        prog_id='EQP-003',
        prog_name='설비이상등록/고장이력조회',
        desc='설비 이상 및 고장 발생을 등록하고, 수리 완료까지의 이력을 관리한다. MTBF, MTTR 등 설비 신뢰성 지표를 제공한다.',
        flow=[
            '이상 발생 (자동: MQTT FAULT / 수동: 작업자 입력)',
            'POST /api/v1/eqp/failures → 이상 등록 → 수리 담당자 알림',
            '수리 진행: PATCH /api/v1/eqp/failures/{id}/repair',
            '수리 완료: PATCH /api/v1/eqp/failures/{id}/resolve → 설비 RUNNING 복구',
            'MTBF/MTTR 자동 계산 및 KPI 연동',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/eqp/failures', 'request': 'eqpCode, failureType, startDate, endDate', 'response': '고장 이력 목록 (failureId, eqpCode, failureType, failureAt, resolvedAt, downtime)', 'desc': '고장 이력 조회'},
            {'method': 'POST', 'url': '/api/v1/eqp/failures', 'request': 'eqpCode*, failureType*, symptom*, reportedBy*, failedAt', 'response': 'failureId, status: OPEN', 'desc': '설비 이상 등록'},
            {'method': 'PATCH', 'url': '/api/v1/eqp/failures/{id}/resolve', 'request': 'repairContent*, resolvedBy*, resolvedAt', 'response': 'failureId, downtime, status: RESOLVED', 'desc': '고장 수리 완료'},
            {'method': 'GET', 'url': '/api/v1/eqp/failures/reliability', 'request': 'eqpCode, period', 'response': 'mtbf(평균고장간격), mttr(평균수리시간), availability(%)', 'desc': '설비 신뢰성 지표 조회'},
        ],
        logic=(
            '- MTBF(Mean Time Between Failures) = 총 가동시간 / 고장 횟수\n'
            '- MTTR(Mean Time To Repair) = 총 수리시간 / 고장 횟수\n'
            '- 가용성(Availability) = MTBF / (MTBF + MTTR) × 100\n'
            '- 고장 유형: MECHANICAL(기계), ELECTRICAL(전기), SENSOR(센서), SOFTWARE(소프트웨어)\n'
            '- 고장 등록 즉시 생산계획 조정 가능 여부 PM 모듈에 알림 발행'
        ),
        exception=(
            '- FailureAlreadyResolvedException: 이미 완료된 고장 재완료 시 409 반환\n'
            '- EqpAlreadyFaultException: 이미 FAULT 상태 설비 중복 이상 등록 시 409 반환'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# 8장. POP작업 (POP)
# ──────────────────────────────────────────────

def add_chapter8_pop(doc: Document):
    add_heading(doc, '8. POP작업 (POP) 프로그램 설계', 1)
    add_body(doc,
        'POP(Point of Production) 작업 모듈은 현장 작업자가 터치스크린 또는 태블릿에서 '
        '직접 생산 실적을 입력하고 조회하는 현장 중심 인터페이스를 제공한다.')
    doc.add_paragraph()

    add_heading(doc, '8.1 POP 작업지시서목록/라인선택/생산진행관리', 2)
    add_program_section(
        doc,
        prog_id='POP-001',
        prog_name='POP작업 핵심 기능',
        desc='POP 화면에서 작업자가 라인을 선택하고 작업지시서를 조회하여 생산 시작/진행/완료 처리를 수행한다.',
        flow=[
            '작업자 로그인: QR코드 스캔 또는 사원번호+PIN 입력',
            '라인 선택: GET /api/v1/pop/lines → 현재 가용 라인 목록 표시',
            '작업지시서 목록: GET /api/v1/pop/work-orders?lineNo={no}',
            '작업 시작: POST /api/v1/pop/work-orders/{id}/start',
            '생산 진행: 실시간 진행률 표시 (WebSocket)',
            '공정별 실적 입력: POST /api/v1/pop/process-results',
            '작업 완료: POST /api/v1/pop/work-orders/{id}/complete',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/pop/lines', 'request': '-', 'response': '가용 라인 목록 (lineNo, lineName, status, currentWorkOrder)', 'desc': 'POP 라인 목록 조회'},
            {'method': 'GET', 'url': '/api/v1/pop/work-orders', 'request': 'lineNo*(필수), date', 'response': '당일 작업지시서 목록 (workOrderId, productName, planQty, status, startTime)', 'desc': 'POP 작업지시서 목록'},
            {'method': 'POST', 'url': '/api/v1/pop/work-orders/{id}/start', 'request': 'workerId*, actualStartTime', 'response': 'workOrderId, status: IN_PROGRESS', 'desc': '작업 시작'},
            {'method': 'POST', 'url': '/api/v1/pop/process-results', 'request': 'workOrderId*, processCode*, resultQty, defectQty, workerId*, recordedAt', 'response': 'resultId, cumulativeQty, progress(%)', 'desc': '공정별 실적 입력'},
            {'method': 'POST', 'url': '/api/v1/pop/work-orders/{id}/complete', 'request': 'actualQty*, defectQty, actualEndTime, notes', 'response': 'workOrderId, status: COMPLETED, yieldRate', 'desc': '작업 완료'},
            {'method': 'GET', 'url': '/api/v1/pop/work-orders/{id}/progress', 'request': 'id (Path)', 'response': '현재 공정, 진행률, 목표 대비 현황, CCP 상태', 'desc': '생산 진행 현황 조회'},
        ],
        logic=(
            '- POP UI: React SPA, 터치 최적화 UI (버튼 크기 최소 44px × 44px)\n'
            '- 오프라인 대응: PWA(Progressive Web App) 구현, 네트워크 단절 시 로컬 스토리지 임시 저장 후 온라인 복구 시 동기화\n'
            '- WebSocket 연결: ws://{host}/ws/pop/progress → 실시간 진행률 Push\n'
            '- 바코드/QR 스캔 지원: 작업자 인식, 제품 확인, LOT 추적\n'
            '- 화면 세션: 30분 무활동 시 자동 로그아웃 (보안)'
        ),
        exception=(
            '- WorkerNotAuthorizedForLineException: 해당 라인 미배정 작업자 접근 시 403 반환\n'
            '- WorkOrderAlreadyCompletedException: 완료된 작업지시서 재시작 시 409 반환\n'
            '- OfflineSyncConflictException: 오프라인 데이터 동기화 충돌 시 수동 확인 요청'
        )
    )

    add_heading(doc, '8.2 POP 설비목록/발주목록/출고지시서목록', 2)
    add_program_section(
        doc,
        prog_id='POP-002',
        prog_name='POP 설비/발주/출고 조회',
        desc='POP 화면에서 현장 작업자가 관련 설비 상태, 자재 발주 현황, 출고 지시서를 간편하게 조회한다.',
        flow=[
            'GET /api/v1/pop/equipments - 현재 라인 내 설비 목록 및 상태 조회',
            'GET /api/v1/pop/purchase-orders - 발주 현황 조회 (자재 부족 확인)',
            'GET /api/v1/pop/shipping-orders - 출고 지시서 목록 조회',
        ],
        apis=[
            {'method': 'GET', 'url': '/api/v1/pop/equipments', 'request': 'lineNo*(필수)', 'response': 'eqpCode, eqpName, status, lastCheckTime', 'desc': 'POP 설비 목록 조회'},
            {'method': 'GET', 'url': '/api/v1/pop/purchase-orders', 'request': 'status (PENDING/ORDERED/RECEIVED)', 'response': '발주 목록 (materialName, orderedQty, expectedDate, status)', 'desc': 'POP 발주 목록 조회'},
            {'method': 'GET', 'url': '/api/v1/pop/shipping-orders', 'request': 'shippingDate, status', 'response': '출고 지시서 목록 (orderId, partnerName, productName, shippingQty, shippingDate)', 'desc': 'POP 출고 지시서 목록'},
        ],
        logic=(
            '- POP 조회 데이터는 Redis 캐시 우선 조회 (최대 2분 지연 허용)\n'
            '- 설비 FAULT 상태 시 빨간색 아이콘 표시 + 터치 시 이상 신고 기능 제공\n'
            '- 발주/출고 조회는 간략 정보만 제공 (상세는 PC 화면에서 처리)'
        ),
        exception=(
            '- LineNotSelectedException: 라인 미선택 상태에서 설비 조회 시 400 반환'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# 9장. AI Agent (AI)
# ──────────────────────────────────────────────

def add_chapter9_ai(doc: Document):
    add_heading(doc, '9. AI Agent (AI) 프로그램 설계', 1)
    add_body(doc,
        'AI Agent 모듈은 MES 데이터를 기반으로 자연어 질의응답, 생산/품질 자동 분석 요약, '
        '일간/주간 리포트 자동 생성 등 지능형 의사결정 지원 기능을 제공한다.')
    doc.add_paragraph()

    add_heading(doc, '9.1 RAG 파이프라인 설계', 2)
    p = doc.add_paragraph()
    r = p.add_run('■ RAG(Retrieval-Augmented Generation) 아키텍처')
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    rag_design = [
        ('문서 임베딩', '- MES 운영 매뉴얼, CCP 기준서, 레시피 문서 → OpenAI text-embedding-ada-002\n- 벡터 DB: PostgreSQL + pgvector (또는 Pinecone)\n- 임베딩 업데이트: 문서 변경 시 자동 재임베딩'),
        ('검색 (Retrieval)', '- 사용자 질의 → 질의 임베딩 생성 → 코사인 유사도 Top-K(5) 문서 검색\n- 하이브리드 검색: 벡터 유사도 + MySQL 키워드 검색 결합\n- 필터링: 날짜, 모듈, 문서 유형별 메타데이터 필터 적용'),
        ('생성 (Generation)', '- LLM: OpenAI GPT-4o-mini (비용 최적화) 또는 Claude claude-sonnet-4-6\n- 프롬프트 구성: [시스템 역할] + [검색된 컨텍스트] + [사용자 질문]\n- 응답 언어: 한국어 기본, 설정으로 영어 전환 가능\n- 스트리밍 응답: Server-Sent Events로 실시간 토큰 전송'),
        ('MES 데이터 연동', '- 실시간 DB 조회: 생산 실적, KPI, 재고 데이터 → 함수 호출(Function Calling)로 처리\n- 캐싱: 동일 질의 24시간 내 재질의 시 캐시 응답 반환 (Redis)\n- 응답 품질 평가: Thumbs Up/Down 피드백 수집 → 모델 파인튜닝 데이터 활용'),
    ]
    tbl = doc.add_table(rows=1 + len(rag_design), cols=2)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    hdr[0].text = '구성요소'
    hdr[1].text = '설계 내용'
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_bg(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, (k, v) in enumerate(rag_design, 1):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl.rows[i].cells[0], 'D9E1F2')
        tbl.rows[i].cells[1].text = v
        tbl.rows[i].cells[0].width = Cm(3.5)
        tbl.rows[i].cells[1].width = Cm(13.5)
    doc.add_paragraph()

    add_heading(doc, '9.2 AI질의등록/응답조회', 2)
    add_program_section(
        doc,
        prog_id='AI-001',
        prog_name='AI질의등록/응답조회',
        desc='사용자가 자연어로 MES 관련 질문을 입력하면 RAG 기반 AI가 답변을 생성한다. 질의 이력 및 피드백을 관리한다.',
        flow=[
            '사용자 질의 입력: POST /api/v1/ai/queries',
            '질의 임베딩 생성 → 벡터 DB 유사 문서 검색',
            'MES DB 조회 (Function Calling): 생산실적, KPI, 재고 데이터 실시간 조회',
            'LLM에 컨텍스트 + 질의 전달 → 응답 스트리밍(SSE)',
            '응답 저장 → GET /api/v1/ai/queries/{id}/response로 조회',
            '사용자 피드백 수집: POST /api/v1/ai/queries/{id}/feedback',
        ],
        apis=[
            {'method': 'POST', 'url': '/api/v1/ai/queries', 'request': 'question*(자연어 질문), contextType (PRODUCTION/QUALITY/INVENTORY/GENERAL, 선택)', 'response': 'queryId, status: PROCESSING', 'desc': 'AI 질의 등록'},
            {'method': 'GET', 'url': '/api/v1/ai/queries/{id}/stream', 'request': 'id (Path)', 'response': 'text/event-stream (LLM 응답 스트리밍)', 'desc': 'AI 응답 실시간 스트리밍'},
            {'method': 'GET', 'url': '/api/v1/ai/queries/{id}', 'request': 'id (Path)', 'response': 'queryId, question, answer, sources[{title, excerpt}], tokens, responseTime', 'desc': 'AI 응답 조회'},
            {'method': 'POST', 'url': '/api/v1/ai/queries/{id}/feedback', 'request': 'rating (GOOD/BAD), comment', 'response': 'feedbackId', 'desc': 'AI 응답 피드백 등록'},
            {'method': 'GET', 'url': '/api/v1/ai/queries/history', 'request': 'userId, startDate, endDate, page, size', 'response': 'AI 질의 이력 목록', 'desc': 'AI 분석 이력 조회'},
        ],
        logic=(
            '- 질의 처리 비동기: @Async + CompletableFuture, 최대 처리 시간 30초\n'
            '- Function Calling 정의: getCurrentStock(), getKpiSummary(), getWorkOrderStatus(), getDefectReport()\n'
            '- 응답 토큰 제한: 최대 2,048 토큰 (비용 제어)\n'
            '- 부적절한 질의 필터: 개인정보, 경쟁사 정보 요청 시 거부 프롬프트 적용\n'
            '- 다국어 지원: 질의 언어 자동 감지 후 동일 언어로 응답'
        ),
        exception=(
            '- LlmTimeoutException: LLM 응답 30초 초과 시 타임아웃 오류 반환\n'
            '- QuestionTooLongException: 질문 1,000자 초과 시 400 반환\n'
            '- RateLimitException: 동일 사용자 분당 10회 초과 질의 시 429 반환'
        )
    )

    add_heading(doc, '9.3 자동 요약 및 리포트 생성', 2)
    add_program_section(
        doc,
        prog_id='AI-002',
        prog_name='자동 요약 및 리포트 생성',
        desc='생산 요약, 불량 요약, 비가동 요약을 자동 생성하고, 일간/주간 리포트를 PDF 또는 이메일로 자동 발송한다.',
        flow=[
            '일일 리포트: Spring Batch Job, 매일 23:55 실행',
            'KPI, 생산실적, 불량, 설비 데이터 집계 → LLM에 데이터 전달',
            'LLM이 자연어 요약 + 이상 항목 하이라이트 → 리포트 생성',
            'PDF 변환: iText 라이브러리 → S3 저장',
            '관리자 이메일/Slack 자동 발송',
            '주간 리포트: 매주 금요일 23:55, 형식 동일하나 7일 데이터 집계',
        ],
        apis=[
            {'method': 'POST', 'url': '/api/v1/ai/reports/daily', 'request': 'targetDate* (YYYY-MM-DD)', 'response': 'reportId, status: GENERATING', 'desc': '일일 리포트 생성 요청'},
            {'method': 'POST', 'url': '/api/v1/ai/reports/weekly', 'request': 'startDate*, endDate*', 'response': 'reportId, status: GENERATING', 'desc': '주간 리포트 생성 요청'},
            {'method': 'GET', 'url': '/api/v1/ai/reports', 'request': 'reportType (DAILY/WEEKLY), startDate, endDate', 'response': '리포트 목록 (reportId, type, targetDate, status, pdfUrl)', 'desc': '리포트 목록 조회'},
            {'method': 'GET', 'url': '/api/v1/ai/reports/{id}', 'request': 'id (Path)', 'response': 'reportId, title, summaryText, highlights, pdfUrl, generatedAt', 'desc': '리포트 상세 조회'},
            {'method': 'GET', 'url': '/api/v1/ai/summaries/production', 'request': 'date', 'response': '생산 요약 (LLM 자연어 요약 + 수치 데이터)', 'desc': '생산 요약 조회'},
            {'method': 'GET', 'url': '/api/v1/ai/summaries/defects', 'request': 'date', 'response': '불량 요약 (불량 유형별 분석 + LLM 원인 추론)', 'desc': '불량 요약 조회'},
            {'method': 'GET', 'url': '/api/v1/ai/summaries/downtime', 'request': 'date', 'response': '비가동 요약 (설비별 비가동 원인 + 영향 분석)', 'desc': '비가동 요약 조회'},
        ],
        logic=(
            '- 리포트 구성: [헤더(날짜/버전)] + [생산 요약] + [품질 현황] + [설비 현황] + [AI 분석 코멘트] + [다음 날 주요 계획]\n'
            '- LLM 프롬프트: 데이터를 JSON으로 제공 → "한국어로 경영자가 이해하기 쉽게 요약하라" 지시\n'
            '- PDF 레이아웃: iText + FreeMarker 템플릿 엔진으로 표/차트 삽입\n'
            '- 이상 하이라이트: KPI 목표 미달 항목 자동 감지 → 빨간색 박스로 강조\n'
            '- 리포트 링크 유효기간: S3 Pre-signed URL 7일\n'
            '- Slack 연동: Webhook URL로 리포트 요약 자동 포스팅'
        ),
        exception=(
            '- ReportGenerationFailedException: LLM 또는 PDF 생성 실패 시 재시도 3회 후 관리자 알림\n'
            '- DataInsufficientException: 집계 데이터 부족(생산 실적 없음) 시 "데이터 없음" 리포트 생성\n'
            '- ReportAlreadyExistsException: 동일 날짜 리포트 중복 생성 시 409 반환'
        )
    )
    doc.add_page_break()


# ──────────────────────────────────────────────
# 10장. 공통 컴포넌트 설계
# ──────────────────────────────────────────────

def add_chapter10_common(doc: Document):
    add_heading(doc, '10. 공통 컴포넌트 설계', 1)

    add_heading(doc, '10.1 인증/인가 (JWT + OAuth2)', 2)
    add_body(doc, '■ 인증 흐름')
    auth_flow = [
        'POST /api/v1/auth/login → ID/PW 검증 → JWT 발급 (Access: 30분, Refresh: 7일)',
        'Access Token 만료 시: POST /api/v1/auth/refresh → Refresh Token으로 재발급',
        'Refresh Token은 Redis에 저장 (키: "refresh::{userId}", TTL 7일)',
        'Logout: POST /api/v1/auth/logout → Refresh Token 삭제 + Access Token 블랙리스트 등록',
        'OAuth2 SSO: Spring Security OAuth2 Client, 사내 AD(Active Directory) 연동 지원',
    ]
    for step in auth_flow:
        p = doc.add_paragraph(step, style='List Number')
        p.paragraph_format.left_indent = Cm(0.5)
    doc.add_paragraph()

    auth_apis = [
        {'method': 'POST', 'url': '/api/v1/auth/login', 'request': 'username*, password*', 'response': 'accessToken, refreshToken, expiresIn, userInfo', 'desc': '로그인'},
        {'method': 'POST', 'url': '/api/v1/auth/refresh', 'request': 'refreshToken*', 'response': 'accessToken, expiresIn', 'desc': 'Access Token 재발급'},
        {'method': 'POST', 'url': '/api/v1/auth/logout', 'request': 'Authorization 헤더', 'response': 'success: true', 'desc': '로그아웃'},
        {'method': 'GET', 'url': '/api/v1/auth/me', 'request': 'Authorization 헤더', 'response': 'userId, username, name, roles, permissions', 'desc': '내 계정 정보 조회'},
    ]
    add_api_table(doc, auth_apis)

    add_body(doc, '■ JWT 토큰 구조')
    jwt_info = (
        '- Header: {"alg": "HS256", "typ": "JWT"}\n'
        '- Payload: {"sub": "{userId}", "username": "...", "roles": [...], "iat": ..., "exp": ...}\n'
        '- Secret Key: 256비트 랜덤 키, 환경변수(JWT_SECRET)로 주입\n'
        '- 알고리즘: HMAC SHA-256 (HS256)\n'
        '- API Gateway 레벨에서 토큰 검증 (JwtAuthenticationFilter)'
    )
    add_body(doc, jwt_info)
    doc.add_paragraph()

    add_heading(doc, '10.2 알람/알림 서비스', 2)
    add_body(doc, '■ 알림 채널 및 우선순위')
    notif_data = [
        ('CRITICAL', 'CCP 이탈, 금속 검출, 냉장고 온도 이상', 'SMS + 이메일 + Slack + 앱 Push'),
        ('WARNING', '안전재고 이하, 납기 임박, 설비 점검 예정', '이메일 + Slack'),
        ('INFO', '작업지시서 생성, 리포트 완성, 배치 완료', '앱 내 알림'),
    ]
    tbl = doc.add_table(rows=1 + len(notif_data), cols=3)
    tbl.style = 'Table Grid'
    hdrs = ['알림 등급', '발생 조건', '전송 채널']
    for i, h in enumerate(hdrs):
        tbl.rows[0].cells[i].text = h
        tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl.rows[0].cells[i], '1F497D')
        tbl.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    colors = {'CRITICAL': 'FCE4D6', 'WARNING': 'FFF2CC', 'INFO': 'E2EFDA'}
    for i, (grade, cond, chan) in enumerate(notif_data, 1):
        tbl.rows[i].cells[0].text = grade
        set_cell_bg(tbl.rows[i].cells[0], colors.get(grade, 'FFFFFF'))
        tbl.rows[i].cells[1].text = cond
        tbl.rows[i].cells[2].text = chan
    doc.add_paragraph()

    notif_apis = [
        {'method': 'GET', 'url': '/api/v1/notifications', 'request': 'userId, isRead(boolean), page, size', 'response': '알림 목록 (notificationId, type, title, message, isRead, createdAt)', 'desc': '내 알림 목록 조회'},
        {'method': 'PATCH', 'url': '/api/v1/notifications/{id}/read', 'request': 'id (Path)', 'response': 'notificationId, isRead: true', 'desc': '알림 읽음 처리'},
        {'method': 'GET', 'url': '/api/v1/notifications/stream', 'request': '-', 'response': 'text/event-stream (신규 알림 실시간)', 'desc': '알림 실시간 수신 (SSE)'},
    ]
    add_api_table(doc, notif_apis)

    add_heading(doc, '10.3 배치 처리 설계', 2)
    batch_data = [
        ('KPI 집계 배치', '매 5분', 'KpiCalculationJob', 'Redis KPI 캐시 갱신'),
        ('일별 KPI 집계', '매일 23:50', 'DailyKpiJob', 'kpi_daily 테이블 적재'),
        ('AI 일간 리포트', '매일 23:55', 'DailyReportJob', 'AI 리포트 생성 + 이메일 발송'),
        ('AI 주간 리포트', '매주 금 23:55', 'WeeklyReportJob', 'AI 주간 리포트 생성'),
        ('안전재고 체크', '매 5분', 'SafetyStockCheckJob', '재고 임계값 비교 + 알림'),
        ('납기 임박 알림', '매일 09:00', 'DeliveryAlertJob', 'D-3 이내 납기 알림 발송'),
        ('DB 백업', '매일 02:00', 'DatabaseBackupJob', 'mysqldump → S3 업로드'),
        ('로그 아카이빙', '매월 1일 03:00', 'LogArchiveJob', '1년 초과 로그 아카이빙'),
    ]
    tbl2 = doc.add_table(rows=1 + len(batch_data), cols=4)
    tbl2.style = 'Table Grid'
    hdr2 = ['배치 잡명', '실행 주기', 'Job 클래스', '처리 내용']
    for i, h in enumerate(hdr2):
        tbl2.rows[0].cells[i].text = h
        tbl2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl2.rows[0].cells[i], '2E75B6')
        tbl2.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row_data in enumerate(batch_data, 1):
        for j, val in enumerate(row_data):
            tbl2.rows[i].cells[j].text = val
        if i % 2 == 0:
            for j in range(4):
                set_cell_bg(tbl2.rows[i].cells[j], 'F2F2F2')
    doc.add_paragraph()

    add_body(doc, '■ 배치 프레임워크: Spring Batch')
    batch_info = (
        '- JobRepository: MySQL 기반 (배치 실행 이력 관리)\n'
        '- JobLauncher: 스케줄러(@Scheduled) + REST API 수동 실행 이중 지원\n'
        '- 청크 처리: ChunkOrientedTasklet, 5,000건 단위 커밋\n'
        '- 재시작: Spring Batch 재시작 기능 활용 (이미 처리된 청크 건너뜀)\n'
        '- 모니터링: Spring Batch Admin 또는 커스텀 배치 현황 API 제공\n'
        '- 장애 알림: 배치 실패 시 Slack + 이메일 즉시 통보'
    )
    add_body(doc, batch_info)
    doc.add_paragraph()

    add_heading(doc, '10.4 공통 응답 형식', 2)
    add_body(doc, '■ 표준 API 응답 구조 (모든 API 공통 적용)')
    response_json = (
        '성공 응답:\n'
        '{\n'
        '  "success": true,\n'
        '  "code": "SUCCESS",\n'
        '  "message": "처리가 완료되었습니다.",\n'
        '  "data": { ... },\n'
        '  "timestamp": "2026-05-12T10:30:00Z"\n'
        '}\n\n'
        '에러 응답:\n'
        '{\n'
        '  "success": false,\n'
        '  "code": "PRODUCT_NOT_FOUND",\n'
        '  "message": "해당 제품을 찾을 수 없습니다.",\n'
        '  "data": null,\n'
        '  "timestamp": "2026-05-12T10:30:00Z"\n'
        '}\n\n'
        '페이징 응답:\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "content": [...],\n'
        '    "page": 0, "size": 20,\n'
        '    "totalElements": 150, "totalPages": 8\n'
        '  }\n'
        '}'
    )
    doc.add_paragraph(response_json).paragraph_format.left_indent = Cm(0.5)
    doc.add_paragraph()

    add_heading(doc, '10.5 성능 및 확장성 설계', 2)
    perf_data = [
        ('API 응답 시간 목표', '일반 조회: 200ms 이내, 복잡 집계: 1,000ms 이내'),
        ('동시 사용자', '최대 50명 동시 접속 (김치 공장 규모 기준)'),
        ('Redis 캐싱 전략', '마스터 데이터: 1시간 TTL, KPI: 5분 TTL, 실시간 상태: 이벤트 기반 갱신'),
        ('DB 인덱스 전략', '주요 조회 컬럼(created_at, status, product_id) 복합 인덱스 적용'),
        ('DB 커넥션 풀', 'HikariCP: 최소 5, 최대 20 커넥션'),
        ('컨테이너 스케일링', 'Docker Compose 기반 (개발/운영), 추후 K8s 전환 고려'),
        ('MQTT 브로커', 'Eclipse Mosquitto, QoS Level 1 (적어도 1회 전송 보장)'),
    ]
    tbl3 = doc.add_table(rows=1 + len(perf_data), cols=2)
    tbl3.style = 'Table Grid'
    for i, h in enumerate(['항목', '내용']):
        tbl3.rows[0].cells[i].text = h
        tbl3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl3.rows[0].cells[i], '1F497D')
        tbl3.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, (k, v) in enumerate(perf_data, 1):
        tbl3.rows[i].cells[0].text = k
        tbl3.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(tbl3.rows[i].cells[0], 'D9E1F2')
        tbl3.rows[i].cells[1].text = v
        tbl3.rows[i].cells[0].width = Cm(4.0)
        tbl3.rows[i].cells[1].width = Cm(13.0)
    doc.add_paragraph()


# ──────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────

def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    add_cover_page(doc)
    add_chapter1_szn(doc)
    add_chapter2_pkg(doc)
    add_chapter3_age(doc)
    add_chapter4_qc(doc)
    add_chapter5_sys(doc)
    add_chapter6_kpi(doc)
    add_chapter7_eqp(doc)
    add_chapter8_pop(doc)
    add_chapter9_ai(doc)
    add_chapter10_common(doc)

    doc.save(OUTPUT_PATH)
    print(f"[완료] 프로그램설계서2 저장: {OUTPUT_PATH}")


if __name__ == '__main__':
    main()
