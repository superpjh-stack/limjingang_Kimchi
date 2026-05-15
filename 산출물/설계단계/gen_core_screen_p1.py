# -*- coding: utf-8 -*-
"""
임진강김치 MES 시스템 화면설계서(코어기능) 생성 스크립트 - Part 1
실행: python gen_core_screen_p1.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r'C:\gerardo\01 SmallSF\Imjingang-Kimchi\산출물\설계단계\06_화면설계서_코어기능.docx'

# ─── 헬퍼 ───────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color='1A73E8'):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_para(doc, text='', bold=False, size=10, color=None, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_ascii(doc, ascii_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F4FF')
    pPr.append(shd)
    run = p.add_run(ascii_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    return p

def add_table(doc, headers, rows, header_color='1A73E8'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        set_cell_bg(c, header_color)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            c = row.cells[i]
            c.text = str(val)
            c.paragraphs[0].runs[0].font.size = Pt(9)
    return table

def screen_block(doc, sid, sname, stype, access, desc, ascii_art,
                 comp_rows, event_rows, valid_rows):
    """화면 1개 전체 블록 출력"""
    # 화면 헤더 표
    doc.add_paragraph()
    t = doc.add_table(rows=2, cols=4)
    t.style = 'Table Grid'
    meta = [('화면 ID', sid), ('화면명', sname), ('화면유형', stype), ('접근권한', access)]
    for i,(k,v) in enumerate(meta):
        c = t.rows[0].cells[i]
        p = c.paragraphs[0]; p.clear()
        r1 = p.add_run(k+': '); r1.bold=True; r1.font.size=Pt(9)
        r2 = p.add_run(v);      r2.font.size=Pt(9)
        set_cell_bg(c, 'E8F0FE')
    # 두번째 행: 화면설명 (4칸 병합)
    merged = t.rows[1].cells[0].merge(t.rows[1].cells[3])
    merged.text = '화면 설명: ' + desc
    merged.paragraphs[0].runs[0].font.size = Pt(9)

    # ASCII 목업
    add_para(doc, '[ 화면 목업 ]', bold=True, size=9, color='1A73E8')
    add_ascii(doc, ascii_art)

    # 주요 컴포넌트
    add_para(doc, '[ 주요 컴포넌트 ]', bold=True, size=9, color='1A73E8')
    add_table(doc, ['컴포넌트명','유형','기능 설명'], comp_rows, '34495E')

    # 이벤트 처리
    add_para(doc, '[ 이벤트 처리 ]', bold=True, size=9, color='1A73E8')
    add_table(doc, ['이벤트','트리거','처리 내용'], event_rows, '34495E')

    # 유효성 검사
    add_para(doc, '[ 유효성 검사 규칙 ]', bold=True, size=9, color='1A73E8')
    add_table(doc, ['필드명','규칙','오류 메시지'], valid_rows, '34495E')

    doc.add_paragraph()
    p = doc.add_paragraph('─'*90)
    p.runs[0].font.size = Pt(7)
    p.runs[0].font.color.rgb = RGBColor(0xCC,0xCC,0xCC)


# ─── 문서 초기화 ─────────────────────────────────────────────

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
sec.left_margin = Cm(2.0); sec.right_margin = Cm(1.5)
sec.top_margin = Cm(2.0);  sec.bottom_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)

# ═══════════════ 표지 ═══════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('화면설계서 (코어기능)')
r.font.size = Pt(28); r.font.bold = True
r.font.color.rgb = RGBColor(0x1A,0x73,0xE8)
r.font.name = '맑은 고딕'

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('임진강김치 MES 시스템 구축')
r2.font.size = Pt(18); r2.font.bold = True

doc.add_paragraph(); doc.add_paragraph()

ct = doc.add_table(rows=4, cols=2)
ct.style = 'Table Grid'
ct.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate([('프로젝트명','임진강김치 MES 시스템 구축'),
                           ('문서명','화면설계서 (코어기능)'),
                           ('작성일','2026-05-12'),('버전','V1.0')]):
    ct.rows[i].cells[0].text = k
    ct.rows[i].cells[1].text = v
    ct.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_bg(ct.rows[i].cells[0], 'E8F0FE')

doc.add_page_break()

# ═══════════════ 목차 ════════════════════════════════════════
add_heading(doc, '목차', 1)
toc = [
  ('1','개요'),('1.1','화면 설계 원칙'),('1.2','공통 UI 컴포넌트'),('1.3','화면 목록 총괄표'),
  ('2','공통 화면'),
  ('2.1','SCR-COM-001  로그인 화면'),
  ('2.2','SCR-COM-002  메인 대시보드'),
  ('2.3','SCR-COM-003  네비게이션 바'),
  ('3','기준정보관리'),
  ('3.1','SCR-BM-001  제품품목기준'),('3.2','SCR-BM-002  레시피 BOM'),
  ('3.3','SCR-BM-003  공정 CCP'),('3.4','SCR-BM-004  설비/탱크'),
  ('3.5','SCR-BM-005  거래처'),('3.6','SCR-BM-006  작업자'),
  ('4','수주생산계획관리'),
  ('4.1','SCR-PM-001  수주 등록/조회'),('4.2','SCR-PM-002  납기 캘린더'),
  ('4.3','SCR-PM-003  생산계획 수립'),('4.4','SCR-PM-004  작업지시서'),
  ('4.5','SCR-PM-005  작업지시서 POP'),
  ('5','자재재고관리'),
  ('5.1','SCR-INV-001  원부자재 입고'),('5.2','SCR-INV-002  재고 현황'),
  ('5.3','SCR-INV-003  소요량 계산'),
  ('6','입고전처리관리'),
  ('6.1','SCR-RCV-001  원물 입고 검사'),('6.2','SCR-RCV-002  전처리 중량'),
  ('7','세척공정관리'),
  ('7.1','SCR-WSH-001  세척 조건'),('7.2','SCR-WSH-002  소독수 농도'),
  ('7.3','SCR-WSH-003  세척 실적 POP'),
  ('8','절임공정관리'),
  ('8.1','SCR-PCK-001  절임 조건'),('8.2','SCR-PCK-002  절임통 운영'),
  ('9','양념버무림공정관리'),
  ('9.1','SCR-SZN-001  양념 계량'),('9.2','SCR-SZN-002  버무림 CCP'),
  ('10','포장출하관리'),
  ('10.1','SCR-PKG-001  중량 검사'),('10.2','SCR-PKG-002  출하 현황'),
  ('11','숙성냉장관리'),
  ('11.1','SCR-AGE-001  숙성 재고'),('11.2','SCR-AGE-002  온도·습도 모니터링'),
  ('12','품질이상관리'),
  ('12.1','SCR-QC-001  공정별 불량'),('12.2','SCR-QC-002  금속검출기'),
  ('12.3','SCR-QC-003  이슈 이력'),
]
for num,title in toc:
    lvl = num.count('.')
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(lvl*0.8)
    r1 = p.add_run(f'{num}  '); r1.bold=(lvl==0); r1.font.size=Pt(10 if lvl==0 else 9)
    r2 = p.add_run(title);      r2.bold=(lvl==0); r2.font.size=Pt(10 if lvl==0 else 9)

doc.add_page_break()

# ═══════════════ 1. 개요 ════════════════════════════════════
add_heading(doc, '1. 개요', 1)

add_heading(doc, '1.1 화면 설계 원칙', 2)
for t,d in [
  ('현장 적합성','장갑 착용 상태에서도 조작 가능한 최소 48px 터치 타겟 적용'),
  ('고대비 색상','현장 밝은 조명 환경에서도 가독성 보장하는 고대비 색상'),
  ('반응형 레이아웃','관리자 웹(1920×1080), POP(1280×800), 패드(1024×768) 최적화'),
  ('직관적 UX','현장 작업자 IT 리터러시 고려, 단순하고 직관적 인터페이스'),
  ('실시간성','재고·공정·온도 데이터 Auto Refresh(5초~30초) 적용'),
  ('오류 방지','입력 유효성 검사, 확인 팝업, 범위 초과 경보로 실수 방지'),
]:
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f'▶ {t}: '); r1.bold=True; r1.font.size=Pt(10)
    r2 = p.add_run(d);          r2.font.size=Pt(10)

doc.add_paragraph()
add_heading(doc, '1.2 공통 UI 컴포넌트', 2)
add_table(doc,
    ['컴포넌트 ID','컴포넌트명','스펙','적용 화면'],
    [
      ('BTN-PRIMARY','주요 버튼','H:40px, #1A73E8, 흰 텍스트, radius 8px','등록/저장/조회'),
      ('BTN-DANGER','위험 버튼','H:40px, #EA4335, 흰 텍스트','삭제/취소'),
      ('BTN-SECONDARY','보조 버튼','H:40px, 흰 배경, 파란 보더','초기화/닫기'),
      ('BTN-POP','POP 대형버튼','H:80px, 최소 W:160px, 폰트 18px','POP/패드'),
      ('TBL-GRID','데이터 그리드','헤더 #1A73E8, 호버 연파랑, 페이지네이션','목록 화면'),
      ('FORM-INPUT','텍스트 입력','H:36px, 보더 1px, 포커스 파란 보더','검색/입력'),
      ('FORM-SELECT','드롭다운','H:36px, 화살표 아이콘','코드 선택'),
      ('FORM-DATE','날짜 선택','달력 팝업, 기간 선택 지원','날짜 조회'),
      ('POPUP-ALERT','경보 팝업','빨간 테두리, 경보음 연동, 확인 버튼','CCP 이상'),
      ('POPUP-CONFIRM','확인 팝업','2-step 확인, 예/아니오','삭제 확인'),
      ('BADGE-STATUS','상태 뱃지','진행중(파랑)/완료(초록)/대기(회색)','공정 상태'),
      ('CARD-KPI','KPI 카드','수치 24px, 단위 12px, 흰 카드','대시보드'),
      ('CHART-LINE','꺾은선 차트','Chart.js, 실시간 갱신','온도/추이'),
      ('CHART-BAR','막대 차트','Chart.js, 색상 범례','생산량'),
      ('PAGINATION','페이지네이션','이전/다음/번호, 기본 20건','목록 화면'),
    ]
)

doc.add_paragraph()
add_heading(doc, '1.3 화면 목록 총괄표', 2)
add_table(doc,
    ['화면 ID','화면명','화면 유형','접근 권한'],
    [
      ('SCR-COM-001','로그인 화면','웹/POP/패드','전체'),
      ('SCR-COM-002','메인 대시보드','관리자 웹','관리자, 생산관리자'),
      ('SCR-COM-003','네비게이션 바','관리자 웹','권한별'),
      ('SCR-BM-001','제품품목기준 목록/등록','관리자 웹','관리자'),
      ('SCR-BM-002','레시피 BOM 관리','관리자 웹','관리자, 품질관리자'),
      ('SCR-BM-003','공정 CCP 기준 관리','관리자 웹','관리자, 품질관리자'),
      ('SCR-BM-004','설비/탱크 기준 관리','관리자 웹','관리자'),
      ('SCR-BM-005','거래처 관리','관리자 웹','관리자'),
      ('SCR-BM-006','작업자 관리','관리자 웹','관리자'),
      ('SCR-PM-001','수주 등록/조회','관리자 웹','영업, 관리자'),
      ('SCR-PM-002','납기 캘린더','관리자 웹','영업, 생산관리자'),
      ('SCR-PM-003','생산계획 수립','관리자 웹','생산관리자'),
      ('SCR-PM-004','작업지시서 관리','관리자 웹','생산관리자'),
      ('SCR-PM-005','작업지시서 POP 표시','POP 현장 단말','현장작업자'),
      ('SCR-INV-001','원부자재 입고 등록','스마트패드','자재담당자'),
      ('SCR-INV-002','실시간 재고 현황','관리자 웹','자재, 관리자'),
      ('SCR-INV-003','소요량 계산','관리자 웹','생산관리자'),
      ('SCR-RCV-001','원물 입고 검사','관리자 웹/패드','품질관리자'),
      ('SCR-RCV-002','전처리 중량 관리','POP 현장 단말','현장작업자'),
      ('SCR-WSH-001','세척 조건 관리','관리자 웹','관리자, 품질관리자'),
      ('SCR-WSH-002','소독수 농도 관리','관리자 웹','품질관리자'),
      ('SCR-WSH-003','세척 실적 입력','POP 현장 단말','현장작업자'),
      ('SCR-PCK-001','절임 조건 설정','관리자 웹','생산관리자'),
      ('SCR-PCK-002','절임통 운영 현황','POP 현장 단말','현장작업자, 관리자'),
      ('SCR-SZN-001','양념 계량 관리','POP/스마트패드','현장작업자'),
      ('SCR-SZN-002','버무림 CCP 관리','POP 현장 단말','품질관리자, 작업자'),
      ('SCR-PKG-001','중량 검사 관리','POP 현장 단말','품질관리자, 작업자'),
      ('SCR-PKG-002','출하 현황','관리자 웹','영업, 관리자'),
      ('SCR-AGE-001','숙성 재고 관리','관리자 웹','생산관리자'),
      ('SCR-AGE-002','냉장고 온도·습도 모니터링','관리자 웹','품질관리자, 관리자'),
      ('SCR-QC-001','공정별 불량 관리','관리자 웹','품질관리자'),
      ('SCR-QC-002','금속검출기 관리','POP 현장 단말','품질관리자, 작업자'),
      ('SCR-QC-003','이슈 이력 관리','관리자 웹','품질관리자, 관리자'),
    ]
)
doc.add_page_break()

# ═══════════════ 2. 공통 화면 ════════════════════════════════
add_heading(doc, '2. 공통 화면', 1)

# SCR-COM-001 로그인
add_heading(doc, '2.1 SCR-COM-001  로그인 화면', 2)
screen_block(doc,
  'SCR-COM-001','로그인 화면','관리자 웹 / POP / 스마트패드','전체 사용자',
  '사용자 인증 화면. 아이디/비밀번호 입력 후 역할(Role)에 따라 메인 화면으로 이동. '
  'POP 단말은 QR코드/바코드 스캔 로그인도 지원.',
  """\
+================================================================+
|              임진강김치 MES  -  로그인                          |
|                                                                |
|              +------------------------------------+            |
|              |        [회사 로고 / 임진강김치]      |            |
|              |  MES 생산관리시스템  v2.0            |            |
|              |                                    |            |
|              |  아이디  [____________________]    |            |
|              |  비밀번호 [____________________]   |            |
|              |                                    |            |
|              |  화면유형: (●)관리자웹 ( )POP ( )패드|            |
|              |                                    |            |
|              |  [      로   그   인      ]         |            |
|              |                                    |            |
|              |  [QR코드 스캔 로그인]  [비밀번호 찾기]|            |
|              +------------------------------------+            |
|                                                                |
|  © 2026 임진강김치. 무단 접근 금지.    접속 IP: 192.168.1.xxx   |
+================================================================+
  (관리자 웹: 1920x1080 / POP: 1280x800 / 패드: 1024x768 공통 사용)""",
  [
    ('로고 이미지','이미지','임진강김치 브랜드 로고 표시'),
    ('아이디 입력','FORM-INPUT','사용자 ID 입력, 최대 20자'),
    ('비밀번호 입력','FORM-INPUT (password)','마스킹 처리, 최대 20자'),
    ('화면유형 선택','라디오 버튼','관리자웹/POP/스마트패드 선택'),
    ('로그인 버튼','BTN-PRIMARY','인증 처리 및 화면 이동'),
    ('QR 스캔 버튼','BTN-SECONDARY','QR코드/바코드 스캔 인증'),
    ('비밀번호 찾기','링크 버튼','비밀번호 재설정 이메일 발송'),
  ],
  [
    ('로그인 버튼 클릭','onClick','ID/PW 서버 인증 → 성공시 역할별 메인화면 이동, 실패시 오류 메시지'),
    ('QR 스캔','onScan','QR/바코드 인식 → 자동 로그인'),
    ('5회 실패','onFailure x5','계정 잠금 처리 + 관리자 알림'),
    ('Enter 키','onKeyDown','로그인 버튼 동작'),
  ],
  [
    ('아이디','필수 입력, 영문/숫자 조합','아이디를 입력해 주세요'),
    ('비밀번호','필수 입력, 8자 이상','비밀번호를 입력해 주세요'),
    ('비밀번호','특수문자 1개 이상 포함','비밀번호 형식이 올바르지 않습니다'),
  ]
)

# SCR-COM-002 메인 대시보드
add_heading(doc, '2.2 SCR-COM-002  메인 대시보드', 2)
screen_block(doc,
  'SCR-COM-002','메인 대시보드','관리자 웹 (1920x1080)','관리자, 생산관리자',
  '공장 전체 생산 현황을 한눈에 파악하는 메인 화면. KPI 카드, 공정별 진행 현황, '
  '재고 알림, 온도 이상 경보를 실시간으로 표시. 30초 자동 갱신.',
  """\
+============================================================================================================================+
| [임진강김치 MES]           [알람 3건 ▲]  [생산관리자: 홍길동]  [설정]  [로그아웃]                  2026-05-12  09:30:15    |
+--------------------+-------------------------------------------------------------------------------------------------------+
| [기준정보관리]     |  ■ 오늘의 생산 현황 대시보드                          [ 자동갱신: 30초 ] [새로고침]                  |
| [수주/생산계획]    +------------------------+------------------------+------------------------+------------------------+   |
| [자재/재고]        | 오늘 생산계획 (kg)     | 생산 실적 (kg)         | 달성률                  | 불량률                 |   |
| [입고/전처리]      |   12,500               |   8,320                |   66.6%                |   0.8%                 |   |
| [세척공정]         +------------------------+------------------------+------------------------+------------------------+   |
| [절임공정]         |                                                                                                       |   |
| [양념공정]         |  ■ 공정별 진행 현황                                                                                   |   |
| [포장/출하]        |  ┌──────────┬──────────┬──────────┬──────────┬──────────┬──────────┬──────────┐                       |   |
| [숙성/냉장]        |  │  입고     │  전처리   │  세척     │  절임     │  버무림   │  포장     │  출하     │                       |   |
| [품질이상]         |  │ [완료]    │ [진행중]  │ [대기]   │ [진행중]  │ [완료]   │ [진행중]  │ [대기]   │                       |   |
|                    |  │ 2,500kg  │ 2,100kg  │ -        │ 8조 운용  │ 1,800kg  │ 450박스  │ -        │                       |   |
| [시스템관리]       |  └──────────┴──────────┴──────────┴──────────┴──────────┴──────────┴──────────┘                       |   |
|                    |                                                                                                       |   |
|                    |  ■ 재고 현황 (주요 원자재)          ■ 냉장창고 온도 모니터링                                          |   |
|                    |  ┌────────────┬──────┬──────┐      ┌──────────────────────────────────────┐                          |   |
|                    |  │ 원자재명   │재고량│상태  │      │ A동 냉장1: -2.1℃ [정상]              │                          |   |
|                    |  ├────────────┼──────┼──────┤      │ A동 냉장2: -1.8℃ [정상]              │                          |   |
|                    |  │ 배추       │5,200 │[정상]│      │ B동 냉장1:  4.2℃ [경고!]▲           │                          |   |
|                    |  │ 고춧가루   │  320 │[부족]│      │ B동 냉장2: -0.5℃ [정상]              │                          |   |
|                    |  │ 멸치액젓   │  180 │[정상]│      └──────────────────────────────────────┘                          |   |
|                    |  └────────────┴──────┴──────┘                                                                        |   |
|                    |                                                                                                       |   |
|                    |  ■ 최근 이슈 / 알람                                                                                   |   |
|                    |  [09:15] ⚠ B동 냉장1 온도 초과 (기준: 0℃, 현재: 4.2℃) - 확인 필요                                   |   |
|                    |  [08:30] ℹ 작업지시서 WO-2026-0512-003 발행 완료                                                     |   |
+--------------------+-------------------------------------------------------------------------------------------------------+""",
  [
    ('KPI 카드 (4종)','CARD-KPI','오늘 생산계획/실적/달성률/불량률 표시'),
    ('공정별 현황 표','상태 배지 테이블','7개 공정 진행 상태를 색상 뱃지로 표시'),
    ('재고 현황 그리드','TBL-GRID','주요 원자재 재고량 및 부족 경고'),
    ('온도 모니터링 패널','실시간 패널','냉장고별 온도 및 이상 상태'),
    ('알람 목록','POPUP-ALERT 연동','최근 이슈 타임라인 표시'),
    ('자동 갱신 타이머','JS setInterval','30초마다 데이터 자동 갱신'),
    ('사이드 메뉴','BREADCRUMB 연동','전체 기능 메뉴 트리'),
  ],
  [
    ('알람 클릭','onClick','해당 이슈 상세 화면으로 이동'),
    ('재고 행 클릭','onClick','재고 상세/이동 화면으로 이동'),
    ('온도 경고 클릭','onClick','냉장고 상세 모니터링 화면 이동'),
    ('공정 배지 클릭','onClick','해당 공정 POP/관리 화면 이동'),
    ('30초 타이머','setInterval','REST API 호출 → 전체 지표 갱신'),
  ],
  [
    ('온도 기준','설정 기준값 초과 시','⚠ 경고 뱃지 표시 및 알람 발송'),
    ('재고량','안전재고 이하 시','[부족] 뱃지 표시'),
    ('달성률','0~100% 범위','100% 초과 시 이상값 표시'),
  ]
)

# SCR-COM-003 네비게이션
add_heading(doc, '2.3 SCR-COM-003  상단 메뉴 / 네비게이션 바', 2)
screen_block(doc,
  'SCR-COM-003','상단 메뉴/네비게이션 바','관리자 웹 (공통 컴포넌트)','권한별 메뉴 노출',
  '모든 관리자 웹 화면 상단에 고정 배치되는 공통 네비게이션. 좌측 사이드바 메뉴와 상단 헤더로 구성. '
  '사용자 권한에 따라 메뉴 노출 여부 제어.',
  """\
+====================================================================================================================+
| [≡]  임진강김치 MES                [검색: 화면/기능 검색...]          [🔔 알람 3]  [홍길동▼]  [로그아웃]            |
+=====================================================================================================================+
| 사이드바 (200px)        | 콘텐츠 영역 (나머지)                                                                      |
|-------------------------|                                                                                           |
| 📊 대시보드             | 브레드크럼: 홈 > 기준정보관리 > 제품품목기준                                                |
|                         |                                                                                           |
| ▼ 기준정보관리          |  < 현재 화면 내용 >                                                                        |
|   - 제품품목기준         |                                                                                           |
|   - 레시피 BOM           |                                                                                           |
|   - 공정 CCP             |                                                                                           |
|   - 설비/탱크            |                                                                                           |
|   - 거래처               |                                                                                           |
|   - 작업자               |                                                                                           |
|                         |                                                                                           |
| ▶ 수주/생산계획          |                                                                                           |
| ▶ 자재/재고              |                                                                                           |
| ▶ 입고/전처리            |                                                                                           |
| ▶ 세척공정               |                                                                                           |
| ▶ 절임공정               |                                                                                           |
| ▶ 양념공정               |                                                                                           |
| ▶ 포장/출하              |                                                                                           |
| ▶ 숙성/냉장              |                                                                                           |
| ▶ 품질이상               |                                                                                           |
| ─────────────────       |                                                                                           |
| ⚙ 시스템관리             |                                                                                           |
+-------------------------+-------------------------------------------------------------------------------------------+""",
  [
    ('햄버거 메뉴 버튼','아이콘 버튼','사이드바 접기/펼치기'),
    ('글로벌 검색','FORM-INPUT','화면명/기능 검색 → 자동완성'),
    ('알람 뱃지','BADGE-STATUS','미확인 알람 수 표시, 클릭시 알람 목록'),
    ('사용자 메뉴','드롭다운','프로필, 비밀번호 변경, 로그아웃'),
    ('사이드바 메뉴','트리 메뉴','대메뉴/소메뉴, 권한별 노출 제어'),
    ('브레드크럼','BREADCRUMB','현재 위치 경로 표시'),
    ('현재 메뉴 하이라이트','CSS active','활성 메뉴 파란색 강조'),
  ],
  [
    ('메뉴 항목 클릭','onClick','해당 화면으로 라우팅'),
    ('대메뉴 클릭','onClick','하위 메뉴 펼치기/접기'),
    ('알람 아이콘 클릭','onClick','알람 드롭다운 표시'),
    ('검색 입력','onKeyUp','자동완성 드롭다운 표시'),
  ],
  [
    ('권한','미허용 메뉴','그레이아웃 처리 또는 비표시'),
    ('알람 수','0~999건','999 초과시 "999+" 표시'),
  ]
)

doc.add_page_break()

doc.save(OUTPUT)
print(f'Part1 저장 완료: {OUTPUT}')
