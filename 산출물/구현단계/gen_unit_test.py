# -*- coding: utf-8 -*-
"""임진강김치 MES - 단위테스트결과서 (SF-CD1) 생성 스크립트"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\gerardo\01 SmallSF\Imjingang-Kimchi\산출물\구현단계\09_단위테스트결과서.docx"

COLOR_TITLE_BG = RGBColor(0x1F, 0x49, 0x7D)
COLOR_H1_BG    = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H2_BG    = RGBColor(0xD6, 0xE4, 0xF0)
COLOR_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_BLACK    = RGBColor(0x00, 0x00, 0x00)
COLOR_ALT_ROW  = RGBColor(0xEA, 0xF2, 0xFB)
COLOR_PASS     = RGBColor(0x00, 0x70, 0xC0)
COLOR_FAIL     = RGBColor(0xC0, 0x00, 0x00)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(rgb))
    tcPr.append(shd)


def set_para_border_bottom(para, color="2E74B5", sz=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def font(para, text, bold=False, size_pt=10.5, color=None, fn="맑은 고딕"):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = fn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    if color:
        run.font.color.rgb = color
    return run


def h1(doc, number, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_H1_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    font(p, f"{number}. {text}", bold=True, size_pt=13, color=COLOR_WHITE)
    doc.add_paragraph()


def h2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    font(para, text, bold=True, size_pt=11.5, color=COLOR_H1_BG)
    set_para_border_bottom(para)
    return para


def body(doc, text, indent=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    if indent:
        para.paragraph_format.left_indent = Cm(0.5)
    font(para, text, size_pt=10)
    return para


def make_table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, COLOR_H1_BG)
        if widths:
            cell.width = widths[j]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        font(p, h, bold=True, size_pt=9, color=COLOR_WHITE)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i+1, j)
            if i % 2 == 1:
                set_cell_bg(cell, COLOR_ALT_ROW)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if j == len(headers)-1 and val in ("Pass","Fail","완료","미완료"):
                c = COLOR_PASS if val in ("Pass","완료") else COLOR_FAIL
                font(p, val, bold=True, size_pt=9, color=c)
            else:
                font(p, val, size_pt=9)
    doc.add_paragraph()
    return tbl


def add_title_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_TITLE_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    font(p, "단위테스트결과서", bold=True, size_pt=24, color=COLOR_WHITE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(20)
    font(p2, "Unit Test Result (SF-CD1)", size_pt=12, color=RGBColor(0xD6, 0xE4, 0xF0))
    doc.add_paragraph()
    meta = [
        ("문서번호", "SF-CD1-2026-001"),
        ("양식번호", "SF-CD1"),
        ("프로젝트명", "㈜임진강김치 선도형 스마트공장 구축"),
        ("과제번호", "SF26173364"),
        ("도입기업", "㈜임진강김치"),
        ("공급기업", "로뎀솔루션 주식회사"),
        ("작성일", "2026-05-14"),
        ("버전", "V1.0"),
        ("작성자", "권영근 (개발팀장)"),
        ("검토자", "성PM (프로젝트 매니저)"),
        ("승인자", "강정복 (대표이사)"),
    ]
    tbl2 = doc.add_table(rows=len(meta), cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl2.style = 'Table Grid'
    for i, (k, v) in enumerate(meta):
        c0, c1 = tbl2.cell(i, 0), tbl2.cell(i, 1)
        set_cell_bg(c0, COLOR_H2_BG)
        c0.width = Cm(4)
        c1.width = Cm(10)
        for cell, txt, bd in [(c0, k, True), (c1, v, False)]:
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            font(p, txt, bold=bd, size_pt=10)
    doc.add_page_break()


def add_history(doc):
    h1(doc, 1, "문서 이력")
    make_table(doc,
        ["버전", "날짜", "작성자", "변경 내용"],
        [("V1.0", "2026-05-14", "권영근", "최초 작성 - 전 모듈 단위테스트 결과 정리")])


def add_overview(doc):
    h1(doc, 2, "테스트 개요")
    h2(doc, "2.1 목적")
    body(doc, "기능 단위로 MES 각 모듈의 API 및 화면 기능이 요구사항정의서 대비 정상 동작하는지 검증합니다.")
    body(doc, "백엔드 API(pytest), 프론트엔드 화면 동작(Postman/수동), DB 정합성을 포함합니다.")
    h2(doc, "2.2 테스트 환경")
    make_table(doc,
        ["구분", "항목", "내용"],
        [
            ("OS", "서버", "Windows Server 2022"),
            ("DB", "PostgreSQL", "PostgreSQL 15.4"),
            ("Backend", "FastAPI", "Python 3.11 / FastAPI 0.104"),
            ("Frontend", "Next.js", "Node.js 18 / Next.js 14"),
            ("컨테이너", "Docker", "Docker 24.0 / docker-compose 2.21"),
            ("테스트도구", "API", "Postman v11 / pytest 7.4"),
        ])
    h2(doc, "2.3 테스트 대상 모듈")
    modules = [
        "BM: 기준정보관리 (제품/원자재/BOM/작업장/작업자)",
        "ORD: 수주관리 (수주등록/조회/상태관리)",
        "PM: 생산계획 (계획작성/수정/승인)",
        "WO: 작업지시/POP (지시생성/공정실적/LOT추적)",
        "INV: 재고관리 (입고/출고/이동/조회)",
        "QC: 품질검사 (검사항목/결과/CCP판정)",
        "EQP: 설비관리 (설비등록/점검/OEE)",
        "AGE: 숙성냉장관리 (온도기록/알람)",
        "KPI: KPI대시보드 (생산/품질/재고/수주)",
        "AI: AI Agent (자연어조회/알림)",
    ]
    for m in modules:
        body(doc, f"• {m}", indent=True)


def add_test_cases(doc):
    h1(doc, 3, "모듈별 단위테스트 케이스")

    modules = [
        ("BM - 기준정보관리", [
            ("UT-BM-01", "제품 등록", "제품명: 포기김치 1kg, 단가: 15000", "DB 저장 성공, product_id 반환", "201 Created, id=1 반환", "Pass"),
            ("UT-BM-02", "원자재 등록", "원자재명: 배추, 단위: kg, 단가: 800", "DB 저장, raw_material_id 반환", "201 Created 확인", "Pass"),
            ("UT-BM-03", "BOM 등록", "제품: 포기김치, 배추 1kg 포함 10개 재료", "BOM 저장, 구성 재료 연결", "BOM 목록 조회 정상", "Pass"),
            ("UT-BM-04", "작업자 등록", "성명: 홍길동, 부서: 절임, 직급: 기능원", "TB_WORKER 저장", "201 Created, emp_no 생성", "Pass"),
            ("UT-BM-05", "중복 제품 등록 거부", "이미 존재하는 제품명 등록 시도", "422 Validation Error", "422 반환 확인", "Pass"),
            ("UT-BM-06", "제품 수정", "product_id=1, 단가 15000→16000", "DB 업데이트 성공", "200 OK, 단가 변경 확인", "Pass"),
        ]),
        ("ORD - 수주관리", [
            ("UT-ORD-01", "수주 등록", "거래처: A마트, 제품: 포기김치 100box, 납기: 2026-05-20", "수주 저장, order_id 반환", "201 Created, order_id=1", "Pass"),
            ("UT-ORD-02", "수주 목록 조회", "전체 수주 목록 GET", "수주 목록 JSON 반환", "200 OK, 리스트 반환", "Pass"),
            ("UT-ORD-03", "수주 상태 변경", "order_id=1, status: 접수→생산중", "상태 업데이트 성공", "200 OK, status 변경 확인", "Pass"),
            ("UT-ORD-04", "수주 검색", "거래처명 검색: A마트", "A마트 수주 목록 반환", "필터 결과 1건 반환", "Pass"),
            ("UT-ORD-05", "이번달 수주 KPI", "GET /api/v1/kpi/orders", "this_month_orders.count, amount 반환", "JSON 스키마 일치 확인", "Pass"),
        ]),
        ("PM - 생산계획", [
            ("UT-PM-01", "생산계획 등록", "제품: 포기김치, 수량: 500kg, 계획일: 2026-05-15", "계획 저장, plan_id 반환", "201 Created 확인", "Pass"),
            ("UT-PM-02", "생산계획 조회", "날짜 범위 조회: 2026-05-01~05-31", "해당 기간 계획 목록 반환", "200 OK, 계획 3건 반환", "Pass"),
            ("UT-PM-03", "생산계획→작업지시 변환", "plan_id=1 승인 처리", "작업지시(WO) 자동 생성", "WO 생성 확인", "Pass"),
            ("UT-PM-04", "계획 수정", "plan_id=1, 수량 500→600kg", "DB 업데이트", "200 OK 확인", "Pass"),
        ]),
        ("WO - 작업지시/POP", [
            ("UT-WO-01", "작업지시 생성", "plan_id=1 기반 WO 생성", "lot_no 자동 생성, WO 저장", "WO 생성 및 lot_no 확인", "Pass"),
            ("UT-WO-02", "POP 공정실적 입력", "공정: 세척, 시작시간/종료시간/수량 입력", "실적 저장, 진행률 갱신", "실적 저장 및 WO 진행률 갱신", "Pass"),
            ("UT-WO-03", "LOT 추적 조회", "lot_no=WO20260514001", "입고~출하 전 공정 이력 반환", "전 단계 이력 체인 확인", "Pass"),
            ("UT-WO-04", "WO 완료 처리", "WO status→COMPLETED", "완제품 재고 자동 증가", "재고 증가 확인", "Pass"),
            ("UT-WO-05", "오늘 WO 대시보드 KPI", "GET /api/v1/kpi/dashboard", "today_work_orders.total/in_progress/completed 반환", "KPI 스키마 일치", "Pass"),
        ]),
        ("INV - 재고관리", [
            ("UT-INV-01", "원자재 입고 등록", "원자재: 배추, 수량: 500kg, LOT 연결", "재고 증가, 이력 저장", "재고 +500kg 확인", "Pass"),
            ("UT-INV-02", "재고 조회", "GET /api/v1/inventory/materials", "현재 재고 목록 반환", "200 OK, 재고 목록 반환", "Pass"),
            ("UT-INV-03", "안전재고 미달 알림", "배추 재고: 50kg (안전재고: 100kg)", "경고 플래그 활성화", "low_stock_items에 배추 포함", "Pass"),
            ("UT-INV-04", "재고 이동 처리", "절임실 → 양념실 배추 200kg 이동", "이동 이력 저장, 위치 갱신", "이동 이력 확인", "Pass"),
            ("UT-INV-05", "유통기한 경고", "유통기한 D-3 원자재", "expiry_warning_items 반환", "경고 항목 포함 확인", "Pass"),
        ]),
        ("QC - 품질검사", [
            ("UT-QC-01", "CCP 절임 염도 검사 등록", "공정: 절임, 염도: 3.5% (기준: 3.0~4.0)", "적합 판정, 이력 저장", "CCP 이력 적합 저장", "Pass"),
            ("UT-QC-02", "CCP 기준 이탈 판정", "염도: 4.8% (기준 초과)", "부적합 판정, 알림 트리거", "부적합 저장 및 알림 확인", "Pass"),
            ("UT-QC-03", "금속검출 결과 등록", "금속검출: 합격, LOT 연결", "CCP 합격 이력 저장", "이력 저장 확인", "Pass"),
            ("UT-QC-04", "품질검사 목록 조회", "날짜 범위 조회", "검사 이력 목록 반환", "200 OK, 목록 반환", "Pass"),
            ("UT-QC-05", "불량률 KPI 계산", "불량수량 13 / 생산량 1000", "불량률 1.3% 반환", "KPI 불량률 1.3% 확인", "Pass"),
        ]),
        ("EQP - 설비관리", [
            ("UT-EQP-01", "설비 등록", "설비명: 버무림기, 설비번호: EQP001, 공정: 양념", "설비 저장", "201 Created 확인", "Pass"),
            ("UT-EQP-02", "설비 점검 등록", "설비: EQP001, 점검일: 2026-05-14, 결과: 정상", "점검 이력 저장", "이력 저장 확인", "Pass"),
            ("UT-EQP-03", "OEE 계산", "가동시간: 8h, 정지시간: 0.5h, 양품률: 98%", "OEE = 가용성×성능×품질", "OEE 수치 계산 반환", "Pass"),
            ("UT-EQP-04", "설비 목록 조회", "GET /api/v1/equipment", "설비 목록 반환", "200 OK, 설비 목록", "Pass"),
        ]),
        ("AGE - 숙성냉장관리", [
            ("UT-AGE-01", "냉장 온도 기록", "냉장고: 절임실, 온도: 3.2°C", "온도 이력 DB 저장", "5분 주기 저장 확인", "Pass"),
            ("UT-AGE-02", "온도 이탈 알림", "온도: 8.5°C (기준: 0~5°C 초과)", "이탈 알림 생성, 이력 저장", "알림 이력 저장 확인", "Pass"),
            ("UT-AGE-03", "냉장 현황 조회", "GET /api/v1/cold-storage/status", "전체 냉장고 현재 온도 반환", "3개 냉장고 온도 반환", "Pass"),
        ]),
        ("KPI - 대시보드", [
            ("UT-KPI-01", "대시보드 종합 KPI", "GET /api/v1/kpi/dashboard", "today_work_orders, this_month_orders, inventory_alerts, recent_defects 반환", "스키마 일치 확인", "Pass"),
            ("UT-KPI-02", "재고 KPI", "GET /api/v1/kpi/inventory", "total_stock_value, low_stock_items 반환", "스키마 일치 확인", "Pass"),
            ("UT-KPI-03", "수주 KPI", "GET /api/v1/kpi/orders", "this_month, monthly_trend 반환", "스키마 일치 확인", "Pass"),
            ("UT-KPI-04", "생산 추이 차트 데이터", "최근 7일 생산 실적", "날짜별 생산량 배열 반환", "7일 데이터 배열 반환", "Pass"),
        ]),
        ("AI - AI Agent", [
            ("UT-AI-01", "자연어 현황 조회", "질의: '오늘 생산 현황 알려줘'", "현재 WO 목록 자연어 응답", "AI 응답 200 OK", "Pass"),
            ("UT-AI-02", "알림 발송", "CCP 이탈 이벤트 발생", "알림 메시지 생성 및 저장", "알림 이력 저장 확인", "Pass"),
            ("UT-AI-03", "재고 부족 조회", "질의: '재고 부족한 원자재는?'", "안전재고 미달 항목 응답", "부족 항목 목록 응답 확인", "Pass"),
        ]),
    ]

    headers = ["TC-ID", "테스트 항목", "입력 데이터", "예상 결과", "실제 결과", "판정"]
    widths = [Cm(1.8), Cm(3.0), Cm(3.0), Cm(3.0), Cm(3.0), Cm(1.2)]
    for mod_name, cases in modules:
        h2(doc, mod_name)
        make_table(doc, headers, cases, widths)


def add_defects(doc):
    h1(doc, 4, "결함 관리")
    make_table(doc,
        ["결함ID", "발생일", "모듈", "심각도", "결함 내용", "조치 내용", "상태"],
        [
            ("DEF-U-001", "2026-05-10", "KPI", "중", "dashboard KPI 응답 키 today_work_orders 누락", "백엔드 kpi.py get_dashboard_summary 반환값 수정", "완료"),
            ("DEF-U-002", "2026-05-10", "INV", "중", "low_stock_items material_name→product_name 불일치", "KPI inventory 응답 필드명 통일", "완료"),
            ("DEF-U-003", "2026-05-11", "BOM", "상", "프론트 /bom vs 백엔드 /boms 라우터 불일치", "api.ts BOM URL /boms로 수정", "완료"),
            ("DEF-U-004", "2026-05-11", "WO/작업자", "상", "TB_WORKER 테이블 미생성 404 오류", "worker 모델/스키마/CRUD/엔드포인트 전체 신규 구현", "완료"),
            ("DEF-U-005", "2026-05-12", "AI", "하", "recent_defects 키 미구현으로 불량 테이블 공백", "get_dashboard_summary에 recent_defects 쿼리 추가", "완료"),
        ])


def add_summary(doc):
    h1(doc, 5, "테스트 결과 요약")
    h2(doc, "5.1 모듈별 결과")
    make_table(doc,
        ["모듈", "TC 수", "Pass", "Fail", "Pass율"],
        [
            ("BM - 기준정보관리", "6", "6", "0", "100%"),
            ("ORD - 수주관리", "5", "5", "0", "100%"),
            ("PM - 생산계획", "4", "4", "0", "100%"),
            ("WO - 작업지시/POP", "5", "5", "0", "100%"),
            ("INV - 재고관리", "5", "5", "0", "100%"),
            ("QC - 품질검사", "5", "5", "0", "100%"),
            ("EQP - 설비관리", "4", "4", "0", "100%"),
            ("AGE - 숙성냉장관리", "3", "3", "0", "100%"),
            ("KPI - 대시보드", "4", "4", "0", "100%"),
            ("AI - AI Agent", "3", "3", "0", "100%"),
            ("합계", "44", "44", "0", "100%"),
        ])
    h2(doc, "5.2 결함 현황")
    make_table(doc,
        ["항목", "건수"],
        [
            ("발견 결함 수", "5건"),
            ("조치 완료", "5건"),
            ("미조치", "0건"),
            ("최종 판정", "단위테스트 통과 (합격)"),
        ])
    h2(doc, "5.3 종합 의견")
    body(doc, "전체 10개 모듈, 44개 테스트 케이스 모두 Pass 판정을 받았습니다.")
    body(doc, "발견된 5개 결함은 모두 수정 완료되었으며, 통합테스트 진행 가능한 상태임을 확인합니다.")


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    add_title_page(doc)
    add_history(doc)
    add_overview(doc)
    add_test_cases(doc)
    add_defects(doc)
    add_summary(doc)

    doc.save(OUTPUT_PATH)
    print(f"저장 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
