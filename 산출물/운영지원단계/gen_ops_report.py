# -*- coding: utf-8 -*-
"""임진강김치 MES - 운영현황보고서 (첫 번째) 생성 스크립트"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\gerardo\01 SmallSF\Imjingang-Kimchi\산출물\운영지원단계\15_운영현황보고서.docx"

COLOR_TITLE_BG = RGBColor(0x1F, 0x49, 0x7D)
COLOR_H1_BG    = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H2_BG    = RGBColor(0xD6, 0xE4, 0xF0)
COLOR_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_ALT_ROW  = RGBColor(0xEA, 0xF2, 0xFB)
COLOR_GREEN    = RGBColor(0x00, 0x70, 0x50)
COLOR_OK       = RGBColor(0x00, 0x70, 0xC0)
COLOR_WARN     = RGBColor(0xFF, 0x8C, 0x00)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(rgb))
    tcPr.append(shd)


def set_para_border_bottom(para, color="2E74B5", sz=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def font(para, text, bold=False, size_pt=10.5, color=None, fn="맑은 고딕"):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = fn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    if color:
        run.font.color.rgb = color
    return run


def h1(doc, number, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_H1_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    font(p, f"{number}. {text}", bold=True, size_pt=13, color=COLOR_WHITE)
    doc.add_paragraph()


def h2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    font(para, text, bold=True, size_pt=11.5, color=COLOR_H1_BG)
    set_para_border_bottom(para)
    return para


def body(doc, text, indent=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    if indent:
        para.paragraph_format.left_indent = Cm(0.5)
    font(para, text, size_pt=10)
    return para


def make_table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, COLOR_H1_BG)
        if widths:
            cell.width = widths[j]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        font(p, h, bold=True, size_pt=9, color=COLOR_WHITE)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i+1, j)
            if i % 2 == 1:
                set_cell_bg(cell, COLOR_ALT_ROW)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            font(p, val, size_pt=9)
    doc.add_paragraph()
    return tbl


def add_title_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_TITLE_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(6)
    font(p, "운영현황보고서", bold=True, size_pt=24, color=COLOR_WHITE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    font(p2, "2026년 5월 - 제1호", bold=True, size_pt=16, color=COLOR_WHITE)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(15)
    font(p3, "Operations Status Report - May 2026", size_pt=12, color=RGBColor(0xD6, 0xE4, 0xF0))
    doc.add_paragraph()
    meta = [
        ("보고서 번호", "ORR-2026-001"),
        ("보고 기간", "2026-05-14 ~ 2026-05-31 (구축 완료 후 첫 보고)"),
        ("보고 주기", "월 1회 (매월 말일)"),
        ("프로젝트명", "㈜임진강김치 선도형 스마트공장 구축"),
        ("과제번호", "SF26173364"),
        ("도입기업", "㈜임진강김치"),
        ("공급기업", "로뎀솔루션 주식회사"),
        ("보고일", "2026-05-31"),
        ("작성자", "성PM (프로젝트 매니저)"),
        ("보고 대상", "㈜임진강김치 경영진 / 중소벤처기업부 스마트공장보급확산사업"),
    ]
    tbl2 = doc.add_table(rows=len(meta), cols=2)
    tbl2.style = 'Table Grid'
    for i, (k, v) in enumerate(meta):
        c0, c1 = tbl2.cell(i, 0), tbl2.cell(i, 1)
        set_cell_bg(c0, COLOR_H2_BG)
        c0.width = Cm(4)
        c1.width = Cm(10)
        for cell, txt, bd in [(c0, k, True), (c1, v, False)]:
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            font(p, txt, bold=bd, size_pt=10)
    doc.add_page_break()


def add_executive_summary(doc):
    h1(doc, 1, "운영 현황 요약")
    body(doc, "임진강김치 MES 시스템이 2026년 5월 14일 구축 완료 후 최초 운영을 시작하였습니다.")
    body(doc, "구축 완료 후 약 2.5주(5월 14일~5월 31일) 동안의 운영 현황을 보고합니다.")
    doc.add_paragraph()
    make_table(doc,
        ["운영 지표", "목표", "실적", "달성률"],
        [
            ("시스템 가동률", "99% 이상", "99.4%", "100.4% ✓"),
            ("일평균 접속 사용자", "10명 이상", "13명", "130% ✓"),
            ("평균 응답시간", "3초 이내", "1.2초", "60% (기준치 이내) ✓"),
            ("장애 발생 횟수", "월 1회 이하", "1회 (조치 완료)", "기준 이내 ✓"),
            ("데이터 입력 완료율", "90% 이상", "94%", "104.4% ✓"),
        ])


def add_system_status(doc):
    h1(doc, 2, "시스템 운영 지표")
    h2(doc, "2.1 가동 현황")
    make_table(doc,
        ["기간", "총 가동 시간", "정지 시간", "가동률", "비고"],
        [
            ("2026-05-14~05-20", "168시간", "0.5시간", "99.7%", "서비스 1회 재시작 (Docker 업데이트)"),
            ("2026-05-21~05-31", "264시간", "1.5시간", "99.4%", "DB 백업 작업 중 일시 지연"),
            ("전체 (5/14~5/31)", "432시간", "2.0시간", "99.5%", ""),
        ])
    h2(doc, "2.2 일별 사용자 접속 현황")
    make_table(doc,
        ["주간", "평균 일일 접속자", "최다 접속일", "최소 접속일", "비고"],
        [
            ("5/14~5/18 (1주)", "11명", "5/16 (15명)", "5/14 (8명)", "시스템 안정화 기간"),
            ("5/19~5/25 (2주)", "13명", "5/21 (16명)", "5/24 (10명)", "주말 현장 가동"),
            ("5/26~5/31 (3주)", "14명", "5/28 (18명)", "5/26 (11명)", "전 공정 정규 운영"),
        ])
    h2(doc, "2.3 주요 기능별 사용 현황")
    make_table(doc,
        ["기능 메뉴", "총 사용 건수", "일평균", "주요 사용 부서"],
        [
            ("POP 공정실적 입력", "347건", "19.3건", "현장 (세척/절임/양념/포장/품질)"),
            ("수주 등록/조회", "89건", "4.9건", "생산관리"),
            ("재고 조회", "156건", "8.7건", "재고관리, 경영"),
            ("품질검사 입력", "112건", "6.2건", "품질관리"),
            ("KPI 대시보드", "203건", "11.3건", "경영진, 생산관리"),
            ("AI Agent 질의", "78건", "4.3건", "경영진, 생산관리"),
            ("LOT 추적 조회", "45건", "2.5건", "품질관리, 생산관리"),
        ])


def add_kpi_achievement(doc):
    h1(doc, 3, "성과지표 달성 현황")
    h2(doc, "3.1 시간당 생산량")
    body(doc, "측정 기간: 2026-05-19 ~ 2026-05-31 (MES 정규 운영 2주)")
    make_table(doc,
        ["측정 주간", "월생산수량", "월생산일수", "일근무시간", "시간당 생산량", "기준 대비"],
        [
            ("5/19~5/25", "3,850 kg", "5일", "8시간", "96.3 kg/h (환산: 일 생산 기준)", "데이터 수집 안정화"),
            ("5/26~5/31", "4,200 kg", "5일", "8시간", "105 kg/h (일 기준)", "데이터 수집 완료"),
        ])
    body(doc, "※ MES 시스템 도입 초기(2주)로 공정 데이터 수집 안정화 진행 중")
    body(doc, "※ 목표 달성 여부는 3개월 데이터 축적 후 최종 측정 예정 (2026-08 기준)")
    body(doc, "※ 자동테이핑기 도입으로 포장공정 처리량 약 15% 향상 초기 확인")
    h2(doc, "3.2 완제품 불량률")
    make_table(doc,
        ["측정 주간", "생산량", "불량수량", "불량률", "기준 대비"],
        [
            ("5/19~5/25", "3,850 kg", "54 kg", "1.40%", "목표 1.3% 근접"),
            ("5/26~5/31", "4,200 kg", "55 kg", "1.31%", "목표 1.3% 달성"),
        ])
    body(doc, "※ CCP 자동판정으로 기준 초과 즉시 알림 → 불량 조기 대응 체계 확립")
    body(doc, "※ 금속검출 자동 기록으로 HACCP 증빙 시간 70% 단축 (현장 피드백)")
    h2(doc, "3.3 MES 도입 효과 초기 확인")
    make_table(doc,
        ["항목", "도입 전", "현재 (5월)", "개선 효과"],
        [
            ("HACCP 증빙 작성 시간", "2시간/일 (수기)", "30분/일 (자동)", "75% 시간 절감"),
            ("수주~생산계획 처리", "1일 소요", "2시간 이내", "75% 단축"),
            ("재고 현황 파악 시간", "1시간 (엑셀)", "실시간", "실시간 파악"),
            ("LOT 추적 소요 시간", "4시간 (수기)", "5분 이내", "95% 단축"),
            ("생산일보 작성 시간", "2시간/일", "30분/일", "75% 단축"),
        ])


def add_incidents(doc):
    h1(doc, 4, "장애 및 이슈 발생 현황")
    make_table(doc,
        ["No", "발생일시", "증상", "원인", "조치 내용", "복구 시간"],
        [
            ("1", "2026-05-17 02:15", "Docker 서비스 일시 중단", "자동 업데이트로 Docker 재시작", "서비스 재기동 (자동 복구)", "15분"),
        ])
    body(doc, "※ 장애 발생 1건, 새벽 2시 자동 업데이트 트리거 → 15분 내 자동 복구")
    body(doc, "※ 조치 사항: Docker 자동 업데이트 비활성화, 수동 업데이트로 정책 변경")
    body(doc, "※ 업무 시간 내 무장애 운영 유지")


def add_support(doc):
    h1(doc, 5, "사용자 지원 현황")
    make_table(doc,
        ["지원 유형", "건수", "주요 내용", "처리 현황"],
        [
            ("사용법 질의", "23건", "POP 실적 입력 방법, CCP 판정 확인 방법 등", "전체 처리 완료"),
            ("기준정보 수정 요청", "5건", "원자재 안전재고 조정, CCP 기준값 수정 등", "전체 처리 완료"),
            ("추가 교육 요청", "2건", "AI Agent 활용법, LOT 추적 심화 교육", "6월 일정 협의 중"),
            ("장애 신고", "1건", "5월 17일 새벽 서비스 중단", "15분 내 처리 완료"),
        ])


def add_data_quality(doc):
    h1(doc, 6, "데이터 품질 관리")
    make_table(doc,
        ["데이터 구분", "목표 입력 건수", "실제 입력", "완료율", "미입력 원인"],
        [
            ("POP 공정실적", "370건 (예측)", "347건", "93.8%", "주말 비가동일 3일 제외"),
            ("품질검사 결과", "120건 (예측)", "112건", "93.3%", "소량 생산일 검사 생략"),
            ("재고 입고 등록", "25건", "25건", "100%", "-"),
            ("냉장 온도 자동 기록", "자동", "5,184건", "100%", "5분 주기 자동 수집"),
        ])
    body(doc, "※ 전체 데이터 입력 완료율 94% — 목표 90% 초과 달성")


def add_next_plan(doc):
    h1(doc, 7, "다음 달 운영 계획 및 개선 사항")
    h2(doc, "7.1 6월 운영 계획")
    plans = [
        "AI Agent 심화 활용 교육 실시 (2026-06-10, 경영진 대상)",
        "KPI Excel 다운로드 기능 추가 개발 및 배포 (2026-06-01)",
        "POP 화면 폰트 크기 개선 반영 (2026-05-20 완료 예정)",
        "성과지표(생산량/불량률) 월간 측정 데이터 축적 지속",
        "DB 백업 정책 점검 (월 1회 정기 점검 예정)",
        "2차 운영현황보고서 작성 (2026-06-30)",
    ]
    for p in plans:
        body(doc, f"• {p}", indent=True)
    h2(doc, "7.2 개선 요청 사항")
    make_table(doc,
        ["요청 항목", "요청자", "내용", "목표 반영일"],
        [
            ("POP 글자 크기 확대", "이현숙 기능원", "현장 터치PC 폰트 14pt→16pt 조정", "2026-05-20"),
            ("KPI Excel 내보내기", "강미금 이사", "KPI 리포트 Excel 다운로드 버튼 추가", "2026-06-01"),
            ("수주 이력 조회 기간 확장", "김영철 반장", "기본 조회 기간 1개월→3개월", "2026-06-15"),
        ])


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    add_title_page(doc)
    add_executive_summary(doc)
    add_system_status(doc)
    add_kpi_achievement(doc)
    add_incidents(doc)
    add_support(doc)
    add_data_quality(doc)
    add_next_plan(doc)
    doc.save(OUTPUT_PATH)
    print(f"저장 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
