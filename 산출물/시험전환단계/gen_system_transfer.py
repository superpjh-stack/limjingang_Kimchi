# -*- coding: utf-8 -*-
"""임진강김치 MES - 시스템전환결과서 (SF-TI2) 생성 스크립트"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\gerardo\01 SmallSF\Imjingang-Kimchi\산출물\시험전환단계\11_시스템전환결과서.docx"

COLOR_TITLE_BG = RGBColor(0x1F, 0x49, 0x7D)
COLOR_H1_BG    = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H2_BG    = RGBColor(0xD6, 0xE4, 0xF0)
COLOR_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_ALT_ROW  = RGBColor(0xEA, 0xF2, 0xFB)
COLOR_OK       = RGBColor(0x00, 0x70, 0xC0)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(rgb))
    tcPr.append(shd)


def set_para_border_bottom(para, color="2E74B5", sz=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def font(para, text, bold=False, size_pt=10.5, color=None, fn="맑은 고딕"):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = fn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    if color:
        run.font.color.rgb = color
    return run


def h1(doc, number, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_H1_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    font(p, f"{number}. {text}", bold=True, size_pt=13, color=COLOR_WHITE)
    doc.add_paragraph()


def h2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    font(para, text, bold=True, size_pt=11.5, color=COLOR_H1_BG)
    set_para_border_bottom(para)
    return para


def body(doc, text, indent=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    if indent:
        para.paragraph_format.left_indent = Cm(0.5)
    font(para, text, size_pt=10)
    return para


def make_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, COLOR_H1_BG)
        if col_widths:
            cell.width = col_widths[j]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        font(p, h, bold=True, size_pt=9, color=COLOR_WHITE)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i+1, j)
            if i % 2 == 1:
                set_cell_bg(cell, COLOR_ALT_ROW)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if val in ("완료", "정상", "✓", "설치완료", "가동중"):
                font(p, val, bold=True, size_pt=9, color=COLOR_OK)
            else:
                font(p, val, size_pt=9)
    doc.add_paragraph()
    return tbl


def add_title_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_TITLE_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    font(p, "시스템전환결과서", bold=True, size_pt=24, color=COLOR_WHITE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(20)
    font(p2, "System Transfer Result (SF-TI2)", size_pt=12, color=RGBColor(0xD6, 0xE4, 0xF0))
    doc.add_paragraph()
    meta = [
        ("문서번호", "SF-TI2-2026-001"),
        ("양식번호", "SF-TI2"),
        ("프로젝트명", "㈜임진강김치 선도형 스마트공장 구축"),
        ("과제번호", "SF26173364"),
        ("도입기업", "㈜임진강김치"),
        ("공급기업", "로뎀솔루션 주식회사"),
        ("설치 장소", "경기도 파주시 임진강김치 공장"),
        ("설치 기간", "2026-04-28 ~ 2026-05-09"),
        ("작성일", "2026-05-14"),
        ("버전", "V1.0"),
        ("작성자", "권영근 (개발팀장)"),
        ("승인자", "강정복 (대표이사)"),
    ]
    tbl2 = doc.add_table(rows=len(meta), cols=2)
    tbl2.style = 'Table Grid'
    for i, (k, v) in enumerate(meta):
        c0, c1 = tbl2.cell(i, 0), tbl2.cell(i, 1)
        set_cell_bg(c0, COLOR_H2_BG)
        c0.width = Cm(4)
        c1.width = Cm(10)
        for cell, txt, bd in [(c0, k, True), (c1, v, False)]:
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            font(p, txt, bold=bd, size_pt=10)
    doc.add_page_break()


def add_hw(doc):
    h1(doc, 1, "H/W 설치 결과")
    h2(doc, "1.1 서버 설치")
    make_table(doc,
        ["장비명", "모델", "수량", "설치 위치", "IP 주소", "결과"],
        [
            ("MES 서버", "Dell PowerEdge R350", "1대", "서버실 랙 1U", "192.168.10.10", "설치완료"),
            ("UPS", "APC Smart-UPS 1000VA", "1대", "서버실", "-", "설치완료"),
        ])
    h2(doc, "1.2 현장 터치PC 설치")
    make_table(doc,
        ["설치 공정", "모델", "IP 주소", "설치 위치", "결과"],
        [
            ("세척공정", "산업용 15인치 Touch PC", "192.168.10.51", "세척라인 작업대 우측", "설치완료"),
            ("절임공정", "산업용 15인치 Touch PC", "192.168.10.52", "절임조 관리실 입구", "설치완료"),
            ("양념공정", "산업용 15인치 Touch PC", "192.168.10.53", "버무림기 조작반 옆", "설치완료"),
            ("포장공정", "산업용 15인치 Touch PC", "192.168.10.54", "포장라인 컨베이어 끝", "설치완료"),
            ("품질관리", "산업용 15인치 Touch PC", "192.168.10.55", "품질검사실 검사대", "설치완료"),
        ])
    h2(doc, "1.3 IoT 센서 설치")
    make_table(doc,
        ["센서 종류", "수량", "설치 위치", "통신방식", "결과"],
        [
            ("온습도 센서", "1개", "절임실 (냉장 온도 모니터링)", "RS485 → 서버 수집", "설치완료"),
            ("온습도 센서", "1개", "숙성냉장실 (발효 온도 관리)", "RS485 → 서버 수집", "설치완료"),
            ("온습도 센서", "1개", "완제품 냉장창고", "RS485 → 서버 수집", "설치완료"),
        ])
    h2(doc, "1.4 자동화 설비 연동")
    make_table(doc,
        ["설비명", "연동방식", "설치 위치", "결과"],
        [
            ("자동테이핑기", "I/O 카운터 신호 → PLC → MES", "포장라인 출구", "설치완료"),
            ("금속검출기", "합격/불합격 신호 → MES 자동기록", "포장라인 금속검출 구간", "설치완료"),
        ])


def add_sw(doc):
    h1(doc, 2, "S/W 설치 결과")
    h2(doc, "2.1 서버 소프트웨어 설치 현황")
    make_table(doc,
        ["소프트웨어", "버전", "설치 경로", "라이선스", "결과"],
        [
            ("Windows Server 2022", "Standard", "C:\\", "정품", "설치완료"),
            ("Docker Desktop", "24.0.7", "C:\\Program Files\\Docker", "무료(CE)", "설치완료"),
            ("PostgreSQL", "15.4", "Docker 컨테이너 (port 5432)", "오픈소스", "설치완료"),
            ("FastAPI (Backend)", "0.104", "Docker 컨테이너 (port 8000)", "오픈소스", "설치완료"),
            ("Next.js (Frontend)", "14.1", "Docker 컨테이너 (port 3000)", "오픈소스", "설치완료"),
            ("Nginx (Reverse Proxy)", "1.25", "Docker 컨테이너 (port 80/443)", "오픈소스", "설치완료"),
        ])
    h2(doc, "2.2 Docker 컨테이너 배포 현황")
    body(doc, "docker-compose up -d 명령으로 전체 서비스 배포 완료")
    make_table(doc,
        ["서비스명", "컨테이너명", "이미지", "포트", "상태"],
        [
            ("DB", "mes_postgres", "postgres:15-alpine", "5432", "가동중"),
            ("Backend API", "mes_backend", "mes-backend:1.0.0", "8000", "가동중"),
            ("Frontend", "mes_frontend", "mes-frontend:1.0.0", "3000", "가동중"),
            ("Nginx", "mes_nginx", "nginx:1.25-alpine", "80, 443", "가동중"),
            ("Redis (세션)", "mes_redis", "redis:7-alpine", "6379", "가동중"),
        ])


def add_nw(doc):
    h1(doc, 3, "N/W 구성 결과")
    h2(doc, "3.1 IP 체계")
    make_table(doc,
        ["장비", "IP 주소", "서브넷", "용도"],
        [
            ("MES 서버", "192.168.10.10", "255.255.255.0", "메인 서버 (DB/API/Web)"),
            ("세척공정 PC", "192.168.10.51", "255.255.255.0", "POP 화면"),
            ("절임공정 PC", "192.168.10.52", "255.255.255.0", "POP 화면"),
            ("양념공정 PC", "192.168.10.53", "255.255.255.0", "POP 화면"),
            ("포장공정 PC", "192.168.10.54", "255.255.255.0", "POP 화면"),
            ("품질관리 PC", "192.168.10.55", "255.255.255.0", "POP 화면"),
            ("사무실 PC", "192.168.10.100~109", "255.255.255.0", "업무용 (수주/생산계획/KPI)"),
        ])
    h2(doc, "3.2 방화벽 설정")
    make_table(doc,
        ["포트", "프로토콜", "허용 대상", "용도"],
        [
            ("80", "TCP", "내부망 전체", "HTTP → HTTPS 리다이렉트"),
            ("443", "TCP", "내부망 전체", "HTTPS MES 웹 접속"),
            ("8000", "TCP", "192.168.10.x 내부망", "FastAPI 직접 접속 (개발용)"),
            ("5432", "TCP", "192.168.10.10 only", "PostgreSQL (서버 내부만)"),
        ])


def add_initial_data(doc):
    h1(doc, 4, "초기데이터 구축 결과")
    h2(doc, "4.1 기준정보 입력 현황")
    make_table(doc,
        ["데이터 구분", "목표 건수", "입력 완료", "비고"],
        [
            ("제품 마스터", "20건", "20건", "포기김치 외 4종 × 4가지 용량"),
            ("원자재 마스터", "30건", "30건", "배추/고추/마늘 등 주요 원부자재"),
            ("BOM", "20건", "20건", "제품별 배합 레시피"),
            ("작업장 마스터", "7개", "7개", "세척/절임/양념/포장/품질/출하/창고"),
            ("작업자 마스터", "15명", "15명", "공정별 현장 인력"),
            ("설비 마스터", "10대", "10대", "버무림기/포장기/테이핑기 등"),
            ("CCP 기준값", "8개 항목", "8개 항목", "절임염도/절임시간/pH/금속검출 등"),
            ("거래처", "25개", "25개", "납품처 및 원자재 공급처"),
        ])
    h2(doc, "4.2 시드 데이터 현황")
    body(doc, "시스템 테스트 및 사용자 교육을 위한 샘플 데이터 입력 완료")
    make_table(doc,
        ["테이블", "데이터 건수", "용도"],
        [
            ("수주 데이터", "20건", "KPI 차트 및 통계 테스트용"),
            ("생산실적", "20건", "대시보드 생산 추이 표시용"),
            ("재고 현황", "30건", "재고 경고 기능 테스트용"),
            ("품질검사 이력", "20건", "불량률 KPI 테스트용"),
            ("작업지시 이력", "20건", "LOT 추적 테스트용"),
        ])


def add_checklist(doc):
    h1(doc, 5, "전환 완료 체크리스트")
    make_table(doc,
        ["점검 항목", "담당자", "완료 여부", "비고"],
        [
            ("서버 H/W 설치 및 전원 확인", "권영근", "완료", ""),
            ("현장 터치PC 5대 설치", "권영근", "완료", "IP 설정 완료"),
            ("IoT 센서 3개소 설치 및 통신 확인", "오개발", "완료", "RS485 통신 확인"),
            ("자동테이핑기 연동 신호 테스트", "오개발", "완료", "카운터 신호 정상"),
            ("금속검출기 연동 신호 테스트", "오개발", "완료", "합격/불합격 신호 정상"),
            ("OS 및 미들웨어 설치", "권영근", "완료", ""),
            ("Docker 컨테이너 전체 기동", "권영근", "완료", "5개 서비스 정상"),
            ("DB 초기화 및 스키마 생성", "권영근", "완료", "44개 테이블"),
            ("기준정보 마스터 데이터 입력", "박PM", "완료", ""),
            ("네트워크 IP 체계 설정", "권영근", "완료", ""),
            ("방화벽 포트 설정", "권영근", "완료", ""),
            ("전체 서비스 정상 가동 확인", "성PM", "완료", ""),
            ("현장 PC 접속 테스트", "박PM", "완료", "5개 PC 모두 정상"),
            ("백업 스케줄 설정", "권영근", "완료", "매일 새벽 2시 자동"),
            ("시스템 전환 완료 확인서 수령", "성PM", "완료", "도입기업 서명 완료"),
        ])
    doc.add_paragraph()
    body(doc, "시스템 전환 완료 확인")
    body(doc, "")
    body(doc, "위의 설치 및 전환 결과를 확인합니다.")
    doc.add_paragraph()
    sign_rows = [
        ("도입기업 확인", "㈜임진강김치    강정복 대표이사    (인)", "2026-05-14"),
        ("공급기업 확인", "로뎀솔루션 주식회사    권영근 대표    (인)", "2026-05-14"),
    ]
    tbl = doc.add_table(rows=len(sign_rows), cols=3)
    tbl.style = 'Table Grid'
    for i, row in enumerate(sign_rows):
        for j, val in enumerate(row):
            p = tbl.cell(i, j).paragraphs[0]
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            font(p, val, size_pt=10)
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    add_title_page(doc)
    add_hw(doc)
    add_sw(doc)
    add_nw(doc)
    add_initial_data(doc)
    add_checklist(doc)
    doc.save(OUTPUT_PATH)
    print(f"저장 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
