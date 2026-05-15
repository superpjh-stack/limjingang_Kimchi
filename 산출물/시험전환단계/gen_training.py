# -*- coding: utf-8 -*-
"""임진강김치 MES - 교육실시관리서 생성 스크립트"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\gerardo\01 SmallSF\Imjingang-Kimchi\산출물\시험전환단계\14_교육실시관리서.docx"

COLOR_TITLE_BG = RGBColor(0x1F, 0x49, 0x7D)
COLOR_H1_BG    = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H2_BG    = RGBColor(0xD6, 0xE4, 0xF0)
COLOR_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_ALT_ROW  = RGBColor(0xEA, 0xF2, 0xFB)
COLOR_OK       = RGBColor(0x00, 0x70, 0xC0)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(rgb))
    tcPr.append(shd)


def set_para_border_bottom(para, color="2E74B5", sz=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def font(para, text, bold=False, size_pt=10.5, color=None, fn="맑은 고딕"):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = fn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    if color:
        run.font.color.rgb = color
    return run


def h1(doc, number, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_H1_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    font(p, f"{number}. {text}", bold=True, size_pt=13, color=COLOR_WHITE)
    doc.add_paragraph()


def h2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    font(para, text, bold=True, size_pt=11.5, color=COLOR_H1_BG)
    set_para_border_bottom(para)
    return para


def body(doc, text, indent=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    if indent:
        para.paragraph_format.left_indent = Cm(0.5)
    font(para, text, size_pt=10)
    return para


def make_table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, COLOR_H1_BG)
        if widths:
            cell.width = widths[j]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        font(p, h, bold=True, size_pt=9, color=COLOR_WHITE)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i+1, j)
            if i % 2 == 1:
                set_cell_bg(cell, COLOR_ALT_ROW)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if val in ("이수", "완료"):
                font(p, val, bold=True, size_pt=9, color=COLOR_OK)
            else:
                font(p, val, size_pt=9)
    doc.add_paragraph()
    return tbl


def add_title_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_TITLE_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    font(p, "교육실시관리서", bold=True, size_pt=24, color=COLOR_WHITE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(20)
    font(p2, "Training Management Report", size_pt=12, color=RGBColor(0xD6, 0xE4, 0xF0))
    doc.add_paragraph()
    meta = [
        ("문서번호", "TRN-2026-001"),
        ("프로젝트명", "㈜임진강김치 선도형 스마트공장 구축"),
        ("교육 기간", "2026-05-12 ~ 2026-05-13 (2일)"),
        ("교육 장소", "㈜임진강김치 1층 회의실 및 생산현장"),
        ("도입기업", "㈜임진강김치"),
        ("교육 주관", "로뎀솔루션 주식회사"),
        ("작성일", "2026-05-14"),
        ("버전", "V1.0"),
        ("작성자", "성PM"),
        ("승인자", "강정복 (대표이사)"),
    ]
    tbl2 = doc.add_table(rows=len(meta), cols=2)
    tbl2.style = 'Table Grid'
    for i, (k, v) in enumerate(meta):
        c0, c1 = tbl2.cell(i, 0), tbl2.cell(i, 1)
        set_cell_bg(c0, COLOR_H2_BG)
        c0.width = Cm(4)
        c1.width = Cm(10)
        for cell, txt, bd in [(c0, k, True), (c1, v, False)]:
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            font(p, txt, bold=bd, size_pt=10)
    doc.add_page_break()


def add_plan(doc):
    h1(doc, 1, "교육 계획 개요")
    h2(doc, "1.1 교육 목표")
    goals = [
        "MES 시스템의 기본 사용법을 습득하여 업무에 즉시 적용 가능하도록 한다.",
        "현장 작업자: POP 화면으로 공정 실적을 정확하게 입력할 수 있도록 한다.",
        "사무직: 수주/생산계획/재고/KPI 등 주요 업무 화면을 독립적으로 운영할 수 있도록 한다.",
        "관리자: 시스템 관리(사용자/기준정보/KPI설정)를 자립적으로 수행할 수 있도록 한다.",
        "AI Agent를 통해 실시간 현황을 자연어로 조회하는 방법을 숙지한다.",
    ]
    for g in goals:
        body(doc, f"• {g}", indent=True)

    h2(doc, "1.2 교육 대상")
    make_table(doc,
        ["대상 그룹", "인원", "교육 과목"],
        [
            ("현장 작업자 (세척/절임/양념/포장)", "8명", "POP 화면 실습, 공정실적 입력, CCP 판정 확인"),
            ("품질관리 담당", "2명", "품질검사 입력, CCP 이력 조회, HACCP 증빙"),
            ("생산·재고 관리 사무직", "3명", "수주/생산계획/재고/KPI 대시보드"),
            ("경영진·관리자", "2명", "KPI 대시보드, 시스템관리, AI Agent"),
        ])


def add_schedule(doc):
    h1(doc, 2, "교육 과정표")
    h2(doc, "2.1 1일차 (2026-05-12, 수요일)")
    make_table(doc,
        ["시간", "과정명", "주요 내용", "강사", "대상"],
        [
            ("09:00~09:30", "시스템 소개", "임진강김치 MES 구축 목적, 전체 메뉴 구조 안내", "성PM", "전체"),
            ("09:30~10:00", "로그인 및 기본 조작", "계정 로그인, 화면 레이아웃, 메뉴 이동", "권영근", "전체"),
            ("10:00~12:00", "대시보드 사용법", "KPI 카드, 생산추이/수주추이 차트, AI Agent 기본 질의", "권영근", "사무직·관리자"),
            ("10:00~12:00", "POP 화면 소개", "POP 화면 접속, WO 선택, 작업시작/종료 방법", "오개발", "현장작업자"),
            ("13:00~15:00", "수주·생산계획 실습", "수주 등록, 생산계획 작성, 작업지시 자동생성 확인", "성PM", "생산·재고 사무직"),
            ("13:00~15:00", "POP 실습 (세척/절임)", "세척 실적 입력, 절임 염도 CCP 입력, 기준 이탈 알림 확인", "오개발", "세척/절임 작업자"),
            ("15:00~17:00", "재고관리 실습", "원자재 입고 등록, 재고 조회, 안전재고 경고 확인", "권영근", "재고 관리"),
            ("15:00~17:00", "POP 실습 (양념/포장)", "배합 BOM 확인, 양념 실적 입력, 포장 자동 집계 확인", "오개발", "양념/포장 작업자"),
        ])
    h2(doc, "2.2 2일차 (2026-05-13, 목요일)")
    make_table(doc,
        ["시간", "과정명", "주요 내용", "강사", "대상"],
        [
            ("09:00~10:30", "품질검사 실습", "검사 결과 입력, CCP 판정, HACCP 이력 조회 및 출력", "권영근", "품질관리"),
            ("09:00~10:30", "LOT 추적 실습", "입고~출하 전 이력 추적, LOT 추적 화면 활용", "오개발", "전체"),
            ("10:30~12:00", "AI Agent 심화", "자연어 질의 예시 실습, 알림 확인, 현황 조회 활용법", "권영근", "전체"),
            ("13:00~14:30", "시스템 관리 (관리자 과정)", "사용자 등록/수정, 기준정보 관리, KPI 기준값 설정", "성PM", "관리자"),
            ("14:30~16:00", "KPI 리포트 활용", "KPI 대시보드 해석, 경영 의사결정 활용 방법", "성PM", "경영진·관리자"),
            ("16:00~17:00", "Q&A 및 마무리", "교육 종합 질의응답, 장애 시 연락처 안내, 평가", "성PM", "전체"),
        ])


def add_attendees(doc):
    h1(doc, 3, "교육 참석자 명단")
    make_table(doc,
        ["No", "성명", "직위", "부서", "이수 구분", "서명"],
        [
            ("1", "강미금", "이사", "경영관리", "이수", ""),
            ("2", "김영철", "반장", "생산관리", "이수", ""),
            ("3", "이현숙", "기능원", "세척/절임공정", "이수", ""),
            ("4", "박지훈", "기능원", "양념/포장공정", "이수", ""),
            ("5", "최수진", "검사원", "품질관리", "이수", ""),
            ("6", "정태호", "사원", "재고관리", "이수", ""),
            ("7", "한승민", "기능원", "세척/절임공정", "이수", ""),
            ("8", "오미선", "기능원", "양념/포장공정", "이수", ""),
            ("9", "윤대현", "기능원", "포장/출하", "이수", ""),
            ("10", "장수빈", "사원", "생산관리", "이수", ""),
            ("11", "임채원", "사원", "재고관리", "이수", ""),
            ("12", "신동호", "기능원", "세척공정", "이수", ""),
            ("13", "고미래", "검사원", "품질관리", "이수", ""),
            ("14", "배성훈", "반장", "포장/출하공정", "이수", ""),
            ("15", "전지현", "사원", "경영관리", "이수", ""),
        ])


def add_content_summary(doc):
    h1(doc, 4, "교육 내용 요약")
    make_table(doc,
        ["과목명", "주요 전달 내용", "실습 여부"],
        [
            ("시스템 소개", "MES 목적, 구축 효과(생산량 증가, 불량률 감소), 전체 메뉴 구조", "X"),
            ("로그인/기본조작", "계정 로그인, 역할별 메뉴, 비밀번호 변경", "O"),
            ("대시보드", "KPI 4개 카드, 생산/수주 추이 차트 해석, 재고경고 확인", "O"),
            ("POP 현장화면", "WO 선택, 작업시작/종료, 공정별 데이터 입력, CCP 기준 확인", "O"),
            ("수주관리", "수주 등록, 상태변경, 납기일 관리", "O"),
            ("생산계획", "계획 등록, WO 자동생성, 교대조 배정", "O"),
            ("재고관리", "입고 등록, 재고 조회, 안전재고 경고, LOT 추적", "O"),
            ("품질검사", "검사결과 입력, CCP 판정, HACCP 이력 조회", "O"),
            ("AI Agent", "자연어 질의, 현황 조회, 알림 확인", "O"),
            ("시스템관리(관리자)", "사용자 등록/수정, 기준정보 관리, KPI 설정", "O"),
        ])


def add_evaluation(doc):
    h1(doc, 5, "교육 결과 및 평가")
    h2(doc, "5.1 수료 현황")
    make_table(doc,
        ["항목", "내용"],
        [
            ("교육 대상 인원", "15명"),
            ("실제 참석 인원", "15명 (100%)"),
            ("이수 인원", "15명 (100%)"),
            ("미이수", "0명"),
        ])
    h2(doc, "5.2 교육 종합 의견")
    body(doc, "교육 참석자 전원이 2일간 교육을 이수하였습니다.")
    body(doc, "주요 피드백:")
    feedbacks = [
        "현장 작업자 전원이 POP 화면 기본 조작을 2시간 내 습득",
        "품질검사 CCP 자동판정 기능에 대한 만족도 높음",
        "AI Agent 자연어 조회 기능에 대한 흥미도 및 활용 의지 높음",
        "관리자 교육(시스템관리) 심화 교육 추가 요청 (차기 예정)",
    ]
    for fb in feedbacks:
        body(doc, f"  • {fb}", indent=True)
    doc.add_paragraph()
    body(doc, "교육 완료 서명")
    doc.add_paragraph()
    sign_data = [
        ("교육 주관자", "로뎀솔루션 주식회사    성PM    (인)", "2026-05-13"),
        ("도입기업 확인", "㈜임진강김치    강정복 대표이사    (인)", "2026-05-13"),
    ]
    tbl = doc.add_table(rows=len(sign_data), cols=3)
    tbl.style = 'Table Grid'
    for i, row in enumerate(sign_data):
        for j, val in enumerate(row):
            p = tbl.cell(i, j).paragraphs[0]
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)
            font(p, val, size_pt=10)
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    add_title_page(doc)
    add_plan(doc)
    add_schedule(doc)
    add_attendees(doc)
    add_content_summary(doc)
    add_evaluation(doc)
    doc.save(OUTPUT_PATH)
    print(f"저장 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
