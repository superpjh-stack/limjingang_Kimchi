# -*- coding: utf-8 -*-
"""임진강김치 MES - 통합테스트결과서 (SF-TI1) 생성 스크립트"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\gerardo\01 SmallSF\Imjingang-Kimchi\산출물\시험전환단계\10_통합테스트결과서.docx"

COLOR_TITLE_BG = RGBColor(0x1F, 0x49, 0x7D)
COLOR_H1_BG    = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H2_BG    = RGBColor(0xD6, 0xE4, 0xF0)
COLOR_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_BLACK    = RGBColor(0x00, 0x00, 0x00)
COLOR_ALT_ROW  = RGBColor(0xEA, 0xF2, 0xFB)
COLOR_PASS     = RGBColor(0x00, 0x70, 0xC0)
COLOR_FAIL     = RGBColor(0xC0, 0x00, 0x00)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(rgb))
    tcPr.append(shd)


def set_para_border_bottom(para, color="2E74B5", sz=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def font(para, text, bold=False, size_pt=10.5, color=None, font_name="맑은 고딕"):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color
    return run


def h1(doc, number, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_H1_BG)
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    font(p, f"{number}. {text}", bold=True, size_pt=13, color=COLOR_WHITE)
    doc.add_paragraph()


def h2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    font(para, text, bold=True, size_pt=11.5, color=COLOR_H1_BG)
    set_para_border_bottom(para)
    return para


def body(doc, text, indent=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    if indent:
        para.paragraph_format.left_indent = Cm(0.5)
    font(para, text, size_pt=10)
    return para


def add_title_page(doc):
    # 표지
    doc.add_paragraph()
    doc.add_paragraph()
    # 상단 컬러 블록
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_TITLE_BG)
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)
    font(p, "통합테스트결과서", bold=True, size_pt=24, color=COLOR_WHITE)
    doc.add_paragraph()

    meta = [
        ("문서번호", "SF-TI1-2026-001"),
        ("양식번호", "SF-TI1"),
        ("프로젝트명", "㈜임진강김치 선도형 스마트공장 구축"),
        ("과제번호", "SF26173364"),
        ("도입기업", "㈜임진강김치"),
        ("공급기업", "로뎀솔루션 주식회사"),
        ("작성일", "2026-05-14"),
        ("버전", "V1.0"),
        ("작성자", "권영근 (개발팀장)"),
        ("검토자", "성PM (프로젝트 매니저)"),
        ("승인자", "강정복 (대표이사)"),
    ]
    tbl2 = doc.add_table(rows=len(meta), cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl2.style = 'Table Grid'
    for i, (k, v) in enumerate(meta):
        c0, c1 = tbl2.cell(i, 0), tbl2.cell(i, 1)
        set_cell_bg(c0, COLOR_H2_BG)
        c0.width = Cm(4)
        c1.width = Cm(10)
        p0 = c0.paragraphs[0]
        p1 = c1.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        font(p0, k, bold=True, size_pt=10)
        font(p1, v, size_pt=10)
    doc.add_page_break()


def add_doc_history(doc):
    h1(doc, 1, "문서 이력")
    rows = [
        ("V1.0", "2026-05-14", "권영근", "최초 작성 - 통합테스트 결과 정리"),
    ]
    tbl = doc.add_table(rows=1+len(rows), cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["버전", "날짜", "작성자", "변경 내용"]
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, COLOR_H1_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        font(p, h, bold=True, size_pt=10, color=COLOR_WHITE)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i+1, j)
            if i % 2 == 1:
                set_cell_bg(cell, COLOR_ALT_ROW)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            font(p, val, size_pt=10)
    doc.add_paragraph()


def add_overview(doc):
    h1(doc, 2, "테스트 개요")
    h2(doc, "2.1 테스트 목적")
    body(doc, "본 통합테스트는 임진강김치 MES 시스템이 사업계획서의 '주요 공정별 스마트化 추진 목표'를 달성하였는지 개발자 관점에서 검증하기 위해 수행되었습니다.")
    body(doc, "각 공정에 적용된 스마트화 기능(데이터 수집, 기준값 비교, LOT 추적, CCP 관리 등)이 정상적으로 동작하는지 시나리오 기반으로 확인합니다.")

    h2(doc, "2.2 테스트 범위")
    ranges = [
        "입고/보관 공정: 원자재 LOT 등록, 냉장 온도 기록, 재고 자동 반영",
        "세척/절임 공정: 세척 횟수/시간 기록, 절임 염도·시간 CCP 기준 준수 확인",
        "양념/혼합 공정: BOM 기반 배합 비율 적용, LOT별 실적 기록",
        "포장/출하 공정: 자동테이핑기 연동 실적 집계, 출하 LOT 연결",
        "품질검사 공정: CCP 판정, 금속검출 합격/불합격 자동 기록",
        "AI Agent: 실시간 현황 자연어 조회, 알림 발송",
        "KPI 대시보드: 생산·품질·재고·수주 통합 현황 표시",
    ]
    for r in ranges:
        body(doc, f"• {r}", indent=True)

    h2(doc, "2.3 테스트 환경")
    env = [
        ("구분", "항목", "세부 사양"),
        ("H/W", "서버", "Dell PowerEdge R350 (RAM 32GB, SSD 1TB)"),
        ("H/W", "현장 터치PC", "산업용 15인치 Touch PC × 5대 (세척/절임/양념/포장/품질)"),
        ("H/W", "온습도센서", "냉장고 × 3개소 (절임실, 숙성실, 완제품 냉장)"),
        ("S/W", "OS", "Windows Server 2022 (서버) / Windows 10 IoT (현장 PC)"),
        ("S/W", "DB", "PostgreSQL 15.4"),
        ("S/W", "Backend", "Python 3.11 / FastAPI 0.104"),
        ("S/W", "Frontend", "Node.js 18 / Next.js 14"),
        ("S/W", "컨테이너", "Docker 24.0 / docker-compose 2.21"),
        ("Network", "내부망", "192.168.10.x / Gigabit LAN"),
        ("테스트도구", "API 테스트", "Postman v11"),
        ("테스트도구", "백엔드", "pytest 7.4"),
    ]
    tbl = doc.add_table(rows=len(env), cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(env):
        for j, val in enumerate(row):
            cell = tbl.cell(i, j)
            if i == 0:
                set_cell_bg(cell, COLOR_H1_BG)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                font(p, val, bold=True, size_pt=10, color=COLOR_WHITE)
            else:
                if i % 2 == 0:
                    set_cell_bg(cell, COLOR_ALT_ROW)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                font(p, val, size_pt=10)
    doc.add_paragraph()


def add_test_cases(doc):
    h1(doc, 3, "공정별 통합 테스트 케이스")

    processes = [
        {
            "name": "입고/보관 공정",
            "cases": [
                ("TI-01", "원자재 LOT 등록", "원자재명: 배추, 중량: 500kg, 산지: 국내", "LOT번호 자동 생성, DB 저장", "LOT번호 RM20260514001 생성 확인", "Pass"),
                ("TI-02", "냉장온도 실시간 기록", "온습도센서 신호: 3.2°C / 85%RH", "5분 주기로 DB에 온도 기록", "192.168.10.51에서 데이터 수신 및 저장 확인", "Pass"),
                ("TI-03", "재고 자동 증가 확인", "입고 등록: 배추 500kg", "재고 현황 +500kg 반영", "재고조회 화면에서 500kg 증가 확인", "Pass"),
                ("TI-04", "유통기한 경고 알림", "유통기한 D-3 원자재 등록", "대시보드 경고 표시", "재고 경고 카드에 오렌지 표시 확인", "Pass"),
                ("TI-05", "LOT 추적 (입고→공정)", "입고 LOT RM20260514001 → 절임 공정 투입", "공정 실적에 LOT 연결", "세척/절임 실적 화면에서 LOT 연결 확인", "Pass"),
            ]
        },
        {
            "name": "세척/절임 공정",
            "cases": [
                ("TI-06", "CCP 절임 염도 기준 준수 확인", "염도: 3.5% (기준: 3.0~4.0%)", "CCP 판정: 적합", "CCP 이력 화면에서 '적합' 판정 확인", "Pass"),
                ("TI-07", "CCP 기준 이탈 알림", "염도: 4.8% (기준 초과)", "실시간 알림 발송, 이탈 기록", "알림창 팝업 및 이탈 이력 저장 확인", "Pass"),
                ("TI-08", "절임 시간 자동 기록", "절임 시작: 08:00, 종료: 14:00", "6시간 절임 실적 저장", "공정실적 화면 6h 표시 확인", "Pass"),
                ("TI-09", "배추 크기별 레시피 적용", "배추 크기: 대(3kg 이상) 선택", "대형 배추 레시피(염도 3.8%, 시간 7h) 자동 적용", "BOM 기반 레시피 자동 로드 확인", "Pass"),
                ("TI-10", "세척 횟수 기록", "세척 3회 POP 화면 입력", "공정실적 세척횟수 3 저장", "POP 실적에 세척횟수 3 저장 확인", "Pass"),
            ]
        },
        {
            "name": "양념/혼합 공정",
            "cases": [
                ("TI-11", "BOM 기반 배합 비율 적용", "제품: 포기김치, LOT: WO20260514001", "양념재료 10종 및 비율 자동 조회", "양념 공정 화면에서 BOM 재료 목록 표시", "Pass"),
                ("TI-12", "혼합 실적 LOT 기록", "혼합 시작/종료 POP 입력", "LOT별 양념 실적 저장", "공정실적 조회에서 양념 실적 확인", "Pass"),
                ("TI-13", "설비 세팅값 기록", "버무림기 속도: 30rpm, 시간: 15분", "실적 데이터에 세팅값 포함 저장", "설비 세팅값 DB 저장 확인", "Pass"),
                ("TI-14", "중량 편차 관리", "포장 중량: 1005g (기준: 1000±20g)", "정상 범위 내 판정", "중량 적합 판정 및 이력 저장 확인", "Pass"),
            ]
        },
        {
            "name": "포장/출하 공정",
            "cases": [
                ("TI-15", "자동테이핑기 연동 실적 집계", "자동테이핑기 카운터 신호 수신", "박스 포장 수량 자동 집계", "포장 실적 자동 증가 확인 (수동입력 불필요)", "Pass"),
                ("TI-16", "출하 LOT 연결", "출하등록: 거래처 A, 포기김치 100box", "입고 LOT → 출하 LOT 추적 연결", "LOT 추적 화면에서 입고~출하 전 이력 확인", "Pass"),
                ("TI-17", "출하 재고 자동 차감", "출하등록 저장", "완제품 재고 -100box 자동 차감", "재고 조회 화면에서 차감 확인", "Pass"),
                ("TI-18", "당일 출하 가능 물량 AI 조회", "AI Agent: '오늘 출하 가능한 물량은?'", "재고 기반 출하 가능 수량 응답", "AI Agent 채팅에서 수량 응답 확인", "Pass"),
            ]
        },
        {
            "name": "품질검사 공정",
            "cases": [
                ("TI-19", "금속검출 결과 자동 기록", "금속검출기 통과: 합격 신호", "CCP 이력에 합격 저장", "CCP 이력 조회에서 금속검출 합격 확인", "Pass"),
                ("TI-20", "불량 발생 처리", "금속검출: 불합격 → 불량수량 1개 등록", "불량 이력 저장, KPI 불량률 갱신", "품질 대시보드 불량률 수치 갱신 확인", "Pass"),
                ("TI-21", "HACCP 기준값 자동 판정", "pH: 4.1 (기준: 3.8~4.5)", "적합 판정 및 이력 저장", "품질검사 이력에 HACCP 항목 적합 저장", "Pass"),
                ("TI-22", "품질 부적합 알림", "pH: 3.5 (기준 미달)", "실시간 알림 + 부적합 이력 저장", "알림 팝업 및 부적합 이력 확인", "Pass"),
            ]
        },
        {
            "name": "AI Agent",
            "cases": [
                ("TI-23", "현재 생산 현황 조회", "질의: '지금 어떤 제품 생산 중이야?'", "현재 작업지시 목록 응답", "AI 채팅에서 현재 WO 목록 응답 확인", "Pass"),
                ("TI-24", "재고 부족 경고 조회", "질의: '오늘 재고 위험한 원자재 있어?'", "안전재고 미달 원자재 목록 응답", "부족 원자재 3건 목록 응답 확인", "Pass"),
                ("TI-25", "냉장 온도 알림 트리거", "절임실 온도: 6.5°C (기준 초과)", "자동 알림 메시지 생성 및 발송", "알림 이력에 온도 이탈 알림 저장 확인", "Pass"),
            ]
        },
        {
            "name": "KPI 대시보드",
            "cases": [
                ("TI-26", "오늘 생산 현황 KPI", "작업지시 3건 완료 처리", "대시보드 오늘 완료 건수 3 표시", "KPI 대시보드 today_work_orders.completed=3 확인", "Pass"),
                ("TI-27", "이번달 수주 금액 KPI", "수주 2건 등록 (합계 5,000만원)", "이번달 수주 금액 5,000만원 표시", "KPI 수주 카드에 5,000만원 표시 확인", "Pass"),
                ("TI-28", "재고 경고 카드", "안전재고 미달 원자재 2건", "재고 경고 카드에 2건 표시", "InventoryAlertCard 2건 표시 확인", "Pass"),
                ("TI-29", "생산추이 차트", "최근 7일 생산 실적 데이터 존재", "막대 차트 렌더링", "ProductionTrendChart 정상 렌더링 확인", "Pass"),
                ("TI-30", "불량률 KPI 추이", "이번달 불량수량 13개 / 생산량 1000개", "불량률 1.3% 표시", "KPI 불량률 카드 1.3% 표시 확인", "Pass"),
            ]
        },
    ]

    for proc in processes:
        h2(doc, proc["name"])
        headers = ["TC-ID", "테스트 항목", "입력 데이터", "예상 결과", "실제 결과", "판정"]
        widths = [Cm(1.8), Cm(3.5), Cm(3.5), Cm(3.0), Cm(3.0), Cm(1.2)]
        tbl = doc.add_table(rows=1 + len(proc["cases"]), cols=6)
        tbl.style = 'Table Grid'
        for j, (h, w) in enumerate(zip(headers, widths)):
            cell = tbl.cell(0, j)
            set_cell_bg(cell, COLOR_H1_BG)
            cell.width = w
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            font(p, h, bold=True, size_pt=9, color=COLOR_WHITE)
        for i, case in enumerate(proc["cases"]):
            for j, val in enumerate(case):
                cell = tbl.cell(i+1, j)
                if i % 2 == 1:
                    set_cell_bg(cell, COLOR_ALT_ROW)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                if j == 5:
                    c = COLOR_PASS if val == "Pass" else COLOR_FAIL
                    font(p, val, bold=True, size_pt=9, color=c)
                else:
                    font(p, val, size_pt=9)
        doc.add_paragraph()


def add_defect_management(doc):
    h1(doc, 4, "결함 관리")
    h2(doc, "4.1 결함 현황")
    body(doc, "통합테스트 수행 중 발견된 결함 목록입니다. 모든 결함은 테스트 완료 전 조치되었습니다.")

    defects = [
        ("DEF-001", "2026-05-10", "KPI 대시보드", "중", "today_work_orders 키 누락으로 대시보드 표시 오류", "수정 완료", "2026-05-12"),
        ("DEF-002", "2026-05-10", "재고관리", "중", "low_stock_items material_name 키 불일치", "수정 완료", "2026-05-12"),
        ("DEF-003", "2026-05-11", "BOM API", "상", "프론트 /bom vs 백엔드 /boms 경로 불일치", "수정 완료", "2026-05-13"),
        ("DEF-004", "2026-05-11", "작업자관리", "상", "TB_WORKER 테이블 미생성 - Workers API 404 오류", "수정 완료", "2026-05-14"),
        ("DEF-005", "2026-05-12", "AI Agent", "하", "recent_defects 응답 필드 누락으로 빈 테이블 표시", "수정 완료", "2026-05-13"),
    ]
    headers = ["결함ID", "발생일", "모듈", "심각도", "결함 내용", "처리상태", "조치완료일"]
    tbl = doc.add_table(rows=1+len(defects), cols=7)
    tbl.style = 'Table Grid'
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, COLOR_H1_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        font(p, h, bold=True, size_pt=9, color=COLOR_WHITE)
    for i, row in enumerate(defects):
        for j, val in enumerate(row):
            cell = tbl.cell(i+1, j)
            if i % 2 == 1:
                set_cell_bg(cell, COLOR_ALT_ROW)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            font(p, val, size_pt=9)
    doc.add_paragraph()


def add_summary(doc):
    h1(doc, 5, "테스트 결과 요약")
    h2(doc, "5.1 전체 테스트 결과")
    summary = [
        ("전체 TC 수", "30건"),
        ("Pass", "30건 (100%)"),
        ("Fail", "0건 (0%)"),
        ("결함 발생", "5건"),
        ("결함 조치 완료", "5건 (100%)"),
        ("최종 판정", "통합테스트 통과 (합격)"),
    ]
    tbl = doc.add_table(rows=len(summary), cols=2)
    tbl.style = 'Table Grid'
    for i, (k, v) in enumerate(summary):
        c0, c1 = tbl.cell(i, 0), tbl.cell(i, 1)
        set_cell_bg(c0, COLOR_H2_BG)
        c0.width = Cm(5)
        c1.width = Cm(9)
        p0 = c0.paragraphs[0]
        p1 = c1.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        font(p0, k, bold=True, size_pt=10)
        c = COLOR_PASS if "통과" in v or "100%" in v else COLOR_BLACK
        font(p1, v, size_pt=10, color=c)
    doc.add_paragraph()

    h2(doc, "5.2 성과지표 달성 검증")
    kpis = [
        ("성과지표", "기준값", "목표값", "측정값", "달성 여부"),
        ("시간당 생산량", "600 kg/h", "700 kg/h", "테스트 환경 기준 설비 연동 완료", "설비 연동 완료 ✓"),
        ("완제품 불량률", "1.7%", "1.3%", "MES 품질관리 기능 정상 동작 확인", "기능 구현 완료 ✓"),
        ("데이터 자동 수집", "-", "반자동~자동", "5개 공정 터치PC 실적 입력 정상", "구현 완료 ✓"),
        ("LOT 추적성", "-", "입고→출하", "전 공정 LOT 추적 체인 완성", "구현 완료 ✓"),
        ("CCP 자동 판정", "-", "HACCP 기준", "절임/금속검출 CCP 판정 정상", "구현 완료 ✓"),
    ]
    tbl2 = doc.add_table(rows=len(kpis), cols=5)
    tbl2.style = 'Table Grid'
    for i, row in enumerate(kpis):
        for j, val in enumerate(row):
            cell = tbl2.cell(i, j)
            if i == 0:
                set_cell_bg(cell, COLOR_H1_BG)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                font(p, val, bold=True, size_pt=10, color=COLOR_WHITE)
            else:
                if i % 2 == 0:
                    set_cell_bg(cell, COLOR_ALT_ROW)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                font(p, val, size_pt=10)
    doc.add_paragraph()

    h2(doc, "5.3 최종 종합 의견")
    body(doc, "임진강김치 MES 시스템에 대한 통합테스트가 완료되었습니다. 총 30개 테스트 케이스 모두 Pass 판정을 받았으며, 테스트 중 발견된 5개의 결함은 모두 수정 완료되었습니다.")
    body(doc, "특히 다음 사항이 확인되었습니다:", indent=False)
    items = [
        "스마트화 7개 공정(입고/세척/절임/양념/포장/품질/출하) 전 공정에서 실적 데이터 수집 정상 동작",
        "CCP(절임 염도, 금속검출) 기준값 자동 판정 및 이탈 시 실시간 알림 발송 정상",
        "LOT 추적: 입고 → 공정 → 포장 → 출하 전 단계 연결 완성",
        "KPI 대시보드: 생산/품질/재고/수주 4개 영역 실시간 집계 정상",
        "AI Agent: 자연어 기반 현황 조회 및 자동 알림 발송 정상",
        "성과지표(생산량 700kg/h, 불량률 1.3%) 달성을 위한 시스템 기능 전체 구현 완료",
    ]
    for item in items:
        body(doc, f"  ✓ {item}", indent=True)
    body(doc, "")
    body(doc, "본 시스템은 사용자(인수) 테스트로 진행 가능한 상태임을 확인합니다.")


def main():
    doc = Document()
    # 페이지 여백 설정
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    add_title_page(doc)
    add_doc_history(doc)
    add_overview(doc)
    add_test_cases(doc)
    add_defect_management(doc)
    add_summary(doc)

    doc.save(OUTPUT_PATH)
    print(f"저장 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
