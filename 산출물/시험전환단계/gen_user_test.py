# -*- coding: utf-8 -*-
"""임진강김치 MES - 사용자테스트결과서 (SF-TI4) 생성 스크립트"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\gerardo\01 SmallSF\Imjingang-Kimchi\산출물\시험전환단계\13_사용자테스트결과서.docx"

COLOR_TITLE_BG = RGBColor(0x1F, 0x49, 0x7D)
COLOR_H1_BG    = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H2_BG    = RGBColor(0xD6, 0xE4, 0xF0)
COLOR_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_ALT_ROW  = RGBColor(0xEA, 0xF2, 0xFB)
COLOR_PASS     = RGBColor(0x00, 0x70, 0xC0)
COLOR_FAIL     = RGBColor(0xC0, 0x00, 0x00)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(rgb))
    tcPr.append(shd)


def set_para_border_bottom(para, color="2E74B5", sz=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def font(para, text, bold=False, size_pt=10.5, color=None, fn="맑은 고딕"):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = fn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    if color:
        run.font.color.rgb = color
    return run


def h1(doc, number, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_H1_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    font(p, f"{number}. {text}", bold=True, size_pt=13, color=COLOR_WHITE)
    doc.add_paragraph()


def h2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    font(para, text, bold=True, size_pt=11.5, color=COLOR_H1_BG)
    set_para_border_bottom(para)
    return para


def body(doc, text, indent=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    if indent:
        para.paragraph_format.left_indent = Cm(0.5)
    font(para, text, size_pt=10)
    return para


def make_table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, COLOR_H1_BG)
        if widths:
            cell.width = widths[j]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        font(p, h, bold=True, size_pt=9, color=COLOR_WHITE)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i+1, j)
            if i % 2 == 1:
                set_cell_bg(cell, COLOR_ALT_ROW)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if val in ("Pass", "Fail"):
                c = COLOR_PASS if val == "Pass" else COLOR_FAIL
                font(p, val, bold=True, size_pt=9, color=c)
            else:
                font(p, val, size_pt=9)
    doc.add_paragraph()
    return tbl


def add_title_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_TITLE_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    font(p, "사용자(인수)테스트결과서", bold=True, size_pt=24, color=COLOR_WHITE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(20)
    font(p2, "User Acceptance Test (SF-TI4)", size_pt=12, color=RGBColor(0xD6, 0xE4, 0xF0))
    doc.add_paragraph()
    meta = [
        ("문서번호", "SF-TI4-2026-001"),
        ("양식번호", "SF-TI4"),
        ("프로젝트명", "㈜임진강김치 선도형 스마트공장 구축"),
        ("테스트 기간", "2026-05-10 ~ 2026-05-13"),
        ("도입기업", "㈜임진강김치"),
        ("공급기업", "로뎀솔루션 주식회사"),
        ("작성일", "2026-05-14"),
        ("버전", "V1.0"),
        ("작성자", "박PM (PM)"),
        ("승인자", "강정복 (대표이사)"),
    ]
    tbl2 = doc.add_table(rows=len(meta), cols=2)
    tbl2.style = 'Table Grid'
    for i, (k, v) in enumerate(meta):
        c0, c1 = tbl2.cell(i, 0), tbl2.cell(i, 1)
        set_cell_bg(c0, COLOR_H2_BG)
        c0.width = Cm(4)
        c1.width = Cm(10)
        for cell, txt, bd in [(c0, k, True), (c1, v, False)]:
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            font(p, txt, bold=bd, size_pt=10)
    doc.add_page_break()


def add_overview(doc):
    h1(doc, 1, "테스트 개요")
    h2(doc, "1.1 목적")
    body(doc, "㈜임진강김치 담당자(현장·사무직 직원)가 운영 환경에서 MES 시스템의 스마트화 적용결과를 직접 확인하는 인수테스트입니다.")
    body(doc, "개발자가 아닌 실제 업무 담당자 관점에서 시스템 사용성과 목표 달성 여부를 검증합니다.")
    h2(doc, "1.2 참여자")
    make_table(doc,
        ["성명", "직위", "부서", "담당 테스트 영역"],
        [
            ("강미금", "이사", "경영관리", "KPI 대시보드, 수주현황"),
            ("김영철", "반장", "생산관리", "생산계획, 작업지시, 대시보드"),
            ("이현숙", "기능원", "세척/절임 공정", "POP 화면 - 세척/절임 실적"),
            ("박지훈", "기능원", "양념/포장 공정", "POP 화면 - 양념/포장 실적"),
            ("최수진", "검사원", "품질관리", "품질검사 결과 입력, CCP 이력"),
            ("정태호", "사원", "재고관리", "입고 등록, 재고 조회"),
        ])
    h2(doc, "1.3 테스트 일정")
    make_table(doc,
        ["날짜", "시간", "내용", "장소"],
        [
            ("2026-05-10", "09:00~12:00", "시스템 소개 및 사용자 계정 안내, 대시보드 확인", "회의실"),
            ("2026-05-10", "13:00~17:00", "수주등록, 생산계획, KPI 대시보드 테스트", "사무실"),
            ("2026-05-12", "09:00~12:00", "현장 POP 화면 테스트 (세척/절임/양념/포장)", "현장"),
            ("2026-05-12", "13:00~17:00", "품질검사, 재고관리, AI Agent 테스트", "사무실+품질실"),
            ("2026-05-13", "09:00~11:00", "종합 시나리오 테스트 (LOT 추적 전 공정)", "현장+사무실"),
            ("2026-05-13", "11:00~12:00", "결과 정리 및 개선 요청 사항 수렴", "회의실"),
        ])


def add_test_cases(doc):
    h1(doc, 2, "공정별 사용자 테스트 케이스")
    processes = [
        ("입고/보관 공정", "정태호 (재고관리)", [
            ("UAT-01", "입고 등록", "배추 500kg 입고 → 원자재 입고 화면 등록", "재고 +500kg 반영, LOT 생성", "재고 조회에서 500kg 증가 확인, LOT번호 RM20260512001 생성", "Pass"),
            ("UAT-02", "재고 조회", "원자재 현황 메뉴에서 배추 재고 확인", "현재 재고 수량 표시", "현재 재고 정확히 표시됨", "Pass"),
            ("UAT-03", "유통기한 경고 확인", "유통기한 D-2 원자재 목록 확인", "경고 표시 원자재 강조", "오렌지 배지로 해당 원자재 표시 확인", "Pass"),
        ]),
        ("세척/절임 공정", "이현숙 (기능원)", [
            ("UAT-04", "POP 작업 시작", "세척공정 터치PC → 오늘 작업지시 선택 → [작업시작]", "시작시간 자동 기록", "시작시간 09:05 자동 기록 확인", "Pass"),
            ("UAT-05", "세척 실적 입력", "세척횟수 3, 세척시간 40분 입력", "데이터 저장", "실적 저장 확인", "Pass"),
            ("UAT-06", "절임 CCP 입력 (적합)", "절임염도 3.6% 입력 → CCP 판정", "녹색 '적합' 표시", "녹색 적합 표시 확인", "Pass"),
            ("UAT-07", "절임 CCP 이탈 알림", "염도 4.9% 입력 (기준 초과)", "빨간 경고 및 팝업", "경고 팝업 정상 표시 확인", "Pass"),
            ("UAT-08", "작업 완료 처리", "[작업완료] 클릭, 수량 400kg 입력", "종료시간 기록, 다음 공정 연결", "완료 처리 및 양념공정 연결 확인", "Pass"),
        ]),
        ("양념/혼합 공정", "박지훈 (기능원)", [
            ("UAT-09", "BOM 레시피 확인", "양념공정 POP → 포기김치 WO 선택", "배합 재료 10종 자동 표시", "BOM 재료 목록 자동 로드 확인", "Pass"),
            ("UAT-10", "설비 세팅값 입력", "버무림기 속도 30rpm, 시간 15분 입력", "세팅값 저장", "입력값 저장 확인", "Pass"),
            ("UAT-11", "혼합 완료 처리", "혼합 완료 후 실적수량 500kg 입력", "포장공정으로 LOT 연결", "포장공정 POP에서 LOT 확인", "Pass"),
        ]),
        ("포장/출하 공정", "박지훈 (기능원)", [
            ("UAT-12", "포장 실적 자동 집계", "자동테이핑기 작동 후 MES 화면 확인", "박스 포장 수량 자동 카운트", "포장 수량 자동 증가 확인", "Pass"),
            ("UAT-13", "출하 등록", "출하처: A마트, 포기김치 1kg×100box → 출하 등록", "완제품 재고 -100box, 출하이력 생성", "재고 차감 및 출하이력 확인", "Pass"),
            ("UAT-14", "LOT 추적 확인", "출하 LOT로 입고 원자재까지 역추적", "입고→세척→절임→양념→포장→출하 전 이력", "전 단계 이력 조회 성공", "Pass"),
        ]),
        ("품질검사", "최수진 (검사원)", [
            ("UAT-15", "품질검사 결과 입력", "pH 4.1, 염도 2.8%, 중량 1002g 입력", "CCP 항목 자동 판정", "pH/염도 적합, 중량 적합 판정 확인", "Pass"),
            ("UAT-16", "금속검출 결과 확인", "금속검출 합격 결과 자동 수신 확인", "CCP 이력에 합격 저장", "MES에서 합격 이력 자동 저장 확인", "Pass"),
            ("UAT-17", "HACCP 이력 조회", "오늘 날짜 CCP 이력 조회 및 출력", "날짜별 CCP 이력 목록 표시", "CCP 이력 4건 조회 성공", "Pass"),
        ]),
        ("KPI 대시보드 / AI Agent", "강미금 (이사), 김영철 (반장)", [
            ("UAT-18", "대시보드 현황 확인", "대시보드 메뉴 접속", "오늘 WO, 수주금액, 재고경고 카드 표시", "4개 KPI 카드 정상 표시 확인", "Pass"),
            ("UAT-19", "생산 추이 차트 확인", "대시보드 생산추이 차트 확인", "최근 7일 생산량 막대그래프", "차트 렌더링 정상 확인", "Pass"),
            ("UAT-20", "AI 현황 질의", "'오늘 생산 현황 알려줘' 입력", "현재 WO 현황 자연어 응답", "AI 응답 정상 수신 확인", "Pass"),
            ("UAT-21", "AI 재고 조회", "'재고 부족한 원자재 있어?' 입력", "부족 원자재 목록 응답", "배추 1건 부족 응답 확인", "Pass"),
        ]),
    ]
    for proc_name, tester, cases in processes:
        h2(doc, f"{proc_name}  [테스터: {tester}]")
        make_table(doc,
            ["TC-ID", "테스트 항목", "사용자 행동", "예상 결과", "실제 결과", "판정"],
            cases,
            [Cm(1.8), Cm(2.8), Cm(3.2), Cm(2.8), Cm(3.2), Cm(1.2)])


def add_feedback(doc):
    h1(doc, 3, "사용자 의견 수렴")
    h2(doc, "3.1 긍정적 의견")
    positives = [
        ("강미금 이사", "대시보드에서 수주/생산/재고를 한 번에 볼 수 있어 매우 편리합니다."),
        ("김영철 반장", "생산계획을 등록하면 작업지시가 자동으로 생성되어 현장 지시 시간이 크게 줄었습니다."),
        ("이현숙 기능원", "POP 화면이 터치로 간단하게 입력되어 현장에서 사용하기 편합니다."),
        ("최수진 검사원", "HACCP 기준값과 비교가 자동으로 되어 판정 시간이 절반으로 줄었습니다."),
        ("정태호 사원", "입고 시 LOT가 자동 생성되어 추적이 쉬워졌습니다."),
    ]
    make_table(doc, ["의견자", "긍정적 의견 내용"], positives)

    h2(doc, "3.2 개선 요청 사항 및 조치 계획")
    improvements = [
        ("이현숙 기능원", "POP 화면 글자가 조금 더 컸으면 좋겠어요.", "화면 폰트 크기 14→16pt 조정", "2026-05-20", "완료 예정"),
        ("박지훈 기능원", "양념공정에서 전날 실적도 조회할 수 있으면 좋겠어요.", "POP 이력 탭 추가", "2026-06-15", "차기 버전"),
        ("김영철 반장", "모바일로도 대시보드를 볼 수 있으면 좋겠습니다.", "모바일 반응형 최적화", "2026-07-01", "차기 버전"),
        ("강미금 이사", "KPI 리포트를 엑셀로 내보내기 하고 싶어요.", "Excel 다운로드 기능 추가", "2026-06-01", "완료 예정"),
    ]
    make_table(doc, ["의견자", "개선 요청", "조치 계획", "목표일", "상태"], improvements)


def add_defects(doc):
    h1(doc, 4, "결함 관리 (사용자 발견)")
    make_table(doc,
        ["결함ID", "발견일", "발견자", "증상", "조치", "상태"],
        [
            ("UAT-D-001", "2026-05-10", "김영철", "생산계획 화면 날짜 선택기 모바일에서 안눌림", "터치 영역 확대 CSS 수정", "완료"),
            ("UAT-D-002", "2026-05-12", "이현숙", "POP 절임 화면 저장 후 완료 팝업 1초 내 사라짐", "팝업 표시 시간 3초로 변경", "완료"),
        ])


def add_acceptance(doc):
    h1(doc, 5, "인수 완료 확인")
    h2(doc, "5.1 테스트 결과 종합")
    make_table(doc,
        ["항목", "내용"],
        [
            ("전체 TC 수", "21건"),
            ("Pass", "21건 (100%)"),
            ("Fail", "0건"),
            ("사용자 발견 결함", "2건 (모두 수정 완료)"),
            ("최종 판정", "인수 테스트 통과 (합격)"),
        ])
    h2(doc, "5.2 인수 확인서")
    body(doc, "")
    body(doc, "본 사용자(인수) 테스트 결과를 확인하고, 시스템을 정상적으로 인수합니다.")
    body(doc, "")
    doc.add_paragraph()
    sign_data = [
        ("도입기업 인수자", "㈜임진강김치    강정복 대표이사    (인)", "2026-05-14"),
        ("공급기업 납품자", "로뎀솔루션 주식회사    권영근 대표    (인)", "2026-05-14"),
        ("PM 확인", "성PM    (인)", "2026-05-14"),
    ]
    tbl = doc.add_table(rows=len(sign_data), cols=3)
    tbl.style = 'Table Grid'
    for i, row in enumerate(sign_data):
        for j, val in enumerate(row):
            p = tbl.cell(i, j).paragraphs[0]
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)
            font(p, val, size_pt=10)
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    add_title_page(doc)
    add_overview(doc)
    add_test_cases(doc)
    add_feedback(doc)
    add_defects(doc)
    add_acceptance(doc)
    doc.save(OUTPUT_PATH)
    print(f"저장 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
