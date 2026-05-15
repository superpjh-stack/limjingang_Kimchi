# -*- coding: utf-8 -*-
"""임진강김치 MES - 메뉴얼 (SF-TI3) 생성 스크립트"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:\gerardo\01 SmallSF\Imjingang-Kimchi\산출물\시험전환단계\12_메뉴얼.docx"

COLOR_TITLE_BG = RGBColor(0x1F, 0x49, 0x7D)
COLOR_H1_BG    = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H2_BG    = RGBColor(0xD6, 0xE4, 0xF0)
COLOR_H3_BG    = RGBColor(0xF2, 0xF7, 0xFF)
COLOR_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_ALT_ROW  = RGBColor(0xEA, 0xF2, 0xFB)
COLOR_TIP      = RGBColor(0x00, 0x70, 0xC0)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(rgb))
    tcPr.append(shd)


def set_para_border_bottom(para, color="2E74B5", sz=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def font(para, text, bold=False, size_pt=10.5, color=None, fn="맑은 고딕"):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = fn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    if color:
        run.font.color.rgb = color
    return run


def h1(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_H1_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    font(p, text, bold=True, size_pt=13, color=COLOR_WHITE)
    doc.add_paragraph()


def h2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    font(para, text, bold=True, size_pt=11.5, color=COLOR_H1_BG)
    set_para_border_bottom(para)
    return para


def h3(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(2)
    font(para, text, bold=True, size_pt=10.5, color=COLOR_H1_BG)
    return para


def body(doc, text, indent=0):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    if indent:
        para.paragraph_format.left_indent = Cm(indent)
    font(para, text, size_pt=10)
    return para


def tip(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_H3_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    font(p, f"Tip. {text}", size_pt=9.5, color=COLOR_TIP)
    doc.add_paragraph()


def make_table(doc, headers, rows):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, COLOR_H1_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        font(p, h, bold=True, size_pt=9, color=COLOR_WHITE)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i+1, j)
            if i % 2 == 1:
                set_cell_bg(cell, COLOR_ALT_ROW)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            font(p, val, size_pt=9)
    doc.add_paragraph()
    return tbl


def add_title_page(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, COLOR_TITLE_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    font(p, "임진강김치 MES 시스템", bold=True, size_pt=20, color=COLOR_WHITE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    font(p2, "사용자 · 운영자 매뉴얼", bold=True, size_pt=24, color=COLOR_WHITE)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(20)
    font(p3, "User & Admin Manual (SF-TI3)", size_pt=12, color=RGBColor(0xD6, 0xE4, 0xF0))
    doc.add_paragraph()
    meta = [
        ("문서번호", "SF-TI3-2026-001"),
        ("양식번호", "SF-TI3"),
        ("프로젝트명", "㈜임진강김치 선도형 스마트공장 구축"),
        ("시스템 URL", "http://192.168.10.10 (내부망)"),
        ("작성일", "2026-05-14"),
        ("버전", "V1.0"),
        ("작성자", "성PM"),
        ("배포 대상", "㈜임진강김치 전 직원"),
    ]
    tbl2 = doc.add_table(rows=len(meta), cols=2)
    tbl2.style = 'Table Grid'
    for i, (k, v) in enumerate(meta):
        c0, c1 = tbl2.cell(i, 0), tbl2.cell(i, 1)
        set_cell_bg(c0, COLOR_H2_BG)
        c0.width = Cm(4)
        c1.width = Cm(10)
        for cell, txt, bd in [(c0, k, True), (c1, v, False)]:
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            font(p, txt, bold=bd, size_pt=10)
    doc.add_page_break()


def add_part1(doc):
    h1(doc, "PART 1. 사용자 매뉴얼")

    h2(doc, "1.1 시스템 접속 방법")
    body(doc, "1. 인터넷 브라우저(Chrome 권장)를 실행합니다.")
    body(doc, "2. 주소창에 http://192.168.10.10 을 입력합니다.")
    body(doc, "3. 로그인 화면에서 지급받은 아이디와 비밀번호를 입력합니다.")
    body(doc, "4. [로그인] 버튼 클릭 시 대시보드로 이동합니다.")
    tip(doc, "비밀번호를 분실한 경우 관리자(내선 101)에게 초기화를 요청하세요.")
    make_table(doc,
        ["역할", "기본 아이디", "접근 가능 메뉴"],
        [
            ("현장작업자", "worker01~15", "POP 현장화면, 공정실적 조회"),
            ("생산관리", "prod01", "수주관리, 생산계획, 공정실적, 재고"),
            ("품질관리", "qc01", "품질검사, CCP이력, KPI"),
            ("경영진/관리자", "admin", "전체 메뉴 + 시스템관리"),
        ])

    h2(doc, "1.2 대시보드 사용법")
    h3(doc, "1.2.1 대시보드 구성")
    items = [
        "오늘 작업지시 현황: 전체/진행중/완료 건수를 카드로 표시",
        "이번 달 수주 현황: 수주건수와 수주금액 표시",
        "재고 경고: 안전재고 미달 원자재 목록 (빨간 배지)",
        "최근 불량 이력: 최근 5건 불량 LOT 정보",
        "생산 추이 차트: 최근 7일 생산량 막대그래프",
        "수주 추이 차트: 월별 수주 건수 꺾은선 그래프",
    ]
    for item in items:
        body(doc, f"• {item}", indent=0.5)
    tip(doc, "대시보드 데이터는 실시간으로 갱신됩니다. F5 또는 새로고침 불필요.")

    h2(doc, "1.3 수주 관리")
    h3(doc, "1.3.1 수주 등록")
    body(doc, "1. 좌측 메뉴 → [수주관리] → [수주 목록] 클릭")
    body(doc, "2. 우측 상단 [+ 수주 등록] 버튼 클릭")
    body(doc, "3. 거래처, 제품, 수량, 납기일을 입력합니다.")
    body(doc, "4. [저장] 버튼 클릭 시 수주번호가 자동 생성됩니다.")
    tip(doc, "납기일이 오늘로부터 3일 이내인 경우 오렌지색 경고가 표시됩니다.")
    h3(doc, "1.3.2 수주 상태 변경")
    make_table(doc,
        ["상태", "설명", "전환 조건"],
        [
            ("접수", "수주 등록 직후 상태", "자동"),
            ("생산중", "생산계획이 수립되어 작업 진행 중", "생산계획 등록 시 자동"),
            ("완료", "출하 처리 완료", "출하 등록 시 자동"),
            ("취소", "수주 취소 처리", "수동 (담당자 확인 필요)"),
        ])

    h2(doc, "1.4 생산계획 작성")
    body(doc, "1. [생산계획] → [계획 목록] → [+ 생산계획 등록]")
    body(doc, "2. 수주 건을 선택하면 제품과 수량이 자동 입력됩니다.")
    body(doc, "3. 계획 날짜, 작업 교대조(1조/2조), 담당 작업자를 선택합니다.")
    body(doc, "4. [승인] 버튼 클릭 시 작업지시(WO)가 자동 생성됩니다.")
    tip(doc, "BOM에 등록된 재료가 현재 재고에 부족한 경우 경고 메시지가 표시됩니다.")

    h2(doc, "1.5 POP 현장작업 화면")
    h3(doc, "1.5.1 접속 방법")
    body(doc, "현장 터치PC에서 자동으로 POP 화면이 실행됩니다. (자동 로그인 설정)")
    body(doc, "POP 화면 URL: http://192.168.10.10/pop/{공정ID}")
    h3(doc, "1.5.2 공정 실적 입력 순서")
    steps = [
        ("1단계", "작업지시(WO) 확인", "오늘 날짜 WO 목록에서 해당 WO 선택"),
        ("2단계", "작업 시작", "[작업 시작] 버튼 클릭 → 시작시간 자동 기록"),
        ("3단계", "공정 데이터 입력", "공정별 항목 입력 (세척: 횟수/시간, 절임: 염도/시간 등)"),
        ("4단계", "CCP 항목 확인", "기준값 범위 내: 녹색 / 이탈: 빨간색 경고"),
        ("5단계", "작업 종료", "[작업 완료] 버튼 클릭 → 종료시간, 실적수량 입력"),
        ("6단계", "저장", "[저장] 클릭 → 다음 공정으로 LOT 연결"),
    ]
    make_table(doc, ["단계", "항목", "설명"], steps)

    h2(doc, "1.6 재고 조회")
    body(doc, "• [재고관리] → [재고 현황]: 현재 모든 원자재/반제품/완제품 재고 조회")
    body(doc, "• 재고 경고 기준: 안전재고 미달 시 빨간 배지 표시")
    body(doc, "• 재고 이력: 입고/출고/이동 이력을 날짜별로 조회 가능")
    tip(doc, "유통기한이 3일 이내인 원자재는 목록에 오렌지색으로 표시됩니다.")

    h2(doc, "1.7 품질검사 결과 입력")
    body(doc, "1. [품질관리] → [검사 등록] → WO 선택")
    body(doc, "2. 검사 항목별 측정값 입력 (pH, 염도, 중량, 금속검출 등)")
    body(doc, "3. CCP 항목은 기준값과 자동 비교 → 적합/부적합 자동 판정")
    body(doc, "4. 부적합 발생 시 처리 방법(폐기/재작업) 선택 후 저장")
    tip(doc, "HACCP 증빙 자료는 [품질관리] → [CCP 이력 조회]에서 날짜별 출력 가능합니다.")

    h2(doc, "1.8 AI Agent 사용법")
    body(doc, "화면 우측 하단 채팅 버튼(💬)을 클릭하면 AI Agent 창이 열립니다.")
    h3(doc, "1.8.1 주요 질의 예시")
    make_table(doc,
        ["질의 예시", "응답 내용"],
        [
            ("오늘 생산 현황 알려줘", "현재 작업 중인 WO 목록과 진행률"),
            ("재고 부족한 원자재 있어?", "안전재고 미달 원자재 목록"),
            ("이번 달 수주 얼마야?", "이번달 수주건수와 금액"),
            ("지금 절임실 온도는?", "절임실 현재 온도 및 기준 대비 상태"),
            ("오늘 불량 발생했어?", "당일 불량 건수 및 LOT 번호"),
            ("포기김치 재고 몇 개야?", "현재 완제품 재고 수량"),
        ])


def add_part2(doc):
    doc.add_page_break()
    h1(doc, "PART 2. 운영자 매뉴얼")

    h2(doc, "2.1 사용자 및 권한 관리")
    body(doc, "1. [시스템관리] → [사용자관리] → [+ 사용자 등록]")
    body(doc, "2. 아이디, 이름, 부서, 권한 그룹 선택 후 저장")
    make_table(doc,
        ["권한 그룹", "메뉴 접근 범위"],
        [
            ("ADMIN", "전체 메뉴 (시스템관리 포함)"),
            ("MANAGER", "수주/생산/재고/품질/KPI (시스템관리 제외)"),
            ("WORKER", "POP 화면, 공정실적 조회만"),
            ("QC", "품질관리, CCP이력, KPI"),
        ])
    tip(doc, "신규 입사자 등록 시 임시 비밀번호(imjin1234!)로 초기 지급 후 첫 로그인 시 변경하도록 안내하세요.")

    h2(doc, "2.2 기준정보 관리")
    body(doc, "[기준정보관리] 메뉴에서 다음 마스터 데이터를 관리합니다:")
    items = [
        "제품 마스터: 생산하는 제품 목록 (포기김치/깍두기/총각김치 등)",
        "원자재 마스터: 배추/고추/마늘 등 원부자재 등록 및 안전재고 설정",
        "BOM: 제품별 레시피 (재료 및 배합 비율)",
        "작업장: 세척/절임/양념/포장/품질/출하/창고",
        "작업자: 현장 인력 등록, 교대조 배정",
        "설비: 버무림기/포장기/테이핑기 등 설비 목록",
    ]
    for item in items:
        body(doc, f"• {item}", indent=0.5)

    h2(doc, "2.3 Docker 서비스 관리")
    h3(doc, "2.3.1 서비스 상태 확인")
    body(doc, "서버 원격 접속 후 아래 명령어 실행:")
    body(doc, '  명령: docker ps', indent=0.5)
    body(doc, '  결과: 5개 컨테이너 모두 "Up" 상태 확인', indent=0.5)
    h3(doc, "2.3.2 서비스 재시작")
    body(doc, '  전체 재시작: docker compose restart', indent=0.5)
    body(doc, '  개별 재시작: docker compose restart backend', indent=0.5)
    h3(doc, "2.3.3 로그 확인")
    body(doc, '  백엔드 로그: docker compose logs backend --tail=100', indent=0.5)
    body(doc, '  DB 로그:     docker compose logs postgres --tail=50', indent=0.5)
    tip(doc, "서비스 이상 시 먼저 로그를 확인하고, 로뎀솔루션(031-XXX-XXXX)에 문의하세요.")

    h2(doc, "2.4 DB 백업 및 복구")
    h3(doc, "2.4.1 자동 백업")
    body(doc, "매일 새벽 2:00 자동 백업 실행 (cron 설정)")
    body(doc, "백업 경로: C:\\mes_backup\\YYYY-MM-DD_backup.sql")
    body(doc, "보관 기간: 최근 30일")
    h3(doc, "2.4.2 수동 백업")
    body(doc, '  명령: docker exec mes_postgres pg_dump -U mesadmin mes_db > backup.sql', indent=0.5)
    h3(doc, "2.4.3 복구 절차")
    body(doc, "1. 로뎀솔루션 기술지원팀에 먼저 연락 (함부로 복구 시도 금지)")
    body(doc, "2. 백업 파일 확인 후 psql 명령으로 복구")

    h2(doc, "2.5 장애 대응 절차 (FAQ)")
    make_table(doc,
        ["증상", "원인", "조치 방법"],
        [
            ("화면이 열리지 않음", "서비스 중단 또는 네트워크 오류", "docker ps 확인 → 서비스 재시작"),
            ("로그인 실패", "비밀번호 오류 또는 계정 잠김", "관리자 계정으로 비밀번호 초기화"),
            ("POP 화면 공백", "네트워크 연결 불량 또는 API 오류", "터치PC 네트워크 확인, 백엔드 로그 확인"),
            ("온도 데이터 미수신", "IoT 센서 연결 불량", "RS485 케이블 및 센서 전원 확인"),
            ("데이터 저장 안됨", "DB 연결 오류", "postgres 컨테이너 상태 확인"),
            ("화면 느림", "서버 부하 증가", "불필요한 Chrome 탭 닫기, 서버 메모리 확인"),
        ])

    h2(doc, "2.6 KPI 기준값 설정")
    body(doc, "[시스템관리] → [KPI 설정]에서 다음 기준값을 관리합니다:")
    make_table(doc,
        ["KPI 항목", "현재 기준값", "설정 방법"],
        [
            ("시간당 생산량 목표", "700 kg/h", "KPI 설정 → 생산 목표값 입력"),
            ("완제품 불량률 목표", "1.3%", "KPI 설정 → 품질 목표값 입력"),
            ("절임 염도 CCP 기준", "3.0 ~ 4.0 %", "기준정보 → CCP 기준값 수정"),
            ("절임 시간 CCP 기준", "최소 6시간", "기준정보 → CCP 기준값 수정"),
            ("냉장온도 경고 기준", "0 ~ 5 °C", "설비관리 → 온도 기준값 수정"),
            ("재고 경고 기준", "원자재별 안전재고 설정", "원자재 마스터 → 안전재고 수량 입력"),
        ])


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    add_title_page(doc)
    add_part1(doc)
    add_part2(doc)
    doc.save(OUTPUT_PATH)
    print(f"저장 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
